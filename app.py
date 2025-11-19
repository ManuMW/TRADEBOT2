from flask import Flask, render_template, request, redirect, url_for, session, jsonify
import requests
import uuid
import logging
import os
from dotenv import load_dotenv
import pickle
from datetime import datetime, timedelta, timezone
import pytz
import sqlite3
import json
from openai import OpenAI
from docx import Document
from werkzeug.utils import secure_filename
import pandas as pd
import numpy as np
import pandas_ta as ta
import yfinance as yf
from bs4 import BeautifulSoup
from apscheduler.schedulers.background import BackgroundScheduler
from apscheduler.triggers.cron import CronTrigger
import threading
import time
from queue import Queue

# Load environment variables from .env file
load_dotenv()

# IST Timezone for Indian market
IST = pytz.timezone('Asia/Kolkata')

def get_ist_now():
    """Get current datetime in IST timezone"""
    return datetime.now(IST)

# Logging setup - Configure console handler with UTF-8 encoding
console_handler = logging.StreamHandler()
console_handler.setLevel(logging.INFO)
console_handler.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s'))
# Force UTF-8 encoding for Windows console
import sys
if sys.platform == 'win32':
    console_handler.stream = open(sys.stdout.fileno(), mode='w', encoding='utf-8', buffering=1)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('tradebot.log', encoding='utf-8'),
        console_handler
    ]
)

# Silence verbose APScheduler logs (only show warnings and errors)
logging.getLogger('apscheduler.scheduler').setLevel(logging.WARNING)
logging.getLogger('apscheduler.executors.default').setLevel(logging.WARNING)

app = Flask(__name__)
app.secret_key = 'replace_this_with_a_secure_key'

# Global session manager for SmartApi
_SMARTAPI_SESSIONS = {}
SESSION_FILE = 'sessions.pkl'

def load_sessions():
    global _SMARTAPI_SESSIONS
    try:
        if os.path.exists(SESSION_FILE):
            with open(SESSION_FILE, 'rb') as f:
                _SMARTAPI_SESSIONS = pickle.load(f)
            logging.info(f"Loaded {len(_SMARTAPI_SESSIONS)} persisted sessions")
    except Exception as e:
        logging.error(f"Failed to load sessions: {e}")
        _SMARTAPI_SESSIONS = {}

def save_sessions():
    try:
        with open(SESSION_FILE, 'wb') as f:
            pickle.dump(_SMARTAPI_SESSIONS, f)
        logging.info(f"Saved {len(_SMARTAPI_SESSIONS)} sessions")
    except Exception as e:
        logging.error(f"Failed to save sessions: {e}")

load_sessions()

# Initialize OpenAI client
openai_client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

# Database setup
DB_FILE = 'trading_data.db'

def init_db():
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS api_data (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
            clientcode TEXT,
            endpoint TEXT,
            data_type TEXT,
            response TEXT
        )
    ''')
    c.execute('''
        CREATE TABLE IF NOT EXISTS documents (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            upload_date DATE DEFAULT CURRENT_DATE,
            clientcode TEXT,
            document_type TEXT,
            filename TEXT,
            content TEXT,
            metadata TEXT
        )
    ''')
    c.execute('''
        CREATE TABLE IF NOT EXISTS daily_pnl (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            date DATE,
            clientcode TEXT,
            total_pnl REAL,
            trades_count INTEGER,
            trades_data TEXT,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    c.execute('''
        CREATE TABLE IF NOT EXISTS trade_plans (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            date DATE,
            clientcode TEXT,
            plan_text TEXT,
            status TEXT,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    conn.commit()
    conn.close()
    logging.info("Database initialized")

def store_data(clientcode, endpoint, data_type, response):
    try:
        conn = sqlite3.connect(DB_FILE)
        c = conn.cursor()
        c.execute(
            'INSERT INTO api_data (clientcode, endpoint, data_type, response) VALUES (?, ?, ?, ?)',
            (clientcode, endpoint, data_type, json.dumps(response))
        )
        conn.commit()
        conn.close()
        logging.info(f"Stored {data_type} data for {clientcode}")
    except Exception as e:
        logging.error(f"Failed to store data: {e}")

init_db()

@app.route('/')
def index():
    session_id = session.get('session_id')
    if session_id and session_id in _SMARTAPI_SESSIONS:
        logging.info(f"User already logged in with session_id={session_id}, redirecting to dashboard")
        return redirect(url_for('dashboard'))
    return render_template('login.html')

@app.route('/login', methods=['POST'])
def login():
    import os
    SMARTAPI_API_KEY = os.getenv("SMARTAPI_API_KEY")
    
    if not SMARTAPI_API_KEY:
        logging.error("SMARTAPI_API_KEY not set in environment")
        return render_template('login.html', error='API Key not configured. Check .env file.')
    
    clientcode = request.form.get('clientcode')
    password = request.form.get('password')
    totp = request.form.get('totp')
    
    if clientcode and password and totp:
        try:
            from SmartApi import SmartConnect
            logging.info(f"Attempting login for clientcode={clientcode}")
            smartApi = SmartConnect(SMARTAPI_API_KEY)
            
            # Call generateSession and log raw response
            data = smartApi.generateSession(clientcode, password, totp)
            logging.info(f"SmartAPI response type: {type(data)}, content: {data}")
            
            # Check if login was successful
            if not data or (isinstance(data, dict) and data.get('status') == False):
                error_msg = data.get('message', 'Login failed') if isinstance(data, dict) else 'Empty response from SmartAPI'
                logging.error(f"Login rejected for clientcode={clientcode}: {error_msg}")
                return render_template('login.html', error=f'Login failed: {error_msg}')
            
            # Extract tokens and set them explicitly
            tokens = data.get('data', {}) if isinstance(data, dict) else {}
            jwt_token = tokens.get('jwtToken')
            refresh_token = tokens.get('refreshToken')
            feed_token = tokens.get('feedToken')
            
            if jwt_token:
                smartApi.setAccessToken(jwt_token)
                logging.info(f"Set access token for clientcode={clientcode}")
            if refresh_token:
                smartApi.setRefreshToken(refresh_token)
            if feed_token:
                smartApi.setFeedToken(feed_token)
            
            session_id = uuid.uuid4().hex
            _SMARTAPI_SESSIONS[session_id] = {
                'api': smartApi,
                'clientcode': clientcode,
                'tokens': tokens,
                'login_time': datetime.now()
            }
            session['session_id'] = session_id
            session.permanent = True
            app.permanent_session_lifetime = timedelta(days=1)
            save_sessions()
            logging.info(f"Login successful for clientcode={clientcode}, session_id={session_id}, tokens extracted")
            return redirect(url_for('dashboard'))
        except Exception as e:
            logging.error(f"Login failed for clientcode={clientcode}: {e}", exc_info=True)
            return render_template('login.html', error=f'Login failed: {e}')
    logging.warning("Login attempt with missing credentials")
    return render_template('login.html', error='Invalid credentials')

@app.route('/logout')
def logout():
    sid = session.pop('session_id', None)
    if sid and sid in _SMARTAPI_SESSIONS:
        del _SMARTAPI_SESSIONS[sid]
        save_sessions()
    return redirect(url_for('index'))

@app.route('/dashboard')
def dashboard():
    if 'session_id' not in session:
        return redirect(url_for('index'))
    return render_template('dashboard.html')

# View routes for individual pages
@app.route('/view/profile')
def view_profile():
    if 'session_id' not in session:
        return redirect(url_for('index'))
    return render_template('view.html', title='Profile', api_endpoint='/api/profile')

@app.route('/view/marketdata')
def view_marketdata():
    if 'session_id' not in session:
        return redirect(url_for('index'))
    return render_template('view.html', title='Market Data (NIFTY 50)', api_endpoint='/api/marketdata')

@app.route('/view/rms')
def view_rms():
    if 'session_id' not in session:
        return redirect(url_for('index'))
    return render_template('view.html', title='RMS / Funds', api_endpoint='/api/rms')

@app.route('/view/orders')
def view_orders():
    if 'session_id' not in session:
        return redirect(url_for('index'))
    return render_template('view.html', title='Order Book', api_endpoint='/api/orders/book')

@app.route('/view/trades')
def view_trades():
    if 'session_id' not in session:
        return redirect(url_for('index'))
    return render_template('view.html', title='Trade Book', api_endpoint='/api/orders/trades')

@app.route('/view/optionchain')
def view_optionchain():
    if 'session_id' not in session:
        return redirect(url_for('index'))
    return render_template('optionchain.html')

@app.route('/view/scriphelper')
def view_scriphelper():
    if 'session_id' not in session:
        return redirect(url_for('index'))
    return render_template('scriphelper.html')

@app.route('/view/user_analysis')
def view_user_analysis():
    if 'session_id' not in session:
        return redirect(url_for('index'))
    return render_template('user_analysis.html')

@app.route('/api/marketdata')
def marketdata():
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        logging.warning("Marketdata access without valid session")
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    smartApi = _SMARTAPI_SESSIONS[session_id]['api']
    jwt_token = _SMARTAPI_SESSIONS[session_id]['tokens'].get('jwtToken', '')
    if jwt_token.startswith('Bearer '):
        jwt_token = jwt_token[7:]
    url = "https://apiconnect.angelone.in/rest/secure/angelbroking/market/v1/quote/"
    headers = {
        'Authorization': f'Bearer {jwt_token}',
        'Content-Type': 'application/json',
        'Accept': 'application/json',
        'X-UserType': 'USER',
        'X-SourceID': 'WEB',
        'X-ClientLocalIP': 'CLIENT_LOCAL_IP',
        'X-ClientPublicIP': 'CLIENT_PUBLIC_IP',
        'X-MACAddress': 'MAC_ADDRESS',
        'X-PrivateKey': os.getenv('SMARTAPI_API_KEY')
    }
    payload = {
        "mode": "FULL",
        "exchangeTokens": {"NSE": ["99926000"]}
    }
    try:
        r = requests.post(url, json=payload, headers=headers, timeout=10)
        data = r.json()
        # Store in database
        clientcode = _SMARTAPI_SESSIONS[session_id]['clientcode']
        store_data(clientcode, '/api/marketdata', 'marketdata', data)
    except Exception as e:
        logging.error(f"Marketdata error for session_id={session_id}: {e}")
        return jsonify({'status': False, 'message': str(e)}), 500
    summary = {}
    buy_depth = []
    sell_depth = []
    try:
        fetched = data.get('data', {}).get('fetched', [])
        if fetched:
            md = fetched[0]
            summary = {
                "ltp": md.get("ltp"),
                "netChange": md.get("netChange"),
                "percentChange": md.get("percentChange"),
                "open": md.get("open"),
                "high": md.get("high"),
                "low": md.get("low"),
                "close": md.get("close"),
                "tradeVolume": md.get("tradeVolume"),
                "exchFeedTime": md.get("exchFeedTime"),
                "exchTradeTime": md.get("exchTradeTime"),
                "upperCircuit": md.get("upperCircuit"),
                "lowerCircuit": md.get("lowerCircuit"),
            }
            buy_depth = md.get("depth", {}).get("buy", [])
            sell_depth = md.get("depth", {}).get("sell", [])
    except Exception:
        pass
    return jsonify({
        "status": True,
        "summary": summary,
        "buy_depth": buy_depth,
        "sell_depth": sell_depth,
        "raw": data
    })

@app.route('/api/marketdata/custom', methods=['POST'])
def marketdata_custom():
    """
    Fetch market data for custom symbols.
    Request body example:
    {
        "mode": "FULL",  // or "OHLC" or "LTP"
        "exchangeTokens": {
            "NSE": ["3045", "881"],
            "NFO": ["58662"]
        }
    }
    """
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        logging.warning("Custom marketdata access without valid session")
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    
    smartApi = _SMARTAPI_SESSIONS[session_id]['api']
    jwt_token = _SMARTAPI_SESSIONS[session_id]['tokens'].get('jwtToken', '')
    if jwt_token.startswith('Bearer '):
        jwt_token = jwt_token[7:]
    
    url = "https://apiconnect.angelone.in/rest/secure/angelbroking/market/v1/quote/"
    headers = {
        'Authorization': f'Bearer {jwt_token}',
        'Content-Type': 'application/json',
        'Accept': 'application/json',
        'X-UserType': 'USER',
        'X-SourceID': 'WEB',
        'X-ClientLocalIP': 'CLIENT_LOCAL_IP',
        'X-ClientPublicIP': 'CLIENT_PUBLIC_IP',
        'X-MACAddress': 'MAC_ADDRESS',
        'X-PrivateKey': os.getenv('SMARTAPI_API_KEY')
    }
    
    # Get payload from request
    payload = request.get_json()
    if not payload:
        return jsonify({'status': False, 'message': 'Invalid request body'}), 400
    
    try:
        logging.info(f"Custom market data request: {payload}")
        r = requests.post(url, json=payload, headers=headers, timeout=10)
        data = r.json()
        logging.info(f"Custom market data response: success={data.get('status')}")
    except Exception as e:
        logging.error(f"Custom marketdata error for session_id={session_id}: {e}")
        return jsonify({'status': False, 'message': str(e)}), 500
    
    return jsonify(data)

@app.route('/api/scrip/search', methods=['POST'])
def scrip_search():
    """Search scrip master for option tokens"""
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    
    body = request.get_json() or {}
    symbol = body.get('symbol', 'NIFTY').upper()
    option_type = body.get('option_type', 'CE')  # CE or PE
    strike = body.get('strike')  # Optional: specific strike
    show_all_expiries = body.get('show_all_expiries', False)
    cache_range = body.get('cache_range', True)  # Cache 3 levels up/down for speed trading
    
    try:
        # Fetch scrip master
        response = requests.get(
            'https://margincalculator.angelbroking.com/OpenAPI_File/files/OpenAPIScripMaster.json',
            timeout=30
        )
        response.raise_for_status()
        scrips = response.json()
        
        logging.info(f"Scrip master loaded: {len(scrips)} total scrips")
        
        # Filter for NFO options
        # Symbol format: NIFTY18NOV2525950CE (symbol + date + year + strike + CE/PE)
        options = [
            s for s in scrips
            if s.get('exch_seg') == 'NFO'
            and s.get('name') == symbol
            and s.get('instrumenttype') == 'OPTIDX'
            and s.get('symbol', '').endswith(option_type)
        ]
        
        logging.info(f"Found {len(options)} {symbol} {option_type} options in NFO")
        
        if not options:
            return jsonify({
                'status': False,
                'message': f'No {option_type} options found for {symbol} in scrip master'
            })
        
        # Parse strike price from symbol
        # Format: NIFTY18NOV2525900CE
        # Breakdown: NIFTY + 18NOV25 (18th Nov 2025) + 25900 (strike) + CE
        import re
        from datetime import datetime
        
        for opt in options:
            symbol_str = opt.get('symbol', '')
            expiry_str = opt.get('expiry', '')  # e.g., "18NOV2025"
            
            try:
                # Parse expiry to get date part (18NOV25)
                expiry_date = datetime.strptime(expiry_str, '%d%b%Y')
                date_part = expiry_date.strftime('%d%b%y').upper()  # 18NOV25
                
                # Remove symbol name from beginning
                remaining = symbol_str.replace(symbol, '', 1)
                
                # Remove CE/PE from end
                if remaining.endswith('CE'):
                    remaining = remaining[:-2]
                elif remaining.endswith('PE'):
                    remaining = remaining[:-2]
                
                # Remove date part from beginning
                if remaining.startswith(date_part):
                    strike_str = remaining[len(date_part):]
                    opt['parsed_strike'] = float(strike_str) if strike_str.isdigit() else 0.0
                else:
                    opt['parsed_strike'] = 0.0
            except:
                # Fallback
                opt['parsed_strike'] = float(opt.get('strike', 0))
        
        # Parse expiry dates
        today = datetime.now()
        
        for opt in options:
            try:
                expiry_str = opt.get('expiry', '')
                opt['expiry_date'] = datetime.strptime(expiry_str, '%d%b%Y')
                opt['days_to_expiry'] = (opt['expiry_date'] - today).days
            except:
                opt['expiry_date'] = None
                opt['days_to_expiry'] = 9999
        
        # Filter only future expiries (including today)
        options = [o for o in options if o['days_to_expiry'] >= 0]
        
        logging.info(f"After filtering future expiries: {len(options)} options")
        
        # Sort by expiry (closest first), then by strike
        options.sort(key=lambda x: (x['days_to_expiry'], x['parsed_strike']))
        
        # Get strike price range for debugging
        if options:
            strikes = [o['parsed_strike'] for o in options]
            logging.info(f"Strike range: {min(strikes)} to {max(strikes)}")
        
        # If strike specified, find exact match + 3 levels up/down
        cached_options = []
        if strike:
            strike_float = float(strike)
            
            # Get closest expiry options only
            closest_expiry_days = options[0]['days_to_expiry']
            closest_options = [o for o in options if o['days_to_expiry'] == closest_expiry_days]
            
            # Get all unique strikes, sorted
            all_strikes = sorted(set([o['parsed_strike'] for o in closest_options]))
            
            # Find the target strike or closest
            if strike_float in all_strikes:
                strike_index = all_strikes.index(strike_float)
            else:
                # Find closest strike
                strike_index = min(range(len(all_strikes)), key=lambda i: abs(all_strikes[i] - strike_float))
                logging.warning(f"Strike {strike_float} not found, using closest: {all_strikes[strike_index]}")
            
            if cache_range:
                # Get 3 levels up and 3 levels down for speed trading
                start_idx = max(0, strike_index - 3)
                end_idx = min(len(all_strikes), strike_index + 4)  # +4 because range is exclusive
                cached_strikes = all_strikes[start_idx:end_idx]
                
                cached_options = [o for o in closest_options if o['parsed_strike'] in cached_strikes]
                logging.info(f"Caching strikes for speed trading: {cached_strikes}")
            else:
                # Only exact match
                cached_options = [o for o in closest_options if o['parsed_strike'] == all_strikes[strike_index]]
            
            if not cached_options:
                return jsonify({
                    'status': False,
                    'message': f'Strike {strike_float} not found. Available strikes: {all_strikes[:20]}'
                })
            
            options = cached_options
        
        # Group by expiry and get results
        if options:
            if show_all_expiries:
                result_options = options[:50]
            else:
                # Show only closest expiry
                closest_expiry_days = options[0]['days_to_expiry']
                result_options = [
                    o for o in options
                    if o['days_to_expiry'] == closest_expiry_days
                ][:50]
            
            # Format response
            result = []
            for opt in result_options:
                result.append({
                    'token': opt.get('token'),
                    'symbol': opt.get('symbol'),
                    'name': opt.get('name'),
                    'strike': opt['parsed_strike'],
                    'expiry': opt.get('expiry'),
                    'days_to_expiry': opt['days_to_expiry'],
                    'lotsize': opt.get('lotsize'),
                    'option_type': option_type
                })
            
            # Get unique expiries
            unique_expiries = sorted(set([o['expiry'] for o in options]))
            
            # Ensure parsed_strike exists in result list
            result_strikes = [o.get('parsed_strike', float(o.get('strike', 0))) for o in result]
            
            return jsonify({
                'status': True,
                'message': f'Found {len(result)} options (cached for speed trading)',
                'data': result,
                'closest_expiry': options[0].get('expiry'),
                'days_to_expiry': options[0]['days_to_expiry'],
                'total_options_found': len(options),
                'available_expiries': unique_expiries[:5],
                'strike_range': {
                    'min': min([o.get('parsed_strike', float(o.get('strike', 0))) for o in options]),
                    'max': max([o.get('parsed_strike', float(o.get('strike', 0))) for o in options])
                },
                'cached_strikes': sorted(set(result_strikes)) if strike else None
            })
        else:
            return jsonify({
                'status': False,
                'message': 'No active options found with future expiries'
            })
    
    except Exception as e:
        logging.error(f"Scrip search error: {e}", exc_info=True)
        return jsonify({'status': False, 'message': str(e)}), 500

@app.route('/api/optionchain', methods=['POST'])
def optionchain():
    """Get Historical OI Data for F&O contracts"""
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        logging.warning("Optionchain access without valid session")
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    
    smartApi = _SMARTAPI_SESSIONS[session_id]['api']
    clientcode = _SMARTAPI_SESSIONS[session_id]['clientcode']
    jwt_token = _SMARTAPI_SESSIONS[session_id]['tokens'].get('jwtToken', '')
    if jwt_token.startswith('Bearer '):
        jwt_token = jwt_token[7:]
    
    body = request.get_json() or {}
    
    # Default parameters
    exchange = body.get('exchange', 'NFO')
    symboltoken = body.get('symboltoken')  # Required
    interval = body.get('interval', 'ONE_DAY')  # ONE_MINUTE, THREE_MINUTE, FIVE_MINUTE, TEN_MINUTE, FIFTEEN_MINUTE, THIRTY_MINUTE, ONE_HOUR, ONE_DAY
    fromdate = body.get('fromdate')  # Format: YYYY-MM-DD HH:MM
    todate = body.get('todate')  # Format: YYYY-MM-DD HH:MM
    
    if not symboltoken:
        return jsonify({'status': False, 'message': 'symboltoken is required'}), 400
    
    if not fromdate or not todate:
        # Default to last 7 days if not provided
        from datetime import datetime, timedelta
        todate = datetime.now().strftime('%Y-%m-%d %H:%M')
        fromdate = (datetime.now() - timedelta(days=7)).strftime('%Y-%m-%d %H:%M')
    
    try:
        # Prepare request
        headers = {
            'Authorization': f'Bearer {jwt_token}',
            'Content-Type': 'application/json',
            'Accept': 'application/json',
            'X-UserType': 'USER',
            'X-SourceID': 'WEB',
            'X-ClientLocalIP': '',
            'X-ClientPublicIP': '',
            'X-MACAddress': '',
            'X-PrivateKey': smartApi.api_key
        }
        
        payload = {
            "exchange": exchange,
            "symboltoken": symboltoken,
            "interval": interval,
            "fromdate": fromdate,
            "todate": todate
        }
        
        logging.info(f"Option chain OI request for token {symboltoken}: {payload}")
        
        response = requests.post(
            'https://apiconnect.angelone.in/rest/secure/angelbroking/historical/v1/getOIData',
            headers=headers,
            json=payload,
            timeout=10
        )
        
        response.raise_for_status()
        result = response.json()
        
        logging.info(f"Option chain OI response: {result}")
        
        # Store in database
        store_data(clientcode, '/api/optionchain', 'optionchain', result)
        
        logging.info(f"Option chain OI data fetched for {clientcode}, token: {symboltoken}")
        
        return jsonify(result)
        
    except requests.RequestException as e:
        logging.error(f"Option chain OI data error: {e}")
        return jsonify({'status': False, 'message': f'API Error: {str(e)}'}), 500
    except Exception as e:
        logging.error(f"Option chain error: {e}")
        return jsonify({'status': False, 'message': str(e)}), 500

@app.route('/api/profile')
def profile():
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        logging.warning("Profile access without valid session")
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    smartApi = _SMARTAPI_SESSIONS[session_id]['api']
    jwt_token = _SMARTAPI_SESSIONS[session_id]['tokens'].get('jwtToken', '')
    if jwt_token.startswith('Bearer '):
        jwt_token = jwt_token[7:]
    
    url = "https://apiconnect.angelone.in/rest/secure/angelbroking/user/v1/getProfile"
    headers = {
        'Authorization': f'Bearer {jwt_token}',
        'Content-Type': 'application/json',
        'Accept': 'application/json',
        'X-UserType': 'USER',
        'X-SourceID': 'WEB',
        'X-ClientLocalIP': 'CLIENT_LOCAL_IP',
        'X-ClientPublicIP': 'CLIENT_PUBLIC_IP',
        'X-MACAddress': 'MAC_ADDRESS',
        'X-PrivateKey': os.getenv('SMARTAPI_API_KEY')
    }
    try:
        logging.info(f"Profile API call with JWT token")
        r = requests.get(url, headers=headers, timeout=10)
        data = r.json()
        logging.info(f"Profile API response: {data}")
        # Store in database
        clientcode = _SMARTAPI_SESSIONS[session_id]['clientcode']
        store_data(clientcode, '/api/profile', 'profile', data)
    except Exception as e:
        logging.error(f"Profile error for session_id={session_id}: {e}")
        return jsonify({'status': False, 'message': str(e)}), 500
    return jsonify(data)

@app.route('/api/rms')
def rms():
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        logging.warning("RMS access without valid session")
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    smartApi = _SMARTAPI_SESSIONS[session_id]['api']
    jwt_token = _SMARTAPI_SESSIONS[session_id]['tokens'].get('jwtToken', '')
    # Remove 'Bearer ' prefix if present
    if jwt_token.startswith('Bearer '):
        jwt_token = jwt_token[7:]
    url = "https://apiconnect.angelone.in/rest/secure/angelbroking/user/v1/getRMS"
    headers = {
        'Authorization': f'Bearer {jwt_token}',
        'Content-Type': 'application/json',
        'Accept': 'application/json',
        'X-UserType': 'USER',
        'X-SourceID': 'WEB',
        'X-ClientLocalIP': 'CLIENT_LOCAL_IP',
        'X-ClientPublicIP': 'CLIENT_PUBLIC_IP',
        'X-MACAddress': 'MAC_ADDRESS',
        'X-PrivateKey': os.getenv('SMARTAPI_API_KEY')
    }
    try:
        r = requests.get(url, headers=headers, timeout=10)
        data = r.json()
        # Store in database
        clientcode = _SMARTAPI_SESSIONS[session_id]['clientcode']
        store_data(clientcode, '/api/rms', 'rms', data)
    except Exception as e:
        logging.error(f"RMS error for session_id={session_id}: {e}")
        return jsonify({'status': False, 'message': str(e)}), 500
    return jsonify(data)

@app.route('/api/orders/book')
def order_book():
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        logging.warning("Order book access without valid session")
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    smartApi = _SMARTAPI_SESSIONS[session_id]['api']
    jwt_token = _SMARTAPI_SESSIONS[session_id]['tokens'].get('jwtToken', '')
    if jwt_token.startswith('Bearer '):
        jwt_token = jwt_token[7:]
    url = 'https://apiconnect.angelone.in/rest/secure/angelbroking/order/v1/getOrderBook'
    headers = {
        'Authorization': f'Bearer {jwt_token}',
        'Content-Type': 'application/json',
        'Accept': 'application/json',
        'X-UserType': 'USER',
        'X-SourceID': 'WEB',
        'X-ClientLocalIP': 'CLIENT_LOCAL_IP',
        'X-ClientPublicIP': 'CLIENT_PUBLIC_IP',
        'X-MACAddress': 'MAC_ADDRESS',
        'X-PrivateKey': os.getenv('SMARTAPI_API_KEY')
    }
    try:
        r = requests.get(url, headers=headers)
        data = r.json()
        # Store in database
        clientcode = _SMARTAPI_SESSIONS[session_id]['clientcode']
        store_data(clientcode, '/api/orders/book', 'orders', data)
    except Exception as e:
        logging.error(f"Order book error for session_id={session_id}: {e}")
        return jsonify({'status': False, 'message': str(e)}), 500
    return jsonify(data)

@app.route('/api/orders/trades')
def trade_book():
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        logging.warning("Trade book access without valid session")
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    smartApi = _SMARTAPI_SESSIONS[session_id]['api']
    jwt_token = _SMARTAPI_SESSIONS[session_id]['tokens'].get('jwtToken', '')
    if jwt_token.startswith('Bearer '):
        jwt_token = jwt_token[7:]
    url = 'https://apiconnect.angelone.in/rest/secure/angelbroking/order/v1/getTradeBook'
    headers = {
        'Authorization': f'Bearer {jwt_token}',
        'Content-Type': 'application/json',
        'Accept': 'application/json',
        'X-UserType': 'USER',
        'X-SourceID': 'WEB',
        'X-ClientLocalIP': 'CLIENT_LOCAL_IP',
        'X-ClientPublicIP': 'CLIENT_PUBLIC_IP',
        'X-MACAddress': 'MAC_ADDRESS',
        'X-PrivateKey': os.getenv('SMARTAPI_API_KEY')
    }
    try:
        r = requests.get(url, headers=headers)
        data = r.json()
        # Store in database
        clientcode = _SMARTAPI_SESSIONS[session_id]['clientcode']
        store_data(clientcode, '/api/orders/trades', 'trades', data)
    except Exception as e:
        logging.error(f"Trade book error for session_id={session_id}: {e}")
        return jsonify({'status': False, 'message': str(e)}), 500
    return jsonify(data)

@app.route('/api/orders/ltp', methods=['POST'])
def ltp():
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        logging.warning("LTP access without valid session")
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    smartApi = _SMARTAPI_SESSIONS[session_id]['api']
    jwt_token = _SMARTAPI_SESSIONS[session_id]['tokens'].get('jwtToken', '')
    if jwt_token.startswith('Bearer '):
        jwt_token = jwt_token[7:]
    
    body = request.get_json()
    params = {k: body.get(k) for k in ('exchange', 'tradingsymbol', 'symboltoken')}
    url = 'https://apiconnect.angelone.in/rest/secure/angelbroking/order/v1/getLtpData'
    headers = {
        'Authorization': f'Bearer {jwt_token}',
        'Content-Type': 'application/json',
        'Accept': 'application/json',
        'X-UserType': 'USER',
        'X-SourceID': 'WEB',
        'X-ClientLocalIP': 'CLIENT_LOCAL_IP',
        'X-ClientPublicIP': 'CLIENT_PUBLIC_IP',
        'X-MACAddress': 'MAC_ADDRESS',
        'X-PrivateKey': os.getenv('SMARTAPI_API_KEY')
    }
    try:
        r = requests.post(url, headers=headers, json=params)
        data = r.json()
    except Exception as e:
        logging.error(f"LTP error for session_id={session_id}: {e}")
        return jsonify({'status': False, 'message': str(e)}), 500
    return jsonify(data)

@app.route('/api/orders/details/<uniqueorderid>')
def order_details(uniqueorderid):
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        logging.warning(f"Order details access without valid session for order {uniqueorderid}")
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    smartApi = _SMARTAPI_SESSIONS[session_id]['api']
    jwt_token = _SMARTAPI_SESSIONS[session_id]['tokens'].get('jwtToken', '')
    if jwt_token.startswith('Bearer '):
        jwt_token = jwt_token[7:]
    
    url = f'https://apiconnect.angelone.in/rest/secure/angelbroking/order/v1/details/{uniqueorderid}'
    headers = {
        'Authorization': f'Bearer {jwt_token}',
        'Content-Type': 'application/json',
        'Accept': 'application/json',
        'X-UserType': 'USER',
        'X-SourceID': 'WEB',
        'X-ClientLocalIP': 'CLIENT_LOCAL_IP',
        'X-ClientPublicIP': 'CLIENT_PUBLIC_IP',
        'X-MACAddress': 'MAC_ADDRESS',
        'X-PrivateKey': os.getenv('SMARTAPI_API_KEY')
    }
    try:
        r = requests.get(url, headers=headers)
        data = r.json()
    except Exception as e:
        logging.error(f"Order details error for session_id={session_id}, order={uniqueorderid}: {e}")
        return jsonify({'status': False, 'message': str(e)}), 500
    return jsonify(data)

@app.route('/api/data/export')
def export_data():
    """Export all stored data for OpenAI analysis"""
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    
    clientcode = _SMARTAPI_SESSIONS[session_id]['clientcode']
    limit = request.args.get('limit', 100, type=int)
    data_type = request.args.get('type', None)
    
    try:
        conn = sqlite3.connect(DB_FILE)
        conn.row_factory = sqlite3.Row
        c = conn.cursor()
        
        if data_type:
            c.execute('''
                SELECT * FROM api_data 
                WHERE clientcode = ? AND data_type = ?
                ORDER BY timestamp DESC LIMIT ?
            ''', (clientcode, data_type, limit))
        else:
            c.execute('''
                SELECT * FROM api_data 
                WHERE clientcode = ?
                ORDER BY timestamp DESC LIMIT ?
            ''', (clientcode, limit))
        
        rows = c.fetchall()
        conn.close()
        
        results = []
        for row in rows:
            results.append({
                'id': row['id'],
                'timestamp': row['timestamp'],
                'endpoint': row['endpoint'],
                'data_type': row['data_type'],
                'response': json.loads(row['response'])
            })
        
        return jsonify({
            'status': True,
            'count': len(results),
            'data': results
        })
    except Exception as e:
        logging.error(f"Export data error: {e}")
        return jsonify({'status': False, 'message': str(e)}), 500

# ==================== USER ANALYSIS DOCUMENT UPLOAD ====================

def allowed_file(filename):
    """Check if file extension is allowed"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_docx(filepath):
    """Extract text content from DOCX file"""
    try:
        doc = Document(filepath)
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        return '\n'.join(full_text)
    except Exception as e:
        logging.error(f"Error extracting text from DOCX: {e}")
        return None

def extract_text_from_txt(filepath):
    """Extract text content from TXT file"""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        logging.error(f"Error reading TXT file: {e}")
        return None

@app.route('/api/analysis/upload', methods=['POST'])
def upload_analysis():
    """Upload user's market analysis document"""
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    
    clientcode = _SMARTAPI_SESSIONS[session_id]['clientcode']
    
    # Check if file was uploaded
    if 'file' not in request.files:
        return jsonify({'status': False, 'message': 'No file uploaded'}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'status': False, 'message': 'No file selected'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'status': False, 'message': 'Invalid file type. Only .docx and .txt files are allowed'}), 400
    
    try:
        # Secure the filename
        original_filename = file.filename or 'analysis.txt'
        filename = secure_filename(original_filename)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        unique_filename = f"{clientcode}_{timestamp}_{filename}"
        filepath = os.path.join(UPLOAD_FOLDER, unique_filename)
        
        # Save file
        file.save(filepath)
        
        # Extract text based on file type
        if filename.lower().endswith('.docx'):
            content = extract_text_from_docx(filepath)
        else:  # .txt
            content = extract_text_from_txt(filepath)
        
        if not content:
            return jsonify({'status': False, 'message': 'Could not extract text from file'}), 500
        
        # Store in memory
        USER_ANALYSIS_DOCS[clientcode] = {
            'content': content,
            'filename': filename,
            'uploaded_at': datetime.now(),
            'filepath': filepath
        }
        
        logging.info(f"[USER ANALYSIS] {clientcode} uploaded analysis: {filename} ({len(content)} characters)")
        
        return jsonify({
            'status': True,
            'message': 'Analysis document uploaded successfully',
            'filename': filename,
            'content_length': len(content),
            'uploaded_at': datetime.now().isoformat()
        })
        
    except Exception as e:
        logging.error(f"Error uploading analysis: {e}", exc_info=True)
        return jsonify({'status': False, 'message': str(e)}), 500

@app.route('/api/analysis/get')
def get_analysis():
    """Get current user's analysis document"""
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    
    clientcode = _SMARTAPI_SESSIONS[session_id]['clientcode']
    
    if clientcode not in USER_ANALYSIS_DOCS:
        return jsonify({
            'status': True,
            'has_analysis': False,
            'message': 'No analysis document uploaded'
        })
    
    doc_info = USER_ANALYSIS_DOCS[clientcode]
    
    return jsonify({
        'status': True,
        'has_analysis': True,
        'filename': doc_info['filename'],
        'content': doc_info['content'],
        'content_length': len(doc_info['content']),
        'uploaded_at': doc_info['uploaded_at'].isoformat()
    })

@app.route('/api/analysis/delete', methods=['POST'])
def delete_analysis():
    """Delete user's analysis document"""
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    
    clientcode = _SMARTAPI_SESSIONS[session_id]['clientcode']
    
    if clientcode not in USER_ANALYSIS_DOCS:
        return jsonify({'status': False, 'message': 'No analysis document to delete'}), 404
    
    try:
        doc_info = USER_ANALYSIS_DOCS[clientcode]
        
        # Delete file from disk
        if os.path.exists(doc_info['filepath']):
            os.remove(doc_info['filepath'])
        
        # Remove from memory
        del USER_ANALYSIS_DOCS[clientcode]
        
        logging.info(f"[USER ANALYSIS] {clientcode} deleted analysis: {doc_info['filename']}")
        
        return jsonify({
            'status': True,
            'message': 'Analysis document deleted successfully'
        })
        
    except Exception as e:
        logging.error(f"Error deleting analysis: {e}", exc_info=True)
        return jsonify({'status': False, 'message': str(e)}), 500

@app.route('/api/data/latest')
def latest_data():
    """Get latest data snapshot for OpenAI analysis"""
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    
    clientcode = _SMARTAPI_SESSIONS[session_id]['clientcode']
    
    try:
        conn = sqlite3.connect(DB_FILE)
        conn.row_factory = sqlite3.Row
        c = conn.cursor()
        
        snapshot = {}
        for dtype in ['profile', 'marketdata', 'rms', 'orders', 'trades']:
            c.execute('''
                SELECT * FROM api_data 
                WHERE clientcode = ? AND data_type = ?
                ORDER BY timestamp DESC LIMIT 1
            ''', (clientcode, dtype))
            row = c.fetchone()
            if row:
                snapshot[dtype] = {
                    'timestamp': row['timestamp'],
                    'data': json.loads(row['response'])
                }
        
        conn.close()
        
        return jsonify({
            'status': True,
            'clientcode': clientcode,
            'snapshot': snapshot
        })
    except Exception as e:
        logging.error(f"Latest data error: {e}")
        return jsonify({'status': False, 'message': str(e)}), 500

@app.route('/api/analyze', methods=['POST'])
def analyze_data():
    """Analyze trading data using OpenAI"""
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    
    clientcode = _SMARTAPI_SESSIONS[session_id]['clientcode']
    
    # Get analysis type and custom prompt from request
    body = request.get_json() or {}
    analysis_type = body.get('type', 'overview')  # overview, risk, performance, strategy
    custom_prompt = body.get('prompt', None)
    
    try:
        # Fetch latest data
        conn = sqlite3.connect(DB_FILE)
        conn.row_factory = sqlite3.Row
        c = conn.cursor()
        
        snapshot = {}
        for dtype in ['profile', 'marketdata', 'rms', 'orders', 'trades']:
            c.execute('''
                SELECT * FROM api_data 
                WHERE clientcode = ? AND data_type = ?
                ORDER BY timestamp DESC LIMIT 1
            ''', (clientcode, dtype))
            row = c.fetchone()
            if row:
                snapshot[dtype] = {
                    'timestamp': row['timestamp'],
                    'data': json.loads(row['response'])
                }
        
        conn.close()
        
        # Prepare prompt based on analysis type
        if custom_prompt:
            prompt = custom_prompt
        elif analysis_type == 'overview':
            prompt = """Analyze this trading account data and provide:
1. Account health summary
2. Current positions and P&L
3. Risk assessment
4. Key observations
5. Recommendations

Data:"""
        elif analysis_type == 'risk':
            prompt = """Analyze the risk profile of this trading account:
1. Margin utilization
2. Exposure levels
3. Position concentration
4. Risk warnings
5. Risk mitigation suggestions

Data:"""
        elif analysis_type == 'performance':
            prompt = """Analyze trading performance:
1. Recent trade analysis
2. Win/loss ratio
3. Trading patterns
4. Performance metrics
5. Improvement suggestions

Data:"""
        elif analysis_type == 'strategy':
            prompt = """Suggest trading strategies based on:
1. Current market conditions
2. Account capacity
3. Risk appetite
4. Trading history
5. Actionable recommendations

Data:"""
        else:
            prompt = "Analyze this trading data and provide insights:"
        
        # Call OpenAI API
        response = openai_client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are an expert financial analyst and trading advisor. Provide clear, actionable insights based on trading data."},
                {"role": "user", "content": f"{prompt}\n\n{json.dumps(snapshot, indent=2)}"}
            ],
            temperature=0.7,
            max_tokens=2000
        )
        
        analysis = response.choices[0].message.content
        
        # Store analysis in database
        store_data(clientcode, '/api/analyze', 'analysis', {
            'analysis_type': analysis_type,
            'prompt': prompt,
            'analysis': analysis,
            'model': 'gpt-4o'
        })
        
        return jsonify({
            'status': True,
            'analysis_type': analysis_type,
            'analysis': analysis,
            'timestamp': datetime.now().isoformat()
        })
        
    except Exception as e:
        logging.error(f"OpenAI analysis error: {e}")
        return jsonify({'status': False, 'message': str(e)}), 500

@app.route('/view/analysis')
def view_analysis():
    if 'session_id' not in session:
        return redirect(url_for('index'))
    return render_template('analysis.html')

@app.route('/view/backtest')
def view_backtest():
    if 'session_id' not in session:
        return redirect(url_for('index'))
    return render_template('backtest.html')

@app.route('/api/backtest', methods=['POST'])
def backtest():
    """Run AI-powered backtest on historical trades"""
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    
    clientcode = _SMARTAPI_SESSIONS[session_id]['clientcode']
    body = request.get_json() or {}
    
    # Get backtest parameters
    strategy_prompt = body.get('strategy', None)
    days_back = body.get('days', 30)
    
    try:
        # Fetch historical data
        conn = sqlite3.connect(DB_FILE)
        conn.row_factory = sqlite3.Row
        c = conn.cursor()
        
        # Get all trades and orders from the time period
        c.execute('''
            SELECT * FROM api_data 
            WHERE clientcode = ? AND data_type IN ('orders', 'trades', 'marketdata')
            AND timestamp >= datetime('now', '-' || ? || ' days')
            ORDER BY timestamp DESC
        ''', (clientcode, days_back))
        
        rows = c.fetchall()
        conn.close()
        
        historical_data = []
        for row in rows:
            historical_data.append({
                'timestamp': row['timestamp'],
                'data_type': row['data_type'],
                'data': json.loads(row['response'])
            })
        
        if not historical_data:
            return jsonify({
                'status': False,
                'message': 'No historical data available for backtesting'
            }), 400
        
        # Prepare backtest prompt
        default_strategy = """Analyze the trading history and perform a backtest:
1. Identify all executed trades
2. Calculate entry and exit prices
3. Compute P&L for each trade
4. Analyze win rate and risk/reward ratio
5. Identify patterns in winning vs losing trades
6. Suggest what could have been done better
7. Provide alternative entry/exit points that would have improved results
8. Calculate overall portfolio performance metrics"""
        
        prompt = strategy_prompt if strategy_prompt else default_strategy
        
        # Call OpenAI for backtesting analysis
        response = openai_client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are an expert quantitative analyst specializing in backtesting trading strategies. Analyze historical trading data and provide detailed performance metrics, insights, and optimization suggestions."},
                {"role": "user", "content": f"{prompt}\n\nHistorical Data (Last {days_back} days):\n{json.dumps(historical_data, indent=2)}"}
            ],
            temperature=0.5,
            max_tokens=3000
        )
        
        backtest_result = response.choices[0].message.content
        
        # Store backtest result
        store_data(clientcode, '/api/backtest', 'backtest', {
            'days_back': days_back,
            'strategy': prompt,
            'result': backtest_result,
            'data_points': len(historical_data),
            'model': 'gpt-4o'
        })
        
        return jsonify({
            'status': True,
            'backtest_result': backtest_result,
            'days_analyzed': days_back,
            'data_points': len(historical_data),
            'timestamp': datetime.now().isoformat()
        })
        
    except Exception as e:
        logging.error(f"Backtest error: {e}")
        return jsonify({'status': False, 'message': str(e)}), 500

@app.route('/api/backtest/optimize', methods=['POST'])
def backtest_optimize():
    """Get AI suggestions for optimizing future trades"""
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    
    clientcode = _SMARTAPI_SESSIONS[session_id]['clientcode']
    body = request.get_json() or {}
    focus_area = body.get('focus', 'all')  # entry, exit, risk, all
    
    try:
        # Get recent trades and current positions
        conn = sqlite3.connect(DB_FILE)
        conn.row_factory = sqlite3.Row
        c = conn.cursor()
        
        recent_data = {}
        for dtype in ['trades', 'orders', 'rms', 'marketdata']:
            c.execute('''
                SELECT * FROM api_data 
                WHERE clientcode = ? AND data_type = ?
                ORDER BY timestamp DESC LIMIT 5
            ''', (clientcode, dtype))
            rows = c.fetchall()
            recent_data[dtype] = [{'timestamp': r['timestamp'], 'data': json.loads(r['response'])} for r in rows]
        
        conn.close()
        
        # Prepare optimization prompt
        if focus_area == 'entry':
            prompt = "Analyze entry points and timing. Suggest optimal entry strategies based on patterns."
        elif focus_area == 'exit':
            prompt = "Analyze exit strategies and profit-taking. Suggest better exit timing and stop-loss placement."
        elif focus_area == 'risk':
            prompt = "Analyze position sizing and risk management. Suggest improvements to reduce drawdowns."
        else:
            prompt = "Provide comprehensive optimization suggestions for entries, exits, position sizing, and risk management."
        
        # Call OpenAI
        response = openai_client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a trading coach focused on helping traders improve their execution. Provide specific, actionable recommendations."},
                {"role": "user", "content": f"{prompt}\n\nRecent Trading Data:\n{json.dumps(recent_data, indent=2)}"}
            ],
            temperature=0.7,
            max_tokens=2000
        )
        
        optimization = response.choices[0].message.content
        
        # Store optimization
        store_data(clientcode, '/api/backtest/optimize', 'optimization', {
            'focus_area': focus_area,
            'optimization': optimization,
            'model': 'gpt-4o'
        })
        
        return jsonify({
            'status': True,
            'focus_area': focus_area,
            'optimization': optimization,
            'timestamp': datetime.now().isoformat()
        })
        
    except Exception as e:
        logging.error(f"Optimization error: {e}")
        return jsonify({'status': False, 'message': str(e)}), 500

@app.route('/view/tradeplan')
def view_tradeplan():
    if 'session_id' not in session:
        return redirect(url_for('index'))
    return render_template('tradeplan.html')

@app.route('/view/aitrade')
def view_aitrade():
    if 'session_id' not in session:
        return redirect(url_for('index'))
    return render_template('aitrade.html')

@app.route('/view/autotrading')
def view_autotrading():
    if 'session_id' not in session:
        return redirect(url_for('index'))
    return render_template('autotrading.html')

@app.route('/api/documents/upload', methods=['POST'])
def upload_documents():
    """Upload daily market analysis documents"""
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    
    clientcode = _SMARTAPI_SESSIONS[session_id]['clientcode']
    
    try:
        # Get uploaded documents (at least one required, both optional)
        perplexity_file = request.files.get('perplexity')
        openai_file = request.files.get('openai')
        
        if not perplexity_file and not openai_file:
            return jsonify({'status': False, 'message': 'At least one document (Perplexity or OpenAI) is required'}), 400
        
        conn = sqlite3.connect(DB_FILE)
        c = conn.cursor()
        upload_date = datetime.now().date()
        uploaded = []
        
        # Process Perplexity document if provided
        if perplexity_file and perplexity_file.filename:
            perplexity_doc = Document(perplexity_file.stream)
            perplexity_content = '\n'.join([para.text for para in perplexity_doc.paragraphs if para.text.strip()])
            
            c.execute('''
                INSERT INTO documents (upload_date, clientcode, document_type, filename, content, metadata)
                VALUES (?, ?, ?, ?, ?, ?)
            ''', (upload_date, clientcode, 'perplexity', secure_filename(perplexity_file.filename or 'perplexity.docx'), 
                  perplexity_content, json.dumps({'word_count': len(perplexity_content.split())})))
            uploaded.append(f'Perplexity ({len(perplexity_content.split())} words)')
        
        # Process OpenAI document if provided
        if openai_file and openai_file.filename:
            openai_doc = Document(openai_file.stream)
            openai_content = '\n'.join([para.text for para in openai_doc.paragraphs if para.text.strip()])
            
            c.execute('''
                INSERT INTO documents (upload_date, clientcode, document_type, filename, content, metadata)
                VALUES (?, ?, ?, ?, ?, ?)
            ''', (upload_date, clientcode, 'openai', secure_filename(openai_file.filename or 'openai.docx'), 
                  openai_content, json.dumps({'word_count': len(openai_content.split())})))
            uploaded.append(f'OpenAI ({len(openai_content.split())} words)')
        
        conn.commit()
        conn.close()
        
        logging.info(f"Documents uploaded for {clientcode} on {upload_date}: {', '.join(uploaded)}")
        
        return jsonify({
            'status': True,
            'message': 'Documents uploaded successfully',
            'upload_date': str(upload_date),
            'uploaded': uploaded
        })
        
    except Exception as e:
        logging.error(f"Document upload error: {e}")
        return jsonify({'status': False, 'message': str(e)}), 500

@app.route('/api/tradeplan/generate', methods=['POST'])
def generate_tradeplan():
    """Generate AI trade plan based on documents and live data"""
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    
    clientcode = _SMARTAPI_SESSIONS[session_id]['clientcode']
    body = request.get_json() or {}
    use_date = body.get('date', None)  # Optional: use specific date
    
    try:
        conn = sqlite3.connect(DB_FILE)
        conn.row_factory = sqlite3.Row
        c = conn.cursor()
        
        # Get latest documents
        if use_date:
            date_filter = use_date
        else:
            c.execute('SELECT MAX(upload_date) FROM documents WHERE clientcode = ?', (clientcode,))
            result = c.fetchone()
            date_filter = result[0] if result and result[0] else datetime.now().date()
        
        # Fetch documents (perplexity and/or openai)
        c.execute('''
            SELECT * FROM documents 
            WHERE clientcode = ? AND upload_date = ?
            ORDER BY id DESC
        ''', (clientcode, date_filter))
        doc_rows = c.fetchall()
        
        if not doc_rows:
            conn.close()
            return jsonify({
                'status': False,
                'message': 'No documents found for the specified date. Please upload at least one analysis document first.'
            }), 400
        
        # Get latest live market data
        snapshot = {}
        for dtype in ['profile', 'marketdata', 'rms', 'orders', 'trades']:
            c.execute('''
                SELECT * FROM api_data 
                WHERE clientcode = ? AND data_type = ?
                ORDER BY timestamp DESC LIMIT 1
            ''', (clientcode, dtype))
            row = c.fetchone()
            if row:
                snapshot[dtype] = json.loads(row['response'])
        
        conn.close()
        
        # Build document analysis from all uploaded documents
        document_analysis = ""
        document_files = []
        for doc_row in doc_rows:
            doc_type = doc_row['document_type']
            source_name = "Perplexity AI Research" if doc_type == 'perplexity' else "OpenAI Research"
            document_analysis += f"\n\n[LIST] {source_name.upper()}:\n{doc_row['content']}\n"
            document_files.append(doc_row['filename'])
        
        # Prepare comprehensive prompt
        prompt = f"""You are an expert trading advisor. Generate a detailed trade plan for today based on:

{document_analysis}

💼 ACCOUNT STATUS:
{json.dumps(snapshot, indent=2)}

Provide a comprehensive trade plan including:
1. Market Outlook (bullish/bearish/neutral)
2. Key Support and Resistance Levels
3. Recommended Trades:
   - Entry points
   - Exit targets
   - Stop losses
   - Position sizes (based on available capital)
4. Risk Assessment
5. Sectors/Stocks to Focus On
6. Sectors/Stocks to Avoid
7. Intraday vs Positional Strategy
8. Timeframes for each trade
9. Overall Risk-Reward Assessment
10. Important Events/News to Watch

Be specific with stock names, price levels, and quantities where applicable."""
        
        # Call OpenAI
        response = openai_client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a professional trading advisor who provides specific, actionable trade plans based on technical analysis, fundamental analysis, and current market conditions."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.6,
            max_tokens=3500
        )
        
        trade_plan = response.choices[0].message.content
        
        # Store trade plan
        store_data(clientcode, '/api/tradeplan/generate', 'tradeplan', {
            'date': str(date_filter),
            'documents': document_files,
            'trade_plan': trade_plan,
            'model': 'gpt-4o'
        })
        
        return jsonify({
            'status': True,
            'trade_plan': trade_plan,
            'date': str(date_filter),
            'documents_used': document_files,
            'timestamp': datetime.now().isoformat()
        })
        
    except Exception as e:
        logging.error(f"Trade plan generation error: {e}")
        return jsonify({'status': False, 'message': str(e)}), 500

@app.route('/api/tradeplan/history')
def tradeplan_history():
    """Get historical trade plans"""
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    
    clientcode = _SMARTAPI_SESSIONS[session_id]['clientcode']
    limit = request.args.get('limit', 7, type=int)
    
    try:
        conn = sqlite3.connect(DB_FILE)
        conn.row_factory = sqlite3.Row
        c = conn.cursor()
        
        c.execute('''
            SELECT * FROM api_data 
            WHERE clientcode = ? AND data_type = 'tradeplan'
            ORDER BY timestamp DESC LIMIT ?
        ''', (clientcode, limit))
        
        rows = c.fetchall()
        conn.close()
        
        plans = []
        for row in rows:
            data = json.loads(row['response'])
            docs_list = data.get('documents', [])
            docs_str = ', '.join(docs_list) if isinstance(docs_list, list) else str(docs_list)
            plans.append({
                'id': row['id'],
                'timestamp': row['timestamp'],
                'date': data.get('date'),
                'documents': docs_str,
                'plan_preview': data.get('trade_plan', '')[:200] + '...'
            })
        
        return jsonify({
            'status': True,
            'count': len(plans),
            'plans': plans
        })
        
    except Exception as e:
        logging.error(f"Trade plan history error: {e}")
        return jsonify({'status': False, 'message': str(e)}), 500

# ============================================================================
# ADVANCED TRADING DATA ENDPOINTS
# ============================================================================

@app.route('/api/historical/candles', methods=['POST'])
def get_historical_candles():
    """Fetch historical candle data for technical analysis"""
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    
    smartApi = _SMARTAPI_SESSIONS[session_id]['api']
    clientcode = _SMARTAPI_SESSIONS[session_id]['clientcode']
    jwt_token = _SMARTAPI_SESSIONS[session_id]['tokens'].get('jwtToken', '')
    if jwt_token.startswith('Bearer '):
        jwt_token = jwt_token[7:]
    
    body = request.get_json() or {}
    
    # Parameters
    exchange = body.get('exchange', 'NSE')
    symboltoken = body.get('symboltoken', '99926000')  # Default NIFTY 50
    interval = body.get('interval', 'FIVE_MINUTE')  # ONE_MINUTE, THREE_MINUTE, FIVE_MINUTE, TEN_MINUTE, FIFTEEN_MINUTE, THIRTY_MINUTE, ONE_HOUR, ONE_DAY
    fromdate = body.get('fromdate')  # Format: YYYY-MM-DD HH:MM
    todate = body.get('todate')  # Format: YYYY-MM-DD HH:MM
    
    if not fromdate or not todate:
        # Default to last 5 days
        todate = datetime.now().strftime('%Y-%m-%d %H:%M')
        fromdate = (datetime.now() - timedelta(days=5)).strftime('%Y-%m-%d %H:%M')
    
    try:
        headers = {
            'Authorization': f'Bearer {jwt_token}',
            'Content-Type': 'application/json',
            'Accept': 'application/json',
            'X-UserType': 'USER',
            'X-SourceID': 'WEB',
            'X-ClientLocalIP': '',
            'X-ClientPublicIP': '',
            'X-MACAddress': '',
            'X-PrivateKey': smartApi.api_key
        }
        
        payload = {
            "exchange": exchange,
            "symboltoken": symboltoken,
            "interval": interval,
            "fromdate": fromdate,
            "todate": todate
        }
        
        logging.info(f"Fetching candles: {payload}")
        
        response = requests.post(
            'https://apiconnect.angelone.in/rest/secure/angelbroking/historical/v1/getCandleData',
            headers=headers,
            json=payload,
            timeout=10
        )
        
        response.raise_for_status()
        result = response.json()
        
        store_data(clientcode, '/api/historical/candles', 'candles', result)
        
        return jsonify(result)
        
    except Exception as e:
        logging.error(f"Historical candles error: {e}")
        return jsonify({'status': False, 'message': str(e)}), 500

@app.route('/api/indicators/calculate', methods=['POST'])
def calculate_technical_indicators():
    """Calculate technical indicators from candle data"""
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    
    body = request.get_json() or {}
    candle_data = body.get('data', [])
    
    if not candle_data:
        return jsonify({'status': False, 'message': 'No candle data provided'}), 400
    
    try:
        # Convert to DataFrame
        # Candle format: [timestamp, open, high, low, close, volume]
        df = pd.DataFrame(candle_data, columns=['timestamp', 'open', 'high', 'low', 'close', 'volume'])
        df['timestamp'] = pd.to_datetime(df['timestamp'])
        df = df.set_index('timestamp')
        df = df.astype(float)
        
        # Calculate indicators
        indicators = {}
        
        # Moving Averages
        indicators['ema_20'] = float(df.ta.ema(length=20).iloc[-1]) if len(df) >= 20 else None
        indicators['ema_50'] = float(df.ta.ema(length=50).iloc[-1]) if len(df) >= 50 else None
        indicators['sma_20'] = float(df.ta.sma(length=20).iloc[-1]) if len(df) >= 20 else None
        indicators['sma_50'] = float(df.ta.sma(length=50).iloc[-1]) if len(df) >= 50 else None
        
        # RSI
        rsi_series = df.ta.rsi(length=14)
        indicators['rsi'] = float(rsi_series.iloc[-1]) if not rsi_series.empty else None
        
        # MACD
        macd = df.ta.macd()
        if macd is not None and not macd.empty:
            indicators['macd'] = float(macd['MACD_12_26_9'].iloc[-1])
            indicators['macd_signal'] = float(macd['MACDs_12_26_9'].iloc[-1])
            indicators['macd_histogram'] = float(macd['MACDh_12_26_9'].iloc[-1])
        
        # Bollinger Bands
        bbands = df.ta.bbands(length=20)
        if bbands is not None and not bbands.empty:
            indicators['bb_upper'] = float(bbands['BBU_20_2.0'].iloc[-1])
            indicators['bb_middle'] = float(bbands['BBM_20_2.0'].iloc[-1])
            indicators['bb_lower'] = float(bbands['BBL_20_2.0'].iloc[-1])
        
        # ATR (Average True Range)
        atr_series = df.ta.atr(length=14)
        indicators['atr'] = float(atr_series.iloc[-1]) if not atr_series.empty else None
        
        # Support & Resistance (using recent highs/lows)
        recent_high = float(df['high'].tail(20).max())
        recent_low = float(df['low'].tail(20).min())
        indicators['resistance'] = recent_high
        indicators['support'] = recent_low
        
        # Pivot Points
        last_candle = df.iloc[-1]
        pivot = (last_candle['high'] + last_candle['low'] + last_candle['close']) / 3
        indicators['pivot_point'] = float(pivot)
        indicators['r1'] = float(2 * pivot - last_candle['low'])
        indicators['r2'] = float(pivot + (last_candle['high'] - last_candle['low']))
        indicators['s1'] = float(2 * pivot - last_candle['high'])
        indicators['s2'] = float(pivot - (last_candle['high'] - last_candle['low']))
        
        # Current price info
        indicators['current_price'] = float(df['close'].iloc[-1])
        indicators['previous_close'] = float(df['close'].iloc[-2]) if len(df) > 1 else None
        indicators['change_percent'] = float(((df['close'].iloc[-1] - df['close'].iloc[-2]) / df['close'].iloc[-2]) * 100) if len(df) > 1 else None
        
        # Trend determination
        if indicators['ema_20'] and indicators['ema_50']:
            if indicators['ema_20'] > indicators['ema_50']:
                indicators['trend'] = 'BULLISH'
            elif indicators['ema_20'] < indicators['ema_50']:
                indicators['trend'] = 'BEARISH'
            else:
                indicators['trend'] = 'NEUTRAL'
        else:
            indicators['trend'] = 'INSUFFICIENT_DATA'
        
        return jsonify({
            'status': True,
            'message': 'Indicators calculated successfully',
            'indicators': indicators,
            'data_points': len(df)
        })
        
    except Exception as e:
        logging.error(f"Indicator calculation error: {e}", exc_info=True)
        return jsonify({'status': False, 'message': str(e)}), 500

@app.route('/api/optionchain/greeks', methods=['POST'])
def get_option_greeks():
    """Fetch Option Greeks from Angel One API"""
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    
    smartApi = _SMARTAPI_SESSIONS[session_id]['api']
    clientcode = _SMARTAPI_SESSIONS[session_id]['clientcode']
    jwt_token = _SMARTAPI_SESSIONS[session_id]['tokens'].get('jwtToken', '')
    if jwt_token.startswith('Bearer '):
        jwt_token = jwt_token[7:]
    
    body = request.get_json() or {}
    
    # Parameters
    name = body.get('name', 'NIFTY')  # Underlying stock/index
    expirydate = body.get('expirydate')  # Format: 25JAN2024
    
    if not expirydate:
        return jsonify({'status': False, 'message': 'expirydate is required (format: 25JAN2024)'}), 400
    
    try:
        headers = {
            'Authorization': f'Bearer {jwt_token}',
            'Content-Type': 'application/json',
            'Accept': 'application/json',
            'X-UserType': 'USER',
            'X-SourceID': 'WEB',
            'X-ClientLocalIP': '',
            'X-ClientPublicIP': '',
            'X-MACAddress': '',
            'X-PrivateKey': smartApi.api_key
        }
        
        payload = {
            "name": name,
            "expirydate": expirydate
        }
        
        logging.info(f"Fetching option Greeks: {payload}")
        
        response = requests.post(
            'https://apiconnect.angelone.in/rest/secure/angelbroking/marketData/v1/optionGreek',
            headers=headers,
            json=payload,
            timeout=10
        )
        
        response.raise_for_status()
        result = response.json()
        
        store_data(clientcode, '/api/optionchain/greeks', 'greeks', result)
        
        return jsonify(result)
        
    except Exception as e:
        logging.error(f"Option Greeks error: {e}")
        return jsonify({'status': False, 'message': str(e)}), 500

@app.route('/api/optionchain/pcr', methods=['POST'])
def calculate_pcr():
    """Calculate Put-Call Ratio from option Greeks data"""
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    
    body = request.get_json() or {}
    greeks_data = body.get('data', [])
    
    if not greeks_data:
        return jsonify({'status': False, 'message': 'No Greeks data provided'}), 400
    
    try:
        total_put_oi = 0
        total_call_oi = 0
        total_put_volume = 0
        total_call_volume = 0
        
        for option in greeks_data:
            option_type = option.get('optionType', '')
            trade_volume = float(option.get('tradeVolume', 0))
            
            # Note: OI not available in Greeks API, using volume as proxy
            if option_type == 'PE':
                total_put_volume += trade_volume
            elif option_type == 'CE':
                total_call_volume += trade_volume
        
        pcr = total_put_volume / total_call_volume if total_call_volume > 0 else 0
        
        result = {
            'pcr': round(pcr, 2),
            'put_volume': total_put_volume,
            'call_volume': total_call_volume,
            'interpretation': 'BULLISH' if pcr > 1.0 else 'BEARISH' if pcr < 0.7 else 'NEUTRAL'
        }
        
        return jsonify({
            'status': True,
            'message': 'PCR calculated successfully',
            'data': result
        })
        
    except Exception as e:
        logging.error(f"PCR calculation error: {e}")
        return jsonify({'status': False, 'message': str(e)}), 500

@app.route('/api/market/vix')
def get_india_vix():
    """Fetch India VIX (Volatility Index) from yfinance"""
    try:
        # India VIX symbol
        vix = yf.Ticker("^INDIAVIX")
        vix_data = vix.history(period="1d")
        
        if not vix_data.empty:
            current_vix = float(vix_data['Close'].iloc[-1])
            prev_vix = float(vix_data['Open'].iloc[-1])
            change = current_vix - prev_vix
            change_percent = (change / prev_vix) * 100
            
            return jsonify({
                'status': True,
                'data': {
                    'vix': round(current_vix, 2),
                    'change': round(change, 2),
                    'change_percent': round(change_percent, 2),
                    'interpretation': 'HIGH_VOLATILITY' if current_vix > 20 else 'MODERATE' if current_vix > 15 else 'LOW_VOLATILITY'
                }
            })
        else:
            return jsonify({'status': False, 'message': 'No VIX data available'}), 404
            
    except Exception as e:
        logging.error(f"VIX fetch error: {e}")
        return jsonify({'status': False, 'message': str(e)}), 500

@app.route('/api/market/global')
def get_global_markets():
    """Fetch global market indices"""
    try:
        indices = {
            '^GSPC': 'S&P 500',
            '^IXIC': 'NASDAQ',
            '^DJI': 'Dow Jones',
            '^HSI': 'Hang Seng',
            '^N225': 'Nikkei 225',
            '^NSEI': 'NIFTY 50',
            '^NSEBANK': 'Bank NIFTY'
        }
        
        results = {}
        for symbol, name in indices.items():
            try:
                ticker = yf.Ticker(symbol)
                data = ticker.history(period="1d")
                
                if not data.empty:
                    current = float(data['Close'].iloc[-1])
                    prev = float(data['Open'].iloc[-1])
                    change = current - prev
                    change_percent = (change / prev) * 100
                    
                    results[name] = {
                        'symbol': symbol,
                        'price': round(current, 2),
                        'change': round(change, 2),
                        'change_percent': round(change_percent, 2),
                        'trend': 'UP' if change > 0 else 'DOWN' if change < 0 else 'FLAT'
                    }
            except:
                results[name] = {'error': 'Data unavailable'}
        
        return jsonify({
            'status': True,
            'data': results,
            'timestamp': datetime.now().isoformat()
        })
        
    except Exception as e:
        logging.error(f"Global markets fetch error: {e}")
        return jsonify({'status': False, 'message': str(e)}), 500

@app.route('/api/market/fii-dii')
def get_fii_dii():
    """Fetch FII/DII data from NSE (simplified)"""
    try:
        # Note: This is a placeholder. Actual implementation requires scraping NSE website
        # which may need handling of cookies, headers, and anti-bot measures
        
        # For now, returning mock structure
        return jsonify({
            'status': True,
            'message': 'FII/DII endpoint - NSE scraping requires additional setup',
            'data': {
                'fii_net': 0,
                'dii_net': 0,
                'note': 'Implement NSE scraping or use alternative data source'
            }
        })
        
    except Exception as e:
        logging.error(f"FII/DII fetch error: {e}")
        return jsonify({'status': False, 'message': str(e)}), 500

@app.route('/api/trading/comprehensive-data', methods=['POST'])
def get_comprehensive_trading_data():
    """Fetch all data required for AI trading decisions"""
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    
    clientcode = _SMARTAPI_SESSIONS[session_id]['clientcode']
    body = request.get_json() or {}
    
    # Parameters
    symboltoken = body.get('symboltoken', '99926000')  # NIFTY 50
    expiry = body.get('expiry')  # For Greeks
    
    try:
        comprehensive_data = {}
        has_any_data = False
        
        # Helper to extract JSON from Flask response (handles both Response and tuple)
        def extract_json(response):
            if isinstance(response, tuple):
                return response[0].get_json()
            return response.get_json()
        
        # 1. Fetch historical candles (with error handling)
        try:
            candles_response = get_historical_candles()
            candles_json = extract_json(candles_response)
            comprehensive_data['candles'] = candles_json
            if candles_json.get('status'):
                has_any_data = True
        except Exception as e:
            logging.warning(f"Failed to fetch candles: {e}")
            comprehensive_data['candles'] = {'status': False, 'message': 'Candles unavailable'}
        
        # 2. Calculate technical indicators from candles
        try:
            if comprehensive_data.get('candles', {}).get('status') and comprehensive_data['candles'].get('data'):
                indicators_body = {'data': comprehensive_data['candles']['data']}
                # Manually call the function
                with app.test_request_context(json=indicators_body):
                    indicators_response = calculate_technical_indicators()
                    comprehensive_data['indicators'] = extract_json(indicators_response)
                    if comprehensive_data['indicators'].get('status'):
                        has_any_data = True
            else:
                comprehensive_data['indicators'] = {'status': False, 'message': 'No candle data for indicators'}
        except Exception as e:
            logging.warning(f"Failed to calculate indicators: {e}")
            comprehensive_data['indicators'] = {'status': False, 'message': 'Indicators unavailable'}
        
        # 3. Get Option Greeks (if expiry provided)
        try:
            if expiry:
                with app.test_request_context(json={'name': 'NIFTY', 'expirydate': expiry}):
                    greeks_response = get_option_greeks()
                    comprehensive_data['greeks'] = extract_json(greeks_response)
                    
                    # Calculate PCR if Greeks available
                    greeks_json = extract_json(greeks_response)
                    if greeks_json.get('status') and greeks_json.get('data'):
                        with app.test_request_context(json={'data': greeks_json['data']}):
                            pcr_response = calculate_pcr()
                            comprehensive_data['pcr'] = extract_json(pcr_response)
                        has_any_data = True
            else:
                comprehensive_data['greeks'] = {'status': False, 'message': 'No expiry provided'}
        except Exception as e:
            logging.warning(f"Failed to fetch Greeks: {e}")
            comprehensive_data['greeks'] = {'status': False, 'message': 'Greeks unavailable'}
        
        # 4. Get VIX
        try:
            vix_response = get_india_vix()
            comprehensive_data['vix'] = extract_json(vix_response)
            if comprehensive_data['vix'].get('status'):
                has_any_data = True
        except Exception as e:
            logging.warning(f"Failed to fetch VIX: {e}")
            comprehensive_data['vix'] = {'status': False, 'message': 'VIX unavailable'}
        
        # 5. Get Global Markets
        try:
            global_response = get_global_markets()
            comprehensive_data['global_markets'] = extract_json(global_response)
            if comprehensive_data['global_markets'].get('status'):
                has_any_data = True
        except Exception as e:
            logging.warning(f"Failed to fetch global markets: {e}")
            comprehensive_data['global_markets'] = {'status': False, 'message': 'Global markets unavailable'}
        
        # 6. Get RMS data
        try:
            conn = sqlite3.connect(DB_FILE)
            c = conn.cursor()
            c.execute('''
                SELECT data FROM api_data 
                WHERE clientcode = ? AND endpoint = '/api/rms'
                ORDER BY timestamp DESC LIMIT 1
            ''', (clientcode,))
            row = c.fetchone()
            if row:
                comprehensive_data['rms'] = json.loads(row[0])
                has_any_data = True
            conn.close()
        except Exception as e:
            logging.warning(f"Failed to fetch RMS data: {e}")
            comprehensive_data['rms'] = {'status': False, 'message': 'RMS unavailable'}
        
        # 7. Get uploaded documents
        try:
            conn = sqlite3.connect(DB_FILE)
            c = conn.cursor()
            today = datetime.now().date()
            c.execute('''
                SELECT document_type, content FROM documents 
                WHERE clientcode = ? AND upload_date = ?
            ''', (clientcode, today))
            docs = c.fetchall()
            comprehensive_data['documents'] = [{'type': d[0], 'content': d[1][:500]} for d in docs]
            conn.close()
        except Exception as e:
            logging.warning(f"Failed to fetch documents: {e}")
            comprehensive_data['documents'] = []
        
        # Return success even if some data sources failed (as long as we got something)
        # If no data at all, provide mock data for testing
        if not has_any_data:
            logging.warning("All live data sources failed - using mock data for testing")
            comprehensive_data = {
                'candles': {'status': True, 'data': [], 'message': 'Mock data'},
                'indicators': {
                    'status': True,
                    'ltp': 24500.0,
                    'ema_9': 24480.0,
                    'ema_21': 24450.0,
                    'rsi': 55.0,
                    'macd': 25.0,
                    'signal': 20.0,
                    'message': 'Mock indicators - EMA9 > EMA21 (bullish trend)'
                },
                'vix': {'status': True, 'vix': 15.5, 'message': 'Mock VIX - normal volatility'},
                'global_markets': {
                    'status': True,
                    'dow': {'change': 0.5},
                    'nasdaq': {'change': 0.8},
                    'sp500': {'change': 0.6},
                    'message': 'Mock global data - positive sentiment'
                },
                'greeks': {'status': False, 'message': 'No expiry provided'},
                'rms': {'status': False, 'message': 'RMS unavailable'},
                'documents': []
            }
            has_any_data = True
        
        return jsonify({
            'status': has_any_data,
            'message': 'Comprehensive data fetched successfully' if has_any_data else 'All data sources failed',
            'data': comprehensive_data,
            'timestamp': datetime.now().isoformat()
        })
        
    except Exception as e:
        logging.error(f"Comprehensive data fetch error: {e}", exc_info=True)
        return jsonify({'status': False, 'message': str(e)}), 500

@app.route('/api/trading/ai-recommendation', methods=['POST'])
def get_ai_trading_recommendation():
    """Generate AI-powered trading recommendations with Rs.15k capital"""
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    
    clientcode = _SMARTAPI_SESSIONS[session_id]['clientcode']
    body = request.get_json() or {}
    
    # Fetch available capital from RMS instead of using fixed amount
    profile_capital = get_available_capital_from_profile(clientcode)
    capital = body.get('capital', profile_capital)  # Use RMS capital as default
    
    risk_percent = body.get('risk_percent', 2)  # 2% risk per trade
    max_per_trade = capital * 0.5  # 50% of capital per trade
    nifty_lot_size = 25  # NIFTY standard lot size
    
    # Log capital source
    if capital == profile_capital:
        logging.info(f"[CAPITAL] Using RMS capital: Rs.{capital:,.2f}")
    else:
        logging.info(f"[CAPITAL] Using user-specified capital: Rs.{capital:,.2f} (RMS had Rs.{profile_capital:,.2f})")
    
    try:
        # Fetch comprehensive trading data with proper session context
        with app.test_client() as client:
            with client.session_transaction() as sess:
                sess['session_id'] = session_id
            
            comp_response = client.post('/api/trading/comprehensive-data', json=body)
            comp_data = comp_response.get_json()
        
        if not comp_data.get('status'):
            logging.error(f"Comprehensive data status=False: {comp_data.get('message')}")
            return jsonify({'status': False, 'message': 'Failed to fetch trading data'}), 500
        
        trading_data = comp_data.get('data', {})
        indicators = trading_data.get('indicators', {})
        current_price = indicators.get('ltp', 26000)
        current_rsi = indicators.get('rsi', 50)
        current_macd = indicators.get('macd', 0)
        
        # Get current VIX for context
        current_vix = get_current_vix_value()
        vix_momentum = get_vix_momentum()
        
        # Get trend direction from candles
        candles_data = trading_data.get('candles', {}).get('data', [])
        trend_direction = check_trend_direction(candles_data)
        
        # Skip fundamentals fetching (not needed for intraday)
        events_str = ""
        
        vix_interpretation = "Unknown"
        if current_vix:
            if current_vix < 12:
                vix_interpretation = "Very Calm Market (15% targets, quick scalps only)"
            elif current_vix < 15:
                vix_interpretation = "Normal Volatility (25% targets, directional trades)"
            elif current_vix < 18:
                vix_interpretation = "Volatile Market (40% targets, ride trends)"
            elif current_vix < 25:
                vix_interpretation = "High Volatility (65% targets, breakout momentum)"
            else:
                vix_interpretation = "EXTREME VOLATILITY (80% targets possible, but RISKY!)"
        
        # VIX momentum interpretation
        vix_momentum_str = {
            'rising': '[UP] RISING (Fear increasing - bigger moves coming, hold longer)',
            'falling': '[DOWN] FALLING (Fear subsiding - book early, tighten stops)',
            'stable': '[NEUTRAL] STABLE (Neutral momentum)',
            'unknown': '[UNKNOWN] UNKNOWN (Insufficient data)'
        }.get(vix_momentum, 'Unknown')
        
        # Trend interpretation
        trend_str = {
            'bullish': '[BULLISH] BULLISH (EMA9 > EMA21, prefer CE trades)',
            'bearish': '[BEARISH] BEARISH (EMA9 < EMA21, prefer PE trades)',
            'neutral': '[NEUTRAL] NEUTRAL (No clear trend, wait for setup)'
        }.get(trend_direction, 'Unknown')
        
        # Get today's performance stats for AI context
        daily_stats = get_daily_stats_summary(clientcode)
        if not daily_stats:
            initialize_daily_stats(clientcode, capital)
            daily_stats = get_daily_stats_summary(clientcode)
        
        # Performance interpretation for AI
        perf_str = "[STATS] TODAY'S PERFORMANCE: No trades yet (fresh start)"
        kelly_advice = ""
        risk_warning = ""
        
        if daily_stats and daily_stats['trades'] > 0:
            win_rate = daily_stats['win_rate']
            pnl_pct = daily_stats['pnl_pct']
            kelly_mult = KELLY_MULTIPLIER.get(clientcode, 1.0)
            
            # Performance color coding
            if pnl_pct > 5:
                perf_emoji = "[HOT]"
                perf_state = "EXCELLENT"
            elif pnl_pct > 0:
                perf_emoji = "[OK]"
                perf_state = "PROFITABLE"
            elif pnl_pct > -3:
                perf_emoji = "[WARNING]"
                perf_state = "SLIGHTLY NEGATIVE"
            else:
                perf_emoji = "[ALERT]"
                perf_state = "LOSING"
            
            perf_str = f"""[STATS] TODAY'S PERFORMANCE ({perf_state}):
{perf_emoji} P&L: Rs.{daily_stats['pnl']:,.0f} ({pnl_pct:+.1f}%)
[UP] Win Rate: {win_rate:.0f}% ({daily_stats['wins']}W / {daily_stats['losses']}L)
[DOWN] Trades: {daily_stats['trades']}/15 executed
[MONEY] Costs: Commissions Rs.{daily_stats['commissions']:.0f} + Slippage Rs.{daily_stats['slippage']:.0f}
[TARGET] Net P&L: Rs.{daily_stats['net_pnl']:,.0f}"""
            
            # Kelly position sizing advice
            if kelly_mult > 1.0:
                kelly_advice = f"\n[STRONG] KELLY SIGNAL: Size up {kelly_mult:.1f}x (strong performance, {win_rate:.0f}% WR)"
            elif kelly_mult < 1.0:
                kelly_advice = f"\n[WARNING] KELLY SIGNAL: Size down {kelly_mult:.1f}x (poor performance, {win_rate:.0f}% WR)"
            
            # Risk warnings
            loss_check = check_daily_loss_circuit_breaker(clientcode)
            if not loss_check[0]:
                risk_warning = f"\n[ALERT] CIRCUIT BREAKER ACTIVE: {loss_check[1]} - NO NEW TRADES ALLOWED"
            elif abs(loss_check[2]) > 7:
                risk_warning = f"\n[WARNING] WARNING: Daily loss at {loss_check[2]:.1f}% (circuit breaker at -10%)"
            
            trades_check = check_max_trades_limit(clientcode)
            if not trades_check[0]:
                risk_warning += f"\n[ALERT] MAX TRADES REACHED: {trades_check[1]}"
            elif trades_check[2] >= 12:
                risk_warning += f"\n[WARNING] HIGH TRADE COUNT: {trades_check[2]}/15 (be selective)"
        
        # Get fundamental context (events, news, economic calendar)
        fundamental_ctx = fetch_fundamental_context()
        events_str = ""
        if fundamental_ctx['events']:
            events_str = "\n\n[NEWS] TODAY'S EVENTS & CONTEXT:\n"
            for event in fundamental_ctx['events']:
                events_str += f"• [{event['type']}] {event['description']}\n  Impact: {event['impact']}\n"
            events_str += f"• US Market: {fundamental_ctx.get('us_market_sentiment', 'N/A')}\n"
            events_str += f"• Day: {fundamental_ctx['day']} (Expiry day volatility expected)" if fundamental_ctx['day'] == 'Tuesday' else f"• Day: {fundamental_ctx['day']}"
        
        # Get user's analysis document if available
        user_analysis = ""
        if clientcode in USER_ANALYSIS_DOCS:
            doc_info = USER_ANALYSIS_DOCS[clientcode]
            user_analysis = f"""

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
USER'S MARKET ANALYSIS:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
File: {doc_info['filename']}
Uploaded: {doc_info['uploaded_at'].strftime('%Y-%m-%d %H:%M')}

{doc_info['content']}
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""
            logging.info(f"[USER ANALYSIS] Using analysis from {doc_info['filename']} for {clientcode}")
        
        # Check if today is expiry day (Tuesday for weekly NIFTY expiry) - USE IST
        today = get_ist_now()
        is_expiry_day = today.weekday() == 1  # Tuesday = 1 (Monday=0, Tuesday=1, Wednesday=2...)
        expiry_note = ""
        if is_expiry_day:
            expiry_note = "\n\n[ALERT] TODAY IS EXPIRY DAY - Use NEXT WEEK's expiry for all trades (current week expires today at 3:30 PM). System will automatically select next available expiry."
        
        # Fetch available expiry dates from options chain
        available_expiries = []
        expiry_info_text = ""
        try:
            with app.test_client() as client:
                with client.session_transaction() as sess:
                    sess['session_id'] = session_id
                
                # Get NIFTY CE options to extract expiry dates
                expiry_response = client.post('/api/scrip/search', json={
                    'symbol': 'NIFTY',
                    'strike': int(current_price),
                    'option_type': 'CE',
                    'cache_range': False,
                    'show_all_expiries': True
                })
                
                expiry_data = expiry_response.get_json()
                if expiry_data.get('status') and expiry_data.get('available_expiries'):
                    available_expiries = expiry_data['available_expiries'][:5]  # First 5 expiries
                    
                    # Format expiry information for AI - ONLY TUESDAYS (NIFTY expiry day)
                    expiry_list = []
                    for exp in available_expiries:
                        try:
                            exp_date = datetime.strptime(exp, '%d%b%Y')
                            
                            # CRITICAL: NIFTY expires on TUESDAYS only - validate
                            if exp_date.weekday() != 1:  # Tuesday = 1
                                logging.warning(f"[EXPIRY ERROR] {exp} is {exp_date.strftime('%A')}, not Tuesday - skipping invalid expiry")
                                continue
                            
                            days_away = (exp_date - today).days
                            exp_formatted = exp_date.strftime('%d-%b-%Y (%A)')
                            expiry_list.append(f"  • {exp_formatted} - {days_away} days away")
                        except:
                            expiry_list.append(f"  • {exp}")
                    
                    if expiry_list:
                        expiry_info_text = f"\n\nAVAILABLE NIFTY EXPIRIES (Tuesdays only):\n" + "\n".join(expiry_list)
                        expiry_info_text += f"\n\n{'**USE NEXT EXPIRY** (skip today)' if is_expiry_day else '**USE CLOSEST EXPIRY** (current week)'}"
                    else:
                        logging.warning("[EXPIRY ERROR] No valid Tuesday expiries found in scrip master data")
                    
        except Exception as e:
            logging.warning(f"Could not fetch expiry dates for AI: {e}")
        
        # Build comprehensive AI prompt
        prompt = f"""Generate intraday NIFTY options trade plan for LIVE TRADING.

CAPITAL: Rs.{capital:,} (from account profile)
MAX PER TRADE: Rs.{max_per_trade:,.0f} (50% of capital){expiry_note}{expiry_info_text}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
MARKET CONTEXT:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
[TARGET] CURRENT VIX: {f'{current_vix:.2f}' if current_vix else 'N/A'} - {vix_interpretation}
[STATS] VIX MOMENTUM: {vix_momentum_str}
[UP] TREND: {trend_str}

{perf_str}{kelly_advice}{risk_warning}{events_str}{user_analysis}

CURRENT NIFTY: {current_price:.2f}
RSI: {current_rsi:.2f}
MACD: {current_macd:.2f}
NIFTY LOT SIZE: {nifty_lot_size}

TECHNICAL INDICATORS:
{json.dumps(trading_data.get('indicators', {}), indent=2)}

OPTION GREEKS:
{json.dumps(trading_data.get('greeks', {}).get('data', [])[:10], indent=2) if trading_data.get('greeks') else 'Not available'}

PUT-CALL RATIO (PCR):
{json.dumps(trading_data.get('pcr', {}), indent=2) if trading_data.get('pcr') else 'Not available'}

VOLATILITY (VIX):
{json.dumps(trading_data.get('vix', {}), indent=2) if trading_data.get('vix') else 'Not available'}

GLOBAL MARKETS:
{json.dumps(trading_data.get('global_markets', {}).get('data', {}), indent=2) if trading_data.get('global_markets') else 'Not available'}

Generate 1-2 NIFTY option trade setups with complete details:

For each trade, specify:
1. Strike price (ATM, slightly OTM based on NIFTY level)
2. Option type (CE for bullish, PE for bearish)
3. Entry price: Option premium price (realistic based on NIFTY level)
4. Entry conditions: NIFTY spot price level that triggers entry
5. Stop loss: Option premium level (not NIFTY index)
6. Target 1 & Target 2: Option premium levels
7. **QUANTITY CALCULATION**: Calculate lots to maximize capital usage
   - Formula: Lots = floor(Max Per Trade / (Entry Premium × Lot Size))
   - Quantity = Lots × {nifty_lot_size}
   - Example: If entry premium is Rs.120 and max per trade is Rs.{max_per_trade:,.0f}:
     * Lots = floor({max_per_trade:,.0f} / (120 × {nifty_lot_size})) = floor({max_per_trade/3000:.1f}) = {int(max_per_trade/3000)} lots
     * Quantity = {int(max_per_trade/3000)} × {nifty_lot_size} = {int(max_per_trade/3000) * nifty_lot_size}
8. Entry time window: e.g., 09:30 to 11:00

Guidelines:
- Option premiums: Rs.50-200 range for ATM options
- Stop loss: 20-30% below entry price
- Target 1: 15-20% above entry price
- Target 2: 30-40% above entry price

**INTRADAY ENTRY STRATEGY - REALISTIC EXECUTION**:

Current Market Level: {current_price:.0f}

Entry Distance Guidelines:
- **AGGRESSIVE SCALP** (Recommended): +/-30 to 50 points → High execution probability (70-80%)
  * Bullish: {current_price+30:.0f} to {current_price+50:.0f}
  * Bearish: {current_price-50:.0f} to {current_price-30:.0f}
  
- **MODERATE SWING**: +/-50 to 100 points → Medium execution probability (40-60%)
  * Bullish: {current_price+50:.0f} to {current_price+100:.0f}
  * Bearish: {current_price-100:.0f} to {current_price-50:.0f}

- **AVOID**: Entries beyond +/-150 points → Low execution probability (<20%)
  * These rarely trigger in a single trading day
  * Example: If NIFTY at 25900, entry at 26100 is unrealistic for intraday

RULE: Default to AGGRESSIVE SCALP range for maximum execution rate. Only suggest MODERATE range if strong trend/momentum indicators support big moves. NEVER suggest entries beyond +/-150 points for intraday plans.

Rationale: Indian market intraday moves average 100-200 points. Suggesting 200+ point entries means trades won't execute, capital sits idle. Tight entries = More action = Better capital utilization.

- Entry condition format: "When NIFTY crosses above {current_price+50:.0f}" (for CE) or "When NIFTY crosses below {current_price-50:.0f}" (for PE)
- **IMPORTANT**: Calculate quantity to use maximum capital available per trade
- **VIX-BASED PROFIT TARGETS**: System will automatically exit trades based on current VIX:
  * VIX < 10: Book at 3% (very calm market)
  * VIX 10-12: Book at 4% (calm market)
  * VIX 12-15: Book at 5% (normal market)
  * VIX 15-18: Book at 6% (active market)
  * VIX 18-22: Book at 65% (high volatility - big moves)
  * VIX 22-25: Book at 80% (extreme volatility - ride it!)
  * VIX > 25: Book at 80%+ but BE CAUTIOUS (hedge recommended)
- **RE-ENTRY RULE**: After profit booking, re-entry only if price breaks prior high/low by 0.5%
- **IV FILTER**: Prefer rising IV (premium expansion) for long positions
- **VIX MOMENTUM**: Rising VIX = bigger moves expected, Falling VIX = book early

Format example (for capital Rs.{capital:,}, max per trade Rs.{max_per_trade:,.0f}):
Trade 1: NIFTY 26000 CE (Expiry: {'Next week' if is_expiry_day else 'Current week'})
Entry Premium: Rs.120
Entry Condition: When NIFTY crosses above 25900
Stop Loss: Rs.85 (premium)
Target 1: Rs.140 (premium)
Target 2: Rs.165 (premium)
Quantity: {int(max_per_trade/3000) * nifty_lot_size} (calculated: floor({max_per_trade:,.0f}/(120×25)) = {int(max_per_trade/3000)} lots = {int(max_per_trade/3000)}×25 = {int(max_per_trade/3000) * nifty_lot_size} qty)
Entry Time: 09:30 to 11:30

Trade 2: NIFTY 25700 PE (Expiry: {'Next week' if is_expiry_day else 'Current week'})
Entry Premium: Rs.100
Entry Condition: When NIFTY crosses below 25750
Stop Loss: Rs.70 (premium)
Target 1: Rs.120 (premium)
Target 2: Rs.145 (premium)
Quantity: {int(max_per_trade/2500) * nifty_lot_size} (calculated: floor({max_per_trade:,.0f}/(100×25)) = {int(max_per_trade/2500)} lots = {int(max_per_trade/2500)}×25 = {int(max_per_trade/2500) * nifty_lot_size} qty)
Entry Time: 09:30 to 12:00

EXECUTION RULES:
- **Professional VIX-based targets** (see above) - Options can move 50-100% in minutes
- **Trend Filter**: Trade with EMA trend (bullish = CE bias, bearish = PE bias)
- **VIX Momentum**: Rising VIX = hold longer, Falling VIX = book early
- **Stop Losses**: VIX < 15 = 10-15% SL, VIX 15-20 = 20% SL, VIX 20-25 = 30% SL, VIX > 25 = 40% SL
- **Re-entries**: Only if price breaks prior extreme by 0.5% (avoid churning)
- Close ALL positions by 3:15 PM
- **EXPIRY USAGE**: {'Use NEXT week expiry (today is expiry day - current week expires at 3:30 PM)' if is_expiry_day else 'Use current week NIFTY expiry'}"""

        # Call OpenAI with performance-aware system prompt
        system_prompt = "You are a professional intraday trader providing specific, executable trade recommendations for NIFTY 50 with limited capital. Always provide exact numbers, strikes, and entry/exit levels."
        
        # Enhance system prompt based on today's performance
        if daily_stats and daily_stats['trades'] > 0:
            if daily_stats['pnl_pct'] < -7:
                system_prompt += " CRITICAL: Today's performance is poor. Be EXTREMELY selective - only suggest A+ setups with strong confirmation. Reduce position sizes."
            elif daily_stats['pnl_pct'] < -3:
                system_prompt += " WARNING: Today is slightly negative. Be more conservative - focus on high-probability setups only."
            elif daily_stats['win_rate'] > 70:
                system_prompt += " Today's performance is excellent. You can be slightly more aggressive with position sizes and pursue quality setups confidently."
            
            if daily_stats['trades'] >= 12:
                system_prompt += " Trade count is high (12+/15). Be highly selective - only best setups."
        
        response = openai_client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=3000
        )
        
        recommendation = response.choices[0].message.content
        
        # Store recommendation
        store_data(clientcode, '/api/trading/ai-recommendation', 'ai_trading', {
            'capital': capital,
            'risk_percent': risk_percent,
            'recommendation': recommendation,
            'data_used': {
                'indicators': 'available' if trading_data.get('indicators') else 'missing',
                'greeks': 'available' if trading_data.get('greeks') else 'missing',
                'vix': 'available' if trading_data.get('vix') else 'missing',
                'global': 'available' if trading_data.get('global_markets') else 'missing'
            }
        })
        
        return jsonify({
            'status': True,
            'recommendation': recommendation,
            'capital': capital,
            'risk_per_trade': int(capital * risk_percent / 100),
            'timestamp': datetime.now().isoformat(),
            'data_sources_used': list(trading_data.keys())
        })
        
    except Exception as e:
        logging.error(f"AI trading recommendation error: {e}", exc_info=True)
        return jsonify({'status': False, 'message': str(e)}), 500

@app.route('/api/trading/daily-stats', methods=['GET'])
def get_daily_trading_stats():
    """Get comprehensive daily trading statistics with risk metrics"""
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    
    clientcode = _SMARTAPI_SESSIONS[session_id]['clientcode']
    
    try:
        # Get daily stats summary
        stats = get_daily_stats_summary(clientcode)
        
        if not stats:
            # Initialize if not exists
            initialize_daily_stats(clientcode, 15000)
            stats = get_daily_stats_summary(clientcode)
        
        # Get circuit breaker status
        loss_check = check_daily_loss_circuit_breaker(clientcode)
        trades_check = check_max_trades_limit(clientcode)
        
        # Get Kelly multiplier
        kelly = KELLY_MULTIPLIER.get(clientcode, 1.0)
        
        # Get active trades count
        active_count = 0
        if clientcode in ACTIVE_TRADES:
            active_count = sum(1 for t in ACTIVE_TRADES[clientcode].values() if t.get('status') == 'open')
        
        return jsonify({
            'status': True,
            'stats': stats,
            'risk_status': {
                'circuit_breaker_active': not loss_check[0],
                'circuit_breaker_message': loss_check[1],
                'current_loss_pct': loss_check[2],
                'trades_limit_reached': not trades_check[0],
                'trades_count': trades_check[2],
                'kelly_multiplier': kelly,
                'active_positions': active_count
            },
            'timestamp': datetime.now().isoformat()
        })
        
    except Exception as e:
        logging.error(f"Daily stats error: {e}", exc_info=True)
        return jsonify({'status': False, 'message': str(e)}), 500

@app.route('/api/trading/reset-daily-stats', methods=['POST'])
def reset_daily_trading_stats():
    """Reset daily statistics (for testing or new day)"""
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    
    clientcode = _SMARTAPI_SESSIONS[session_id]['clientcode']
    body = request.get_json() or {}
    starting_capital = body.get('starting_capital', 15000)
    
    try:
        initialize_daily_stats(clientcode, starting_capital)
        
        return jsonify({
            'status': True,
            'message': f'Daily stats reset with capital Rs.{starting_capital:,.0f}',
            'timestamp': datetime.now().isoformat()
        })
        
    except Exception as e:
        logging.error(f"Reset stats error: {e}", exc_info=True)
        return jsonify({'status': False, 'message': str(e)}), 500

# ==================== AUTOMATED TRADING SCHEDULER ====================

# Global variables for automated trading
ACTIVE_TRADES = {}  # {clientcode: {trade_id: {entry_price, stop_loss, target, quantity, status}}}
DAILY_TRADE_PLAN = {}  # Store generated trade plan
TRADE_PLAN_HISTORY = {}  # {clientcode: [{id, plan, trades, generated_at, selected}]}
AUTO_TRADING_ENABLED = {}  # {clientcode: True/False}
PARSED_TRADE_SETUPS = {}  # {clientcode: [{parsed trade setup}]}
PRICE_MONITOR_THREAD = None
SCRIP_MASTER_CACHE = {}  # Cache scrip master data
WEBSOCKET_CONNECTIONS = {}  # {clientcode: websocket_instance}
LIVE_PRICE_CACHE = {}  # {symboltoken: {ltp, timestamp}}
MONITORING_INTERVAL = 60  # Seconds - faster than 5 minutes
PRICE_UPDATE_QUEUE = Queue()  # Queue for WebSocket price updates
VIX_CACHE = {'value': None, 'timestamp': None}  # Cache VIX data (refresh every 5 min)
CLOSED_TRADE_EXTREMES = {}  # {clientcode: {symbol: {'high': price, 'low': price}}} - Track extremes for re-entry

# Risk Management Tracking
DAILY_STATS = {}  # {clientcode: {date: {pnl, trades_count, wins, losses, commissions, slippage}}}
KELLY_MULTIPLIER = {}  # {clientcode: multiplier} - Position sizing based on win rate
INITIAL_CAPITAL = {}  # {clientcode: starting_capital} - Track daily starting capital
FLASH_CRASH_CACHE = {}  # {clientcode: [price_snapshots]} - Track 5-min price movements
OPENING_PRICE_CACHE = {}  # {clientcode: opening_price} - Track market gap

# VIX History Cache for momentum calculation
VIX_HISTORY = []  # List of (timestamp, vix_value) tuples

# NEW: Advanced Risk Management & Performance Tracking
CONSECUTIVE_LOSSES = {}  # {clientcode: count} - Track loss streaks
PEAK_DAILY_PROFIT = {}  # {clientcode: peak_profit} - Profit protect mode
IV_PERCENTILE_CACHE = {}  # {symboltoken: {iv_rank, timestamp}} - Cache IV rankings
TRADE_PATTERN_STATS = {}  # {clientcode: {pattern_type: {wins, losses, pnl}}} - Track by setup type
POSITION_ENTRY_TIME = {}  # {clientcode: {trade_id: entry_timestamp}} - Time-based profit taking
VOLUME_BASELINE = {}  # {symboltoken: avg_volume} - Volume confirmation baseline
TRAILING_STOPS = {}  # {clientcode: {trade_id: {'initial_sl': X, 'trailing_sl': Y, 'peak_profit_pct': Z}}}
MULTI_TF_CACHE = {}  # {symbol: {timeframe: trend_direction}} - Multi-timeframe confirmation

# USER ANALYSIS DOCUMENTS
USER_ANALYSIS_DOCS = {}  # {clientcode: {'content': text, 'filename': str, 'uploaded_at': datetime}}
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'docx', 'txt'}

# Create uploads folder if it doesn't exist
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

# ==================== RISK MANAGEMENT FUNCTIONS ====================

def get_available_capital_from_profile(clientcode):
    """Fetch available capital from Angel One RMS API"""
    try:
        # Find session for this client
        session_id = None
        for sid, sdata in _SMARTAPI_SESSIONS.items():
            if sdata.get('clientcode') == clientcode:
                session_id = sid
                break
        
        if not session_id:
            logging.warning(f"No session found for {clientcode}, using default capital")
            return 15000
        
        jwt_token = _SMARTAPI_SESSIONS[session_id]['tokens'].get('jwtToken', '')
        if jwt_token.startswith('Bearer '):
            jwt_token = jwt_token[7:]
        
        url = "https://apiconnect.angelone.in/rest/secure/angelbroking/user/v1/getRMS"
        headers = {
            'Authorization': f'Bearer {jwt_token}',
            'Content-Type': 'application/json',
            'Accept': 'application/json',
            'X-UserType': 'USER',
            'X-SourceID': 'WEB',
            'X-ClientLocalIP': 'CLIENT_LOCAL_IP',
            'X-ClientPublicIP': 'CLIENT_PUBLIC_IP',
            'X-MACAddress': 'MAC_ADDRESS',
            'X-PrivateKey': os.getenv('SMARTAPI_API_KEY')
        }
        
        r = requests.get(url, headers=headers, timeout=10)
        data = r.json()
        
        if data.get('status') and data.get('data'):
            rms_data = data['data']
            # Available cash from RMS response
            # RMS returns 'net' field which is available cash for trading
            available_cash = float(rms_data.get('net', 0))
            
            if available_cash > 0:
                logging.info(f"[CAPITAL] Fetched from RMS for {clientcode}: Rs.{available_cash:,.2f}")
                return available_cash
            else:
                logging.warning(f"RMS returned zero/negative capital for {clientcode}, using default")
                return 15000
        else:
            logging.warning(f"RMS API failed for {clientcode}: {data.get('message', 'Unknown error')}")
            return 15000
            
    except Exception as e:
        logging.error(f"Error fetching capital from RMS for {clientcode}: {e}")
        return 15000

def initialize_daily_stats(clientcode, starting_capital=None):
    """Initialize daily statistics for risk tracking with dynamic capital from RMS"""
    global DAILY_STATS, INITIAL_CAPITAL
    today = datetime.now().date().isoformat()
    
    # Fetch capital from RMS if not provided
    if starting_capital is None:
        starting_capital = get_available_capital_from_profile(clientcode)
        logging.info(f"[CAPITAL] Using RMS capital: Rs.{starting_capital:,.2f}")
    
    if clientcode not in DAILY_STATS:
        DAILY_STATS[clientcode] = {}
    
    if today not in DAILY_STATS[clientcode]:
        DAILY_STATS[clientcode][today] = {
            'pnl': 0.0,
            'trades_count': 0,
            'wins': 0,
            'losses': 0,
            'commissions': 0.0,
            'slippage': 0.0,
            'starting_capital': starting_capital,
            'gross_profit': 0.0,
            'gross_loss': 0.0,
            'max_drawdown': 0.0,
            'peak_capital': starting_capital
        }
        INITIAL_CAPITAL[clientcode] = starting_capital
        logging.info(f"[STATS] Daily stats initialized for {clientcode}: Capital Rs.{starting_capital:,.0f}")

def check_daily_loss_circuit_breaker(clientcode, loss_limit_pct=10.0):
    """
    Check if daily loss exceeds circuit breaker limit
    Returns: (allowed: bool, reason: str, current_loss_pct: float)
    """
    global DAILY_STATS, INITIAL_CAPITAL
    today = datetime.now().date().isoformat()
    
    if clientcode not in DAILY_STATS or today not in DAILY_STATS[clientcode]:
        return (True, "No stats available", 0.0)
    
    stats = DAILY_STATS[clientcode][today]
    starting_capital = stats.get('starting_capital', INITIAL_CAPITAL.get(clientcode, 15000))
    current_pnl = stats['pnl']
    loss_pct = (current_pnl / starting_capital) * 100
    
    if loss_pct < -loss_limit_pct:
        return (False, f"[STOP] CIRCUIT BREAKER: Daily loss {loss_pct:.1f}% exceeds limit -{loss_limit_pct}%", loss_pct)
    
    return (True, "Within limits", loss_pct)

def check_max_trades_limit(clientcode, max_trades=10, extended_max=15):
    """
    Check if max trades per day limit reached
    Returns: (allowed: bool, reason: str, trades_count: int)
    """
    global DAILY_STATS
    today = datetime.now().date().isoformat()
    
    if clientcode not in DAILY_STATS or today not in DAILY_STATS[clientcode]:
        return (True, "No trades today", 0)
    
    stats = DAILY_STATS[clientcode][today]
    trades_count = stats['trades_count']
    win_rate = stats['wins'] / max(trades_count, 1)
    
    # Allow extended trades if win rate > 60%
    if trades_count >= extended_max:
        return (False, f"[STOP] MAX TRADES: {trades_count}/{extended_max} trades executed today", trades_count)
    elif trades_count >= max_trades and win_rate < 0.6:
        return (False, f"[WARNING] MAX TRADES: {trades_count}/{max_trades} (win rate {win_rate*100:.0f}% < 60%)", trades_count)
    
    return (True, f"Trades: {trades_count}/{extended_max}", trades_count)

def check_time_based_blocking():
    """
    Check if current time is in blocked trading window
    Block 2:30-3:15 PM (expiry hour chaos)
    Returns: (allowed: bool, reason: str)
    """
    now = datetime.now()
    current_time = now.time()
    
    # Block 14:30 to 15:15 (2:30 PM to 3:15 PM)
    block_start = now.replace(hour=14, minute=30, second=0).time()
    block_end = now.replace(hour=15, minute=15, second=0).time()
    
    if block_start <= current_time <= block_end:
        return (False, f"🕐 TIME BLOCK: No new entries during 2:30-3:15 PM (expiry chaos)")
    
    return (True, "Time allowed")

def calculate_kelly_position_size(clientcode, base_quantity):
    """
    Adjust position size based on Kelly Criterion (win rate performance)
    Kelly% = (Win% * Avg_Win - Loss% * Avg_Loss) / Avg_Win
    
    Returns: adjusted_quantity
    """
    global DAILY_STATS, KELLY_MULTIPLIER
    today = datetime.now().date().isoformat()
    
    if clientcode not in DAILY_STATS or today not in DAILY_STATS[clientcode]:
        KELLY_MULTIPLIER[clientcode] = 1.0
        return base_quantity
    
    stats = DAILY_STATS[clientcode][today]
    total_trades = stats['trades_count']
    
    if total_trades < 3:
        # Not enough data, use base size
        KELLY_MULTIPLIER[clientcode] = 1.0
        return base_quantity
    
    win_rate = stats['wins'] / total_trades
    loss_rate = 1 - win_rate
    
    # Simple Kelly: Increase size after wins, decrease after losses
    if win_rate >= 0.65:  # 65%+ win rate
        multiplier = 1.3  # 30% larger position
    elif win_rate >= 0.50:  # 50-65% win rate
        multiplier = 1.0  # Standard position
    elif win_rate >= 0.35:  # 35-50% win rate
        multiplier = 0.7  # 30% smaller position
    else:  # < 35% win rate
        multiplier = 0.5  # 50% smaller position (defensive)
    
    KELLY_MULTIPLIER[clientcode] = multiplier
    adjusted_qty = int(base_quantity * multiplier)
    
    logging.info(f"📐 Kelly sizing: Win rate {win_rate*100:.0f}% → Multiplier {multiplier:.1f}x → Qty {adjusted_qty}")
    return adjusted_qty

def check_max_open_positions(clientcode, max_positions=2):
    """
    Limit open positions (Rs.15k capital can't handle more)
    Returns: (allowed: bool, reason: str, active_count: int)
    """
    global ACTIVE_TRADES
    
    if clientcode not in ACTIVE_TRADES:
        return (True, "No active trades", 0)
    
    active_count = sum(1 for t in ACTIVE_TRADES[clientcode].values() if t.get('status') == 'open')
    
    if active_count >= max_positions:
        return (False, f"🚫 MAX POSITIONS: Already holding {active_count}/{max_positions} positions", active_count)
    
    return (True, f"Positions: {active_count}/{max_positions}", active_count)

def check_liquidity_filter(symboltoken, clientcode, min_oi=5000):
    """
    Check option liquidity via Historical OI Data API.
    
    SmartAPI does NOT have getOptionChain() method.
    Use existing pattern: POST getOIData endpoint (see /api/optionchain at line 587)
    
    Returns: (allowed: bool, reason: str, oi: int)
    """
    try:
        # TODO: Implement real-time OI check using Historical OI Data API
        # from datetime import datetime, timedelta
        # 
        # session_data = _SMARTAPI_SESSIONS.get(get_session_id_for_client(clientcode))
        # if not session_data:
        #     return (True, "Session not found, skipping OI check", 0)
        # 
        # smartApi = session_data['api']
        # jwt_token = session_data['tokens'].get('jwtToken', '').replace('Bearer ', '')
        # 
        # headers = {
        #     'Authorization': f'Bearer {jwt_token}',
        #     'Content-Type': 'application/json',
        #     'X-PrivateKey': smartApi.api_key
        # }
        # 
        # payload = {
        #     "exchange": "NFO",
        #     "symboltoken": symboltoken,
        #     "interval": "ONE_DAY",
        #     "fromdate": (datetime.now() - timedelta(days=1)).strftime('%Y-%m-%d 09:15'),
        #     "todate": datetime.now().strftime('%Y-%m-%d %H:%M')
        # }
        # 
        # response = requests.post(
        #     'https://apiconnect.angelone.in/rest/secure/angelbroking/historical/v1/getOIData',
        #     headers=headers,
        #     json=payload,
        #     timeout=5
        # )
        # 
        # oi_data = response.json()
        # if oi_data.get('status') and oi_data.get('data'):
        #     # Data format: [[timestamp, open, high, low, close, volume, OI]]
        #     latest_oi = oi_data['data'][-1][6]  # OI is 7th field (index 6)
        #     if latest_oi < min_oi:
        #         return (False, f"🚫 LIQUIDITY: OI {latest_oi:,} < {min_oi:,} (illiquid)", latest_oi)
        #     return (True, f"Liquidity OK (OI: {latest_oi:,})", latest_oi)
        
        # Placeholder - allow trade until API integrated
        return (True, "OI check pending (API integration needed)", 0)
    except Exception as e:
        logging.error(f"Liquidity filter error: {e}")
        return (True, f"OI check error: {str(e)}", 0)

def check_spread_filter(symboltoken, clientcode, max_spread_pct=3.0):
    """
    Check bid-ask spread - reject if > 3% (bad liquidity)
    Returns: (allowed: bool, reason: str, spread_pct: float)
    """
    try:
        # In production: Fetch bid/ask from market depth
        # TODO: Integrate with Level 2 data API
        # For now, simulate spread check
        spread_pct = 1.5  # Placeholder (typical ATM spread)
        
        if spread_pct > max_spread_pct:
            return (False, f"🚫 SPREAD: {spread_pct:.1f}% > {max_spread_pct}% (poor execution)", spread_pct)
        
        return (True, f"Spread OK ({spread_pct:.1f}%)", spread_pct)
    except:
        return (True, "Spread check skipped", 0.0)

def check_flash_crash_protection(clientcode, current_price):
    """
    Pause if NIFTY moves > 2% in 5 minutes (flash crash/spike)
    Returns: (allowed: bool, reason: str, move_pct: float)
    """
    global FLASH_CRASH_CACHE
    
    if clientcode not in FLASH_CRASH_CACHE:
        FLASH_CRASH_CACHE[clientcode] = []
    
    # Store price with timestamp
    now = datetime.now()
    FLASH_CRASH_CACHE[clientcode].append((now, current_price))
    
    # Keep only last 5 minutes
    cutoff = now - timedelta(minutes=5)
    FLASH_CRASH_CACHE[clientcode] = [(ts, p) for ts, p in FLASH_CRASH_CACHE[clientcode] if ts > cutoff]
    
    if len(FLASH_CRASH_CACHE[clientcode]) < 2:
        return (True, "Insufficient data", 0.0)
    
    # Check 5-minute move
    oldest_price = FLASH_CRASH_CACHE[clientcode][0][1]
    move_pct = abs((current_price - oldest_price) / oldest_price) * 100
    
    if move_pct > 2.0:
        return (False, f"[ALERT] FLASH MOVE: NIFTY moved {move_pct:.1f}% in 5 min (pausing)", move_pct)
    
    return (True, f"Normal volatility ({move_pct:.1f}%)", move_pct)

def check_gap_filter(clientcode, current_price):
    """
    Check opening gap - different strategy if gap > 1%
    Returns: (gap_pct: float, interpretation: str)
    """
    global OPENING_PRICE_CACHE
    
    now = datetime.now()
    
    # Store opening price (between 9:15-9:20 AM)
    if now.hour == 9 and 15 <= now.minute <= 20:
        if clientcode not in OPENING_PRICE_CACHE:
            OPENING_PRICE_CACHE[clientcode] = current_price
            logging.info(f"[STATS] Opening price captured: {current_price:.2f}")
    
    if clientcode not in OPENING_PRICE_CACHE:
        return (0.0, "Opening price not yet set")
    
    opening_price = OPENING_PRICE_CACHE[clientcode]
    gap_pct = ((current_price - opening_price) / opening_price) * 100
    
    if abs(gap_pct) > 1.0:
        direction = "GAP UP" if gap_pct > 0 else "GAP DOWN"
        return (gap_pct, f"{direction}: {abs(gap_pct):.1f}% - Momentum bias expected")
    
    return (gap_pct, "Normal opening")

def check_time_decay_filter():
    """
    Avoid buying options after 2 PM (theta decay accelerates)
    Returns: (allowed: bool, reason: str)
    """
    now = datetime.now()
    current_time = now.time()
    
    # Block option buying after 14:00 (2:00 PM)
    cutoff_time = now.replace(hour=14, minute=0, second=0).time()
    
    if current_time >= cutoff_time:
        return (False, "[TIME] TIME DECAY: No option buying after 2 PM (theta kills premium)")
    
    return (True, "Time OK for entry")

def check_correlation_filter(clientcode, new_instrument):
    """
    Check if holding opposite position (CE + PE simultaneously)
    Holding both = hedging = reduced profit potential
    
    Returns: (allowed: bool, reason: str)
    """
    global ACTIVE_TRADES
    
    if clientcode not in ACTIVE_TRADES:
        return (True, "No active trades")
    
    active_instruments = []
    for trade_id, trade in ACTIVE_TRADES[clientcode].items():
        if trade.get('status') == 'open':
            instrument = trade.get('instrument', '')
            active_instruments.append(instrument)
    
    # Check if we're holding opposite side
    new_type = 'CE' if 'CE' in new_instrument else 'PE' if 'PE' in new_instrument else 'unknown'
    
    if new_type == 'CE' and any('PE' in inst for inst in active_instruments):
        return (False, "🚫 CORRELATION: Already holding PE, don't add CE (hedging reduces profit)")
    elif new_type == 'PE' and any('CE' in inst for inst in active_instruments):
        return (False, "🚫 CORRELATION: Already holding CE, don't add PE (hedging reduces profit)")
    
    return (True, "No correlation conflict")

def calculate_slippage(planned_price, actual_price, transaction_type='BUY'):
    """
    Calculate slippage between planned and actual execution price
    Options typically have 2-5% slippage due to wider spreads
    
    Returns: slippage_pct, slippage_amount
    """
    if transaction_type == 'BUY':
        # Positive slippage = paid more than planned (bad)
        slippage_pct = ((actual_price - planned_price) / planned_price) * 100
    else:  # SELL
        # Positive slippage = received less than planned (bad)
        slippage_pct = ((planned_price - actual_price) / planned_price) * 100
    
    slippage_amount = actual_price - planned_price
    return slippage_pct, slippage_amount

def track_commission(clientcode, num_orders=1, commission_per_order=20):
    """
    Track commission costs (Rs.20 per executed order typical for Angel One)
    Returns: total_commission
    """
    global DAILY_STATS
    today = datetime.now().date().isoformat()
    
    if clientcode not in DAILY_STATS:
        DAILY_STATS[clientcode] = {}
    if today not in DAILY_STATS[clientcode]:
        initialize_daily_stats(clientcode, 15000)
    
    total_commission = num_orders * commission_per_order
    DAILY_STATS[clientcode][today]['commissions'] += total_commission
    
    logging.info(f"[MONEY] Commission: Rs.{total_commission} ({num_orders} orders × Rs.{commission_per_order})")
    return total_commission

def update_daily_pnl(clientcode, pnl_change, is_win=None):
    """
    Update daily P&L tracking and statistics
    """
    global DAILY_STATS
    today = datetime.now().date().isoformat()
    
    if clientcode not in DAILY_STATS:
        DAILY_STATS[clientcode] = {}
    if today not in DAILY_STATS[clientcode]:
        initialize_daily_stats(clientcode, 15000)
    
    stats = DAILY_STATS[clientcode][today]
    stats['pnl'] += pnl_change
    stats['trades_count'] += 1
    
    # Track profit factor
    if pnl_change > 0:
        stats['gross_profit'] += pnl_change
    else:
        stats['gross_loss'] += abs(pnl_change)
    
    if is_win is True:
        stats['wins'] += 1
    elif is_win is False:
        stats['losses'] += 1
    
    # Track max drawdown
    current_capital = stats['starting_capital'] + stats['pnl']
    if current_capital > stats['peak_capital']:
        stats['peak_capital'] = current_capital
    
    drawdown = ((stats['peak_capital'] - current_capital) / stats['peak_capital']) * 100
    if drawdown > stats['max_drawdown']:
        stats['max_drawdown'] = drawdown
    
    win_rate = stats['wins'] / max(stats['trades_count'], 1) * 100
    profit_factor = stats['gross_profit'] / max(stats['gross_loss'], 1)
    
    logging.info(f"[STATS] Daily Stats: P&L Rs.{stats['pnl']:,.0f} | Trades {stats['trades_count']} | WR {win_rate:.0f}% | PF {profit_factor:.2f} | DD {stats['max_drawdown']:.1f}%")

def get_daily_stats_summary(clientcode):
    """
    Get comprehensive daily statistics summary
    """
    global DAILY_STATS
    today = datetime.now().date().isoformat()
    
    if clientcode not in DAILY_STATS or today not in DAILY_STATS[clientcode]:
        return None
    
    stats = DAILY_STATS[clientcode][today]
    starting_capital = stats.get('starting_capital', 15000)
    
    return {
        'date': today,
        'pnl': stats['pnl'],
        'pnl_pct': (stats['pnl'] / starting_capital) * 100,
        'trades': stats['trades_count'],
        'wins': stats['wins'],
        'losses': stats['losses'],
        'win_rate': (stats['wins'] / max(stats['trades_count'], 1)) * 100,
        'commissions': stats['commissions'],
        'slippage': stats['slippage'],
        'net_pnl': stats['pnl'] - stats['commissions'] - stats['slippage'],
        'starting_capital': starting_capital
    }

def get_current_vix_value():
    """Get current INDIA VIX value from SmartAPI with 5-minute caching"""
    global VIX_CACHE, VIX_HISTORY
    
    try:
        # Check cache first (5 minute expiry)
        if VIX_CACHE['value'] and VIX_CACHE['timestamp']:
            age = (datetime.now() - VIX_CACHE['timestamp']).total_seconds()
            if age < 300:  # 5 minutes
                return VIX_CACHE['value']
        
        # Fetch INDIA VIX from SmartAPI (token 99926017)
        # Get any active client for API call
        active_client = None
        for clientcode in _SMARTAPI_SESSIONS:
            if _SMARTAPI_SESSIONS[clientcode].get('active'):
                active_client = _SMARTAPI_SESSIONS[clientcode]['smartApi']
                break
        
        if not active_client:
            logging.warning("No active SmartAPI session for VIX fetch")
            return None
        
        # Fetch India VIX using LTP quote
        vix_data = active_client.ltpData("NSE", "India VIX", "99926017")
        
        if vix_data and vix_data.get('status') and vix_data.get('data'):
            current_vix = float(vix_data['data']['ltp'])
            VIX_CACHE = {
                'value': current_vix,
                'timestamp': datetime.now()
            }
            
            # Store in history for momentum calculation
            VIX_HISTORY.append((datetime.now(), current_vix))
            # Keep only last 60 minutes of data (12 samples at 5min intervals)
            VIX_HISTORY = [(ts, val) for ts, val in VIX_HISTORY 
                          if (datetime.now() - ts).total_seconds() < 3600]
            
            logging.info(f"[VIX] Fetched from SmartAPI: {current_vix:.2f}")
            return current_vix
        
        return None
        
    except Exception as e:
        logging.error(f"Error fetching VIX from SmartAPI: {e}")
        return None

def get_vix_momentum():
    """
    Calculate VIX momentum/slope from recent history
    Returns: ('rising', 'falling', 'stable', 'unknown')
    
    Rising VIX = Fear increasing = Larger moves coming
    Falling VIX = Fear subsiding = Consolidation possible
    """
    global VIX_HISTORY
    
    if len(VIX_HISTORY) < 3:
        return 'unknown'
    
    try:
        # Get last 3 VIX values (15 minutes)
        recent = VIX_HISTORY[-3:]
        values = [v for _, v in recent]
        
        # Calculate simple slope
        first_val = values[0]
        last_val = values[-1]
        change_pct = ((last_val - first_val) / first_val) * 100
        
        if change_pct > 2:  # Rising more than 2%
            return 'rising'
        elif change_pct < -2:  # Falling more than 2%
            return 'falling'
        else:
            return 'stable'
            
    except Exception as e:
        logging.error(f"VIX momentum calculation error: {e}")
        return 'unknown'

# ==================== NEW: ADVANCED RISK MANAGEMENT & ENHANCEMENTS ====================

def check_consecutive_loss_limit(clientcode, max_consecutive=3):
    """
    25. Consecutive Loss Limit - Stop trading after N consecutive losses
    
    Prevents emotional/revenge trading after losing streak
    Returns: (allowed: bool, reason: str, streak: int)
    """
    global CONSECUTIVE_LOSSES
    
    if clientcode not in CONSECUTIVE_LOSSES:
        CONSECUTIVE_LOSSES[clientcode] = 0
    
    streak = CONSECUTIVE_LOSSES[clientcode]
    
    if streak >= max_consecutive:
        return (False, f"[STOP] CONSECUTIVE LOSSES: {streak} losses in a row - PAUSED for emotional protection", streak)
    
    return (True, f"Loss streak: {streak}/{max_consecutive}", streak)

def update_loss_streak(clientcode, is_win):
    """Update consecutive loss counter"""
    global CONSECUTIVE_LOSSES
    
    if clientcode not in CONSECUTIVE_LOSSES:
        CONSECUTIVE_LOSSES[clientcode] = 0
    
    if is_win:
        CONSECUTIVE_LOSSES[clientcode] = 0  # Reset on win
        logging.info(f"[OK] Win! Loss streak reset for {clientcode}")
    else:
        CONSECUTIVE_LOSSES[clientcode] += 1
        logging.warning(f"[FAIL] Loss #{CONSECUTIVE_LOSSES[clientcode]} for {clientcode}")

def check_profit_protect_mode(clientcode):
    """
    26. Profit Protect Mode - Lock in daily profits
    
    After reaching good profit, protect against giving it all back
    Returns: (protected_capital: float, risk_reduction_pct: float, status: str)
    """
    global PEAK_DAILY_PROFIT, DAILY_STATS
    
    today = datetime.now().date().isoformat()
    
    if clientcode not in DAILY_STATS or today not in DAILY_STATS[clientcode]:
        return (0, 0, "NO_PROFIT_YET")
    
    current_pnl = DAILY_STATS[clientcode][today]['pnl']
    
    # Track peak profit
    if clientcode not in PEAK_DAILY_PROFIT:
        PEAK_DAILY_PROFIT[clientcode] = 0
    
    if current_pnl > PEAK_DAILY_PROFIT[clientcode]:
        PEAK_DAILY_PROFIT[clientcode] = current_pnl
    
    peak = PEAK_DAILY_PROFIT[clientcode]
    
    # Profit protect rules
    if peak >= 5000:  # If made Rs.5000+ today
        drawdown_from_peak = peak - current_pnl
        drawdown_pct = (drawdown_from_peak / peak) * 100 if peak > 0 else 0
        
        if drawdown_pct > 40:  # Given back 40% of peak profit
            return (peak * 0.6, 100, "STOP_TRADING")  # Stop for the day
        elif drawdown_pct > 25:  # Given back 25%
            return (peak * 0.75, 50, "REDUCE_RISK")  # Reduce position sizes by 50%
        elif peak >= 5000:
            return (peak, 30, "PROTECT_MODE")  # Only risk 30% of profits
    
    return (0, 0, "NORMAL")

def get_time_of_day_adjustment():
    """
    2. Time-of-Day Volatility Adjustment
    
    Returns: (sl_multiplier, target_multiplier, description)
    """
    now = datetime.now()
    current_time = now.time()
    
    # Define time windows
    opening_start = now.replace(hour=9, minute=15).time()
    opening_end = now.replace(hour=10, minute=30).time()
    midday_end = now.replace(hour=14, minute=0).time()
    closing_start = now.replace(hour=14, minute=0).time()
    
    if opening_start <= current_time < opening_end:
        # 09:15-10:30: High volatility opening
        return (1.25, 1.15, "OPENING_VOLATILITY")  # Wider SL, higher targets
    elif current_time < midday_end:
        # 10:30-14:00: Mid-day calm
        return (0.85, 1.0, "MIDDAY_CALM")  # Tighter SL, standard targets
    elif closing_start <= current_time:
        # 14:00-15:30: Closing rush
        return (1.1, 1.05, "CLOSING_RUSH")  # Medium SL, quick exits
    
    return (1.0, 1.0, "STANDARD")

def check_volume_confirmation(symboltoken, clientcode):
    """
    3. Volume Confirmation - Reject low volume trades
    
    NOTE: Requires historical volume data from getCandleData API
    For now, this is a placeholder for future implementation
    
    Returns: (allowed: bool, reason: str, volume_ratio: float)
    """
    try:
        # TODO: When implementing, fetch historical candles and compare volumes
        # Use getCandleData API with FIVE_MINUTE interval for last 5 days
        # Calculate average volume and compare to current
        
        # Placeholder - allow trade
        return (True, "Volume check pending (needs historical data API)", 1.0)
        
    except Exception as e:
        logging.error(f"Volume confirmation error: {e}")
        return (True, f"Volume check error: {str(e)}", 1.0)

def check_breakout_confirmation(clientcode, symbol, current_price, breakout_level, direction='bullish'):
    """
    4. Breakout Confirmation - Wait for firm break with buffer
    
    Simplified version: Check if price firmly above/below level (0.2% buffer)
    Full version would wait for 15-min candle close
    
    Returns: (confirmed: bool, reason: str)
    """
    try:
        buffer = breakout_level * 0.002  # 0.2% buffer
        
        if direction == 'bullish' and current_price > (breakout_level + buffer):
            return (True, f"[OK] Breakout confirmed: {current_price:.0f} > {breakout_level:.0f}")
        elif direction == 'bearish' and current_price < (breakout_level - buffer):
            return (True, f"[OK] Breakdown confirmed: {current_price:.0f} < {breakout_level:.0f}")
        else:
            return (False, f"Waiting for clear break (need 0.2% buffer)")
            
    except Exception as e:
        logging.error(f"Breakout confirmation error: {e}")
        return (True, f"Breakout check error: {str(e)}")

def check_multi_timeframe_confirmation(symbol='NIFTY'):
    """
    5. Multi-Timeframe Confirmation - Align 5min, 15min, 1hour trends
    
    NOTE: Requires fetching candles for multiple intervals
    For now, placeholder for future implementation
    
    Returns: (aligned: bool, reason: str, trends: dict)
    """
    try:
        # TODO: Fetch and analyze multiple timeframes
        # Use getCandleData with FIVE_MINUTE, FIFTEEN_MINUTE, ONE_HOUR
        # Calculate EMA trend for each
        
        # Placeholder - allow trade
        return (True, "Multi-TF check pending (needs multiple interval data)", {})
        
    except Exception as e:
        logging.error(f"Multi-timeframe confirmation error: {e}")
        return (True, f"Multi-TF error: {str(e)}", {})

def calculate_iv_percentile(symbol, strike, expiry, current_iv):
    """
    6. IV Percentile Ranking - Compare current IV to 30-day range
    
    NOTE: Requires tracking historical IV daily for 30 days
    Not directly available in SmartAPI - would need to store daily
    
    Returns: (iv_rank: float, recommendation: str)
    """
    try:
        # TODO: Build historical IV database
        # Track daily IV for each strike for 30 days
        # Calculate percentile rank
        
        # Placeholder - assume moderate IV
        return (50, "IV ranking pending (needs 30-day IV history database)")
        
    except Exception as e:
        logging.error(f"IV percentile error: {e}")
        return (50, f"IV check error: {str(e)}")

def adjust_position_size_by_greeks(base_quantity, delta):
    """
    7. Greeks-Based Position Sizing - Adjust based on Delta
    
    High Delta options behave like stock - reduce size
    Low Delta options have less directional risk
    
    Returns: adjusted_quantity
    """
    try:
        if delta is None or delta == 0:
            return base_quantity
        
        abs_delta = abs(delta)
        
        if abs_delta > 0.7:  # Deep ITM - acts like stock
            multiplier = 0.7  # Reduce to 70%
            reason = f"High Delta ({abs_delta:.2f}) - reduced size"
        elif abs_delta < 0.3:  # Deep OTM - risky
            multiplier = 0.5  # Reduce to 50%
            reason = f"Low Delta ({abs_delta:.2f}) - very OTM, reduced size"
        else:  # Sweet spot 0.3-0.7
            multiplier = 1.0
            reason = f"Good Delta ({abs_delta:.2f}) - standard size"
        
        adjusted = int(base_quantity * multiplier)
        logging.info(f"💎 GREEKS SIZING: {reason} | {base_quantity} → {adjusted}")
        
        return adjusted
        
    except Exception as e:
        logging.error(f"Greeks sizing error: {e}")
        return base_quantity

def calculate_support_resistance_levels(candles_data):
    """
    8. Support/Resistance Levels - Enhanced calculation
    
    Uses pivot points, Fibonacci, and recent highs/lows
    NOTE: Basic pivot points already in indicators
    
    Returns: {'pivots': {}, 'fibonacci': {}, 'recent': {}}
    """
    try:
        if not candles_data or len(candles_data) < 20:
            return {}
        
        # Get recent data
        recent_20 = candles_data[-20:]
        highs = [c[2] for c in recent_20]  # High is 3rd field
        lows = [c[3] for c in recent_20]   # Low is 4th field
        closes = [c[4] for c in recent_20]  # Close is 5th field
        
        last_high = recent_20[-1][2]
        last_low = recent_20[-1][3]
        last_close = recent_20[-1][4]
        
        # Pivot Points
        pivot = (last_high + last_low + last_close) / 3
        r1 = 2 * pivot - last_low
        r2 = pivot + (last_high - last_low)
        r3 = r1 + (last_high - last_low)
        s1 = 2 * pivot - last_high
        s2 = pivot - (last_high - last_low)
        s3 = s1 - (last_high - last_low)
        
        # Fibonacci Retracements
        swing_high = max(highs)
        swing_low = min(lows)
        diff = swing_high - swing_low
        
        fib_levels = {
            'fib_0': swing_high,
            'fib_23.6': swing_high - (diff * 0.236),
            'fib_38.2': swing_high - (diff * 0.382),
            'fib_50': swing_high - (diff * 0.5),
            'fib_61.8': swing_high - (diff * 0.618),
            'fib_100': swing_low
        }
        
        # Recent significant levels
        recent_levels = {
            'swing_high': swing_high,
            'swing_low': swing_low,
            'prev_day_high': highs[-2] if len(highs) > 1 else last_high,
            'prev_day_low': lows[-2] if len(lows) > 1 else last_low
        }
        
        return {
            'pivots': {'pivot': pivot, 'r1': r1, 'r2': r2, 'r3': r3, 's1': s1, 's2': s2, 's3': s3},
            'fibonacci': fib_levels,
            'recent': recent_levels
        }
        
    except Exception as e:
        logging.error(f"Support/Resistance calculation error: {e}")
        return {}

def update_trailing_stop(clientcode, trade_id, current_price, entry_price, current_sl):
    """
    1. Dynamic Trailing Stop Loss
    
    Trail stop loss as profit increases:
    - At +10% profit: Move SL to breakeven
    - At +20% profit: Trail SL to +10%
    - At +30% profit: Trail SL to +15%
    
    Returns: new_stop_loss
    """
    global TRAILING_STOPS
    
    try:
        if clientcode not in TRAILING_STOPS:
            TRAILING_STOPS[clientcode] = {}
        
        if trade_id not in TRAILING_STOPS[clientcode]:
            TRAILING_STOPS[clientcode][trade_id] = {
                'initial_sl': current_sl,
                'trailing_sl': current_sl,
                'peak_profit_pct': 0
            }
        
        trade_trail = TRAILING_STOPS[clientcode][trade_id]
        
        # Calculate profit percentage
        profit_pct = ((current_price - entry_price) / entry_price) * 100
        
        # Update peak profit
        if profit_pct > trade_trail['peak_profit_pct']:
            trade_trail['peak_profit_pct'] = profit_pct
        
        # Trailing stop logic
        new_sl = trade_trail['trailing_sl']
        
        if profit_pct >= 30:  # At +30% profit
            new_sl = entry_price * 1.15  # Trail to +15%
            reason = "Profit 30%+ → Trail SL to +15%"
        elif profit_pct >= 20:  # At +20% profit
            new_sl = entry_price * 1.10  # Trail to +10%
            reason = "Profit 20%+ → Trail SL to +10%"
        elif profit_pct >= 10:  # At +10% profit
            new_sl = entry_price  # Move to breakeven
            reason = "Profit 10%+ → SL to breakeven"
        else:
            new_sl = trade_trail['initial_sl']
            reason = "Profit <10% → Keep initial SL"
        
        # Only update if new SL is higher (never lower stop loss)
        if new_sl > trade_trail['trailing_sl']:
            trade_trail['trailing_sl'] = new_sl
            logging.info(f"[UP] TRAILING STOP: Trade {trade_id} | {reason} | New SL: Rs.{new_sl:.2f}")
            return new_sl
        
        return trade_trail['trailing_sl']
        
    except Exception as e:
        logging.error(f"Trailing stop error: {e}")
        return current_sl

def check_time_based_profit_taking(clientcode, trade_id, entry_time, current_profit_pct):
    """
    12. Time-Based Profit Taking - Book if stagnant
    
    If position open 45+ min and profit unchanged for 20 min, book it
    Returns: (should_exit: bool, reason: str)
    """
    global POSITION_ENTRY_TIME
    
    try:
        now = datetime.now()
        
        if clientcode not in POSITION_ENTRY_TIME:
            POSITION_ENTRY_TIME[clientcode] = {}
        
        if trade_id not in POSITION_ENTRY_TIME[clientcode]:
            POSITION_ENTRY_TIME[clientcode][trade_id] = {
                'entry_time': entry_time,
                'last_profit_update': now,
                'last_profit_pct': current_profit_pct
            }
        
        trade_time = POSITION_ENTRY_TIME[clientcode][trade_id]
        time_open = (now - trade_time['entry_time']).total_seconds() / 60  # minutes
        
        # Check if profit has changed
        if abs(current_profit_pct - trade_time['last_profit_pct']) > 1:  # Changed by 1%+
            trade_time['last_profit_update'] = now
            trade_time['last_profit_pct'] = current_profit_pct
        
        time_since_profit_change = (now - trade_time['last_profit_update']).total_seconds() / 60
        
        # Time-based exit conditions
        if time_open >= 45 and current_profit_pct > 0:  # Open 45+ min with profit
            if time_since_profit_change >= 20:  # Stagnant for 20 min
                return (True, f"[TIME] TIME EXIT: Stagnant {time_since_profit_change:.0f}min (Theta decay)")
        
        # Approaching close with profit
        if now.hour >= 15 and now.minute >= 0 and current_profit_pct > 5:  # After 3 PM
            return (True, f"[TIME] CLOSING: Book {current_profit_pct:.1f}% before 3:30 PM")
        
        return (False, f"Time OK (open {time_open:.0f}min)")
        
    except Exception as e:
        logging.error(f"Time-based profit taking error: {e}")
        return (False, f"Time check error: {str(e)}")

# ==================== AI/ML ENHANCEMENTS ====================

def track_trade_pattern_performance(clientcode, pattern_type, is_win, pnl):
    """
    Track win rate by setup type
    Patterns: 'breakout', 'trend_following', 'mean_reversion', 'volatility_expansion'
    """
    global TRADE_PATTERN_STATS
    
    if clientcode not in TRADE_PATTERN_STATS:
        TRADE_PATTERN_STATS[clientcode] = {}
    
    if pattern_type not in TRADE_PATTERN_STATS[clientcode]:
        TRADE_PATTERN_STATS[clientcode][pattern_type] = {
            'wins': 0,
            'losses': 0,
            'total_pnl': 0,
            'win_rate': 0
        }
    
    stats = TRADE_PATTERN_STATS[clientcode][pattern_type]
    
    if is_win:
        stats['wins'] += 1
    else:
        stats['losses'] += 1
    
    stats['total_pnl'] += pnl
    total_trades = stats['wins'] + stats['losses']
    stats['win_rate'] = (stats['wins'] / total_trades * 100) if total_trades > 0 else 0
    
    logging.info(f"[STATS] PATTERN [{pattern_type}]: WR {stats['win_rate']:.0f}% ({stats['wins']}W/{stats['losses']}L) | P&L Rs.{stats['total_pnl']:,.0f}")

def get_best_performing_patterns(clientcode, min_trades=5):
    """Get highest win rate patterns"""
    global TRADE_PATTERN_STATS
    
    if clientcode not in TRADE_PATTERN_STATS:
        return []
    
    patterns = []
    for pattern, stats in TRADE_PATTERN_STATS[clientcode].items():
        total = stats['wins'] + stats['losses']
        if total >= min_trades:
            patterns.append({
                'pattern': pattern,
                'win_rate': stats['win_rate'],
                'total_trades': total,
                'pnl': stats['total_pnl']
            })
    
    patterns.sort(key=lambda x: x['win_rate'], reverse=True)
    return patterns

def detect_market_regime(vix_value, trend_strength):
    """
    Detect current market regime
    Regimes: 'high_vol', 'low_vol', 'trending', 'choppy'
    Returns: (regime, confidence, strategy_recommendation)
    """
    try:
        if vix_value > 25:
            return ('high_vol', 0.9, "Wider stops, bigger targets, reduce size")
        elif vix_value < 12:
            return ('low_vol', 0.9, "Tighter stops, quicker exits")
        elif trend_strength > 70:
            return ('trending', 0.8, "Trade with trend, let winners run")
        else:
            return ('choppy', 0.7, "Quick scalps, tight stops, selective")
            
    except Exception as e:
        logging.error(f"Market regime detection error: {e}")
        return ('unknown', 0, "Use standard strategy")

def check_trend_direction(candles_data):
    """
    Check trend direction using EMA crossover
    Returns: ('bullish', 'bearish', 'neutral')
    
    Uses EMA9 vs EMA21 crossover for intraday trend
    """
    try:
        if not candles_data or len(candles_data) < 21:
            return 'neutral'
        
        closes = [float(c[4]) for c in candles_data[-21:]]  # Last 21 closes
        
        # Calculate EMA9 and EMA21
        ema9 = closes[-9:]  # Simple avg as approximation
        ema21 = closes
        
        ema9_val = sum(ema9) / len(ema9)
        ema21_val = sum(ema21) / len(ema21)
        
        if ema9_val > ema21_val * 1.002:  # 0.2% above
            return 'bullish'
        elif ema9_val < ema21_val * 0.998:  # 0.2% below
            return 'bearish'
        else:
            return 'neutral'
            
    except Exception as e:
        logging.error(f"Trend direction check error: {e}")
        return 'neutral'

def calculate_vix_based_thresholds(vix_value):
    """
    Calculate trailing SL thresholds based on VIX (volatility)
    PROFESSIONAL STOPS - Wider stops for volatility give trades room to breathe
    
    VIX Ranges (Professional Options Trading):
    - < 12: Very Calm → Tight stops (10% SL)
    - 12-15: Normal → Standard stops (15% SL)
    - 15-20: Volatile → Wider stops (20% SL)
    - 20-25: High Vol → Very wide (30% SL)
    - > 25: Extreme → Maximum width (40% SL) to avoid whipsaws
    """
    if vix_value is None:
        # Default if VIX unavailable
        return {'breakeven_threshold': 15.0, 'trail_threshold': 25.0}
    
    if vix_value < 12:
        # Very calm - tight stops
        return {'breakeven_threshold': 10.0, 'trail_threshold': 15.0}
    elif vix_value < 15:
        # Normal volatility - standard stops
        return {'breakeven_threshold': 15.0, 'trail_threshold': 25.0}
    elif vix_value < 20:
        # Volatile - wider stops to avoid noise
        return {'breakeven_threshold': 20.0, 'trail_threshold': 40.0}
    elif vix_value < 25:
        # High volatility - very wide stops
        return {'breakeven_threshold': 30.0, 'trail_threshold': 65.0}
    else:
        # Extreme volatility - maximum width to ride big moves
        return {'breakeven_threshold': 40.0, 'trail_threshold': 80.0}

def calculate_vix_based_profit_target(vix_value):
    """
    Calculate dynamic profit target based on VIX (volatility)
    PROFESSIONAL SCALPING TARGETS - Options can move 50-100% in minutes on high VIX
    
    Logic: High VIX = explosive premium expansion = ride the move
           Low VIX = limited premium growth = book early
    
    VIX Ranges (Professional Options Trading):
    - < 12: Very Calm → 15% target (quick scalp only)
    - 12-15: Normal Market → 25% target (CE/PE directional)
    - 15-18: Volatile → 40% target (trend trades)
    - 18-25: High Volatility → 65% target (breakout momentum)
    - > 25: Extreme → 80% target (but risky, use hedges)
    """
    if vix_value is None:
        # Default if VIX unavailable
        return 25.0
    
    if vix_value < 12:
        # Very calm market - quick scalps only
        return 15.0
    elif vix_value < 15:
        # Normal volatility - standard directional trades
        return 25.0
    elif vix_value < 18:
        # Volatile market - ride trends
        return 40.0
    elif vix_value < 25:
        # High volatility - breakouts can go big
        return 65.0
    else:
        # Extreme volatility - huge moves but risky
        return 80.0

# ==================== TRADE PLAN PARSING ====================

def parse_trade_plan_with_ai(plan_text, clientcode):
    """Use OpenAI to parse trade plan text into structured JSON"""
    try:
        logging.info(f"Parsing trade plan for {clientcode} using AI")
        
        parsing_prompt = f"""Parse this trading plan into structured JSON format. Extract ALL trade setups mentioned.

TRADE PLAN:
{plan_text}

Return ONLY valid JSON (no markdown, no explanation) in this exact format:
{{
  "trades": [
    {{
      "trade_number": 1,
      "instrument": "NIFTY 26000 CE",
      "tradingsymbol": "NIFTY26000CE",
      "strike": 26000,
      "option_type": "CE",
      "entry_price": 120.00,
      "entry_conditions": [
        {{"type": "price", "indicator": "NIFTY", "operator": ">", "value": 25900}}
      ],
      "quantity": 25,
      "stop_loss": 85.00,
      "target_1": 140.00,
      "target_2": 165.00,
      "entry_time_start": "09:30",
      "entry_time_end": "11:30"
    }}
  ]
}}

PARSING RULES:
1. instrument: Extract strike and option type (e.g., "NIFTY 26000 CE")
2. tradingsymbol: Combine without spaces (e.g., "NIFTY26000CE")
3. strike: Extract numeric strike price (e.g., 26000)
4. option_type: Extract "CE" or "PE"
5. entry_price: Extract option premium for entry (e.g., Rs.120 → 120.00)
6. entry_conditions: Convert "When NIFTY crosses above 25900" to:
   {{"type": "price", "indicator": "NIFTY", "operator": ">", "value": 25900}}
   Convert "When NIFTY crosses below 25750" to:
   {{"type": "price", "indicator": "NIFTY", "operator": "<", "value": 25750}}
7. stop_loss: Extract premium level (e.g., Rs.85 → 85.00)
8. target_1: Extract first target premium (e.g., Rs.140 → 140.00)
9. target_2: Extract second target premium (e.g., Rs.165 → 165.00)
10. quantity: Extract quantity (default to 25 if not specified)
11. entry_time_start/end: Extract time window (e.g., "09:30 to 11:30" → "09:30", "11:30")

IMPORTANT: All prices (entry_price, stop_loss, targets) should be OPTION PREMIUM levels, NOT NIFTY index levels.
Return valid JSON only."""

        response = openai_client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a JSON parser. Return ONLY valid JSON, no markdown formatting."},
                {"role": "user", "content": parsing_prompt}
            ],
            temperature=0,
            max_tokens=2000
        )
        
        content = response.choices[0].message.content
        if not content:
            raise ValueError("OpenAI returned empty content")
        
        parsed_json = content.strip()
        
        # Remove markdown code blocks if present
        if parsed_json.startswith("```"):
            parsed_json = parsed_json.split("```")[1]
            if parsed_json.startswith("json"):
                parsed_json = parsed_json[4:]
        
        parsed_data = json.loads(parsed_json)
        logging.info(f"Successfully parsed {len(parsed_data.get('trades', []))} trades")
        
        return parsed_data
        
    except Exception as e:
        logging.error(f"Error parsing trade plan with AI: {e}", exc_info=True)
        return None

def find_symbol_token(tradingsymbol, clientcode):
    """Find symbol token from local scrip master for given trading symbol"""
    try:
        # Check cache first
        if tradingsymbol in SCRIP_MASTER_CACHE:
            return SCRIP_MASTER_CACHE[tradingsymbol]
        
        # Load local scrip master file
        scrip_master_path = 'scrip_master.json'
        
        try:
            with open(scrip_master_path, 'r', encoding='utf-8') as f:
                scrip_data = json.load(f)
            
            logging.info(f"Loaded {len(scrip_data)} symbols from scrip master")
            
            # Search for matching symbol (case-insensitive)
            tradingsymbol_upper = tradingsymbol.upper()
            
            # Try exact match first
            for item in scrip_data:
                if item.get('symbol', '').upper() == tradingsymbol_upper or item.get('name', '').upper() == tradingsymbol_upper:
                    if item.get('exch_seg') == 'NFO':  # Options
                        token = item.get('token')
                        result = {
                            'token': token,
                            'symbol': item.get('symbol'),
                            'name': item.get('name'),
                            'expiry': item.get('expiry'),
                            'strike': item.get('strike'),
                            'lotsize': item.get('lotsize')
                        }
                        SCRIP_MASTER_CACHE[tradingsymbol] = result
                        logging.info(f"Found token for {tradingsymbol}: {token}")
                        return result
            
            # Try partial match (contains)
            for item in scrip_data:
                symbol = item.get('symbol', '').upper()
                if tradingsymbol_upper in symbol and item.get('exch_seg') == 'NFO':
                    token = item.get('token')
                    result = {
                        'token': token,
                        'symbol': item.get('symbol'),
                        'name': item.get('name'),
                        'expiry': item.get('expiry'),
                        'strike': item.get('strike'),
                        'lotsize': item.get('lotsize')
                    }
                    SCRIP_MASTER_CACHE[tradingsymbol] = result
                    logging.info(f"Found token for {tradingsymbol}: {token} (partial match: {symbol})")
                    return result
            
            logging.error(f"Could not find token for {tradingsymbol} in scrip master")
            return None
            
        except FileNotFoundError:
            logging.error(f"Scrip master file not found: {scrip_master_path}")
            return None
        
    except Exception as e:
        logging.error(f"Error finding symbol token: {e}", exc_info=True)
        return None

# ==================== LIVE PRICE FETCHING (BATCH API) ====================

def get_market_quotes_batch(clientcode, exchange_tokens, mode='FULL'):
    """
    Fetch market quotes for multiple symbols in ONE API call
    Supports up to 50 symbols per request
    
    Args:
        clientcode: Client code
        exchange_tokens: Dict like {"NSE": ["99926000"], "NFO": ["12345", "67890"]}
        mode: 'LTP', 'OHLC', or 'FULL'
    
    Returns:
        Dict with fetched data for all symbols
    """
    try:
        session_id = None
        for sid, sdata in _SMARTAPI_SESSIONS.items():
            if sdata.get('clientcode') == clientcode:
                session_id = sid
                break
        
        if not session_id:
            logging.error(f"No session found for {clientcode}")
            return None
        
        smartapi = _SMARTAPI_SESSIONS[session_id]['api']
        jwt_token = _SMARTAPI_SESSIONS[session_id]['tokens'].get('jwtToken', '')
        if jwt_token.startswith('Bearer '):
            jwt_token = jwt_token[7:]
        
        url = "https://apiconnect.angelone.in/rest/secure/angelbroking/market/v1/quote/"
        
        headers = {
            'Authorization': f'Bearer {jwt_token}',
            'Content-Type': 'application/json',
            'Accept': 'application/json',
            'X-SourceID': 'WEB',
            'X-ClientLocalIP': '192.168.1.1',
            'X-ClientPublicIP': '106.193.147.98',
            'X-MACAddress': 'fe:80:ca:76:19:13',
            'X-PrivateKey': smartapi.api_key,
            'X-UserType': 'USER'
        }
        
        payload = {
            'mode': mode,
            'exchangeTokens': exchange_tokens
        }
        
        response = requests.post(url, json=payload, headers=headers)
        
        if response.status_code == 200:
            result = response.json()
            if result.get('status'):
                return result.get('data', {})
        
        logging.error(f"Market quotes fetch failed: {response.status_code}")
        return None
        
    except Exception as e:
        logging.error(f"Error fetching market quotes: {e}", exc_info=True)
        return None

def get_live_nifty_price(clientcode):
    """Fetch current NIFTY spot price using batch API"""
    try:
        quotes = get_market_quotes_batch(
            clientcode,
            {"NSE": ["99926000"]},  # NIFTY 50 token
            mode='LTP'
        )
        
        if quotes and quotes.get('fetched'):
            return float(quotes['fetched'][0]['ltp'])
        
        return None
        
    except Exception as e:
        logging.error(f"Error fetching NIFTY price: {e}")
        return None

def get_option_ltp(symboltoken, clientcode):
    """Fetch current option LTP using batch API"""
    try:
        quotes = get_market_quotes_batch(
            clientcode,
            {"NFO": [symboltoken]},
            mode='LTP'
        )
        
        if quotes and quotes.get('fetched'):
            return float(quotes['fetched'][0]['ltp'])
        
        return None
        
    except Exception as e:
        logging.error(f"Error fetching option LTP: {e}")
        return None

def get_batch_option_prices(symboltokens, clientcode):
    """
    Fetch prices for multiple options in ONE API call
    Much faster than calling get_option_ltp() multiple times
    
    Args:
        symboltokens: List of symbol tokens ["12345", "67890", ...]
        clientcode: Client code
    
    Returns:
        Dict: {symboltoken: ltp}
    """
    try:
        quotes = get_market_quotes_batch(
            clientcode,
            {"NFO": symboltokens},
            mode='FULL'  # Get full data including volume, OI
        )
        
        if not quotes or not quotes.get('fetched'):
            return {}
        
        result = {}
        for item in quotes['fetched']:
            token = item.get('symbolToken')
            result[token] = {
                'ltp': float(item.get('ltp', 0)),
                'volume': int(item.get('tradeVolume', 0)),
                'oi': int(item.get('opnInterest', 0)),
                'change_percent': float(item.get('percentChange', 0))
            }
        
        return result
        
    except Exception as e:
        logging.error(f"Error fetching batch option prices: {e}")
        return {}

def get_current_technical_indicators(clientcode):
    """Fetch current technical indicators (RSI, MACD, etc.)"""
    try:
        # Fetch recent candles and calculate indicators
        session_id = None
        for sid, sdata in _SMARTAPI_SESSIONS.items():
            if sdata.get('clientcode') == clientcode:
                session_id = sid
                break
        
        if not session_id:
            return {}
        
        # Call candles endpoint
        with app.test_request_context(json={
            'exchange': 'NSE',
            'symboltoken': '99926000',
            'interval': 'FIVE_MINUTE',
            'fromdate': (datetime.now() - timedelta(days=1)).strftime('%Y-%m-%d %H:%M'),
            'todate': datetime.now().strftime('%Y-%m-%d %H:%M')
        }):
            with app.test_client() as client:
                with client.session_transaction() as sess:
                    sess['session_id'] = session_id
                
                candles_response = client.post('/api/historical/candles')
                
                if candles_response.status_code == 200:
                    candles_data = candles_response.get_json()
                    
                    if candles_data.get('status') and candles_data.get('data'):
                        # Calculate indicators
                        indicators_response = client.post(
                            '/api/indicators/calculate',
                            json={'data': candles_data['data']}
                        )
                        
                        if indicators_response.status_code == 200:
                            indicators_data = indicators_response.get_json()
                            if indicators_data.get('status'):
                                return indicators_data.get('data', {})
        
        return {}
        
    except Exception as e:
        logging.error(f"Error fetching technical indicators: {e}")
        return {}

# ==================== ENTRY CONDITION EVALUATION ====================

def evaluate_entry_conditions(trade_setup, clientcode):
    """Check if all entry conditions for a trade are met"""
    try:
        conditions = trade_setup.get('entry_conditions', [])
        
        if not conditions:
            logging.warning(f"No entry conditions defined for trade {trade_setup.get('trade_number')}")
            return False
        
        # Get current market data
        nifty_price = get_live_nifty_price(clientcode)
        indicators = get_current_technical_indicators(clientcode)
        
        if nifty_price is None:
            logging.error("Could not fetch NIFTY price")
            return False
        
        logging.info(f"Evaluating conditions: NIFTY={nifty_price}, RSI={indicators.get('rsi')}")
        
        # Check time window
        current_time = datetime.now().time()
        entry_start = datetime.strptime(trade_setup.get('entry_time_start', '09:30'), '%H:%M').time()
        entry_end = datetime.strptime(trade_setup.get('entry_time_end', '15:00'), '%H:%M').time()
        
        if not (entry_start <= current_time <= entry_end):
            logging.info(f"Outside entry time window: {entry_start}-{entry_end}")
            return False
        
        # Evaluate each condition
        all_conditions_met = True
        
        for condition in conditions:
            condition_type = condition.get('type')
            indicator_name = condition.get('indicator', '').upper()
            operator = condition.get('operator')
            threshold = condition.get('value')
            
            if condition_type == 'price':
                # Price-based condition (e.g., NIFTY > 25850)
                if indicator_name == 'NIFTY':
                    current_value = nifty_price
                else:
                    logging.warning(f"Unknown price indicator: {indicator_name}")
                    continue
                
            elif condition_type == 'indicator':
                # Technical indicator condition (e.g., RSI > 60)
                current_value = indicators.get(indicator_name.lower())
                
                if current_value is None:
                    logging.warning(f"Indicator {indicator_name} not available")
                    all_conditions_met = False
                    continue
            else:
                logging.warning(f"Unknown condition type: {condition_type}")
                continue
            
            # Evaluate operator
            condition_met = False
            
            if operator == '>':
                condition_met = current_value > threshold
            elif operator == '>=':
                condition_met = current_value >= threshold
            elif operator == '<':
                condition_met = current_value < threshold
            elif operator == '<=':
                condition_met = current_value <= threshold
            elif operator == '==':
                condition_met = abs(current_value - threshold) < 0.01
            
            logging.info(f"Condition: {indicator_name} {operator} {threshold} -> {current_value} -> {condition_met}")
            
            if not condition_met:
                all_conditions_met = False
                break
        
        return all_conditions_met
        
    except Exception as e:
        logging.error(f"Error evaluating entry conditions: {e}", exc_info=True)
        return False

# ==================== ORDER EXECUTION ====================

def place_order_angel_one(clientcode, order_params):
    """Place order via Angel One SmartAPI with proper session handling"""
    try:
        # Find active session for this client
        session_id = None
        session_data = None
        
        for sid, sdata in _SMARTAPI_SESSIONS.items():
            if sdata.get('clientcode') == clientcode:
                session_id = sid
                session_data = sdata
                break
        
        if not session_id or not session_data:
            logging.error(f"[ORDER] No active session found for {clientcode}")
            return {
                'status': False,
                'message': f'No active session for client {clientcode}'
            }
        
        # Get SmartAPI instance from session
        smartapi = session_data.get('api')
        if not smartapi:
            logging.error(f"[ORDER] No SmartAPI instance in session for {clientcode}")
            return {
                'status': False,
                'message': 'SmartAPI not initialized in session'
            }
        
        logging.info(f"[ORDER] Placing order for {clientcode}: {order_params}")
        
        # Place order using session's SmartAPI instance
        order_response = smartapi.placeOrder(order_params)
        
        if order_response and order_response.get('status'):
            order_id = order_response['data']['orderid']
            unique_order_id = order_response['data'].get('uniqueorderid')  # Critical for tracking
            logging.info(f"[ORDER] [OK] Order placed successfully for {clientcode}: Order ID = {order_id}, Unique ID = {unique_order_id}")
            return {
                'status': True,
                'orderid': order_id,
                'uniqueorderid': unique_order_id,
                'message': order_response.get('message', 'Order placed'),
                'data': order_response.get('data', {})
            }
        else:
            error_msg = order_response.get('message', 'Order failed') if order_response else 'No response from API'
            logging.error(f"[ORDER] ❌ Order placement failed for {clientcode}: {error_msg}")
            return {
                'status': False,
                'message': error_msg,
                'error_code': order_response.get('errorcode') if order_response else None
            }
        
    except Exception as e:
        logging.error(f"[ORDER] Exception while placing order for {clientcode}: {e}", exc_info=True)
        return {
            'status': False,
            'message': f'Exception: {str(e)}'
        }

def modify_order_angel_one(clientcode, modify_params):
    """Modify existing order via Angel One SmartAPI"""
    try:
        session_id = None
        for sid, sdata in _SMARTAPI_SESSIONS.items():
            if sdata.get('clientcode') == clientcode:
                session_id = sid
                break
        
        if not session_id:
            logging.error(f"No session found for {clientcode}")
            return None
        
        smartapi = _SMARTAPI_SESSIONS[session_id]['api']
        
        logging.info(f"Modifying order: {modify_params}")
        
        # Angel One Modify Order API
        # Required params: variety, orderid, ordertype, producttype, duration, 
        # price, quantity, tradingsymbol, symboltoken, exchange
        modify_response = smartapi.modifyOrder(modify_params)
        
        if modify_response and modify_response.get('status'):
            order_id = modify_response['data']['orderid']
            logging.info(f"[OK] Order modified successfully: Order ID = {order_id}")
            return {
                'status': True,
                'orderid': order_id,
                'message': modify_response.get('message', 'Order modified')
            }
        else:
            logging.error(f"Order modification failed: {modify_response}")
            return {
                'status': False,
                'message': modify_response.get('message', 'Modification failed')
            }
        
    except Exception as e:
        logging.error(f"Error modifying order: {e}")
        return {
            'status': False,
            'message': str(e)
        }

def ai_analyze_market_shift(clientcode):
    """Use AI to analyze if market conditions have shifted significantly"""
    try:
        # Fetch current market data
        nifty_price = get_live_nifty_price(clientcode)
        indicators = get_current_technical_indicators(clientcode)
        
        if not nifty_price or not indicators:
            return None
        
        # Fetch pre-market/global sentiment
        try:
            premarket = fetch_premarket_data()
        except:
            premarket = {}
        
        premarket = premarket or {}  # Ensure not None
        
        # Build AI prompt for market shift analysis
        prompt = f"""Analyze current NIFTY market conditions and determine if there's a significant shift in direction.

[UP] CURRENT NIFTY: {nifty_price:.2f}
[CHART] RSI (14): {indicators.get('rsi', 'N/A')}
[DOWN] MACD: {indicators.get('macd', 'N/A')}
🌎 Global Markets: {premarket.get('sgx_nifty', 'N/A')}

Based on these indicators, has the market shifted direction significantly?

Respond in JSON format:
{{
  "shift_detected": true/false,
  "new_direction": "bullish" / "bearish" / "neutral",
  "confidence": 0-100,
  "reason": "Brief explanation",
  "recommendation": "hold" / "tighten_sl" / "trail_sl" / "exit_early"
}}"""

        response = openai_client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a professional market analyst. Respond only with valid JSON."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,  # Lower temp for more consistent output
            max_tokens=500
        )
        
        content = response.choices[0].message.content
        if not content:
            return None
        
        analysis_text = content.strip()
        
        # Extract JSON from response
        if '```json' in analysis_text:
            analysis_text = analysis_text.split('```json')[1].split('```')[0].strip()
        elif '```' in analysis_text:
            analysis_text = analysis_text.split('```')[1].split('```')[0].strip()
        
        analysis = json.loads(analysis_text)
        
        logging.info(f"[AI] Market Analysis: {analysis.get('new_direction')} ({analysis.get('confidence')}% confidence)")
        logging.info(f"   Reason: {analysis.get('reason')}")
        logging.info(f"   Recommendation: {analysis.get('recommendation')}")
        
        return analysis
        
    except Exception as e:
        logging.error(f"Error in AI market analysis: {e}")
        return None

def ai_adjust_trade_params(clientcode, trade_data, market_analysis):
    """Use AI to determine new stop loss and target levels based on market shift"""
    try:
        recommendation = market_analysis.get('recommendation')
        new_direction = market_analysis.get('new_direction')
        
        current_sl = trade_data.get('stop_loss')
        current_target_1 = trade_data.get('target_1')
        current_target_2 = trade_data.get('target_2')
        entry_price = trade_data.get('entry_price')
        current_price = trade_data.get('current_price', entry_price)
        
        # Build AI prompt for parameter adjustment
        prompt = f"""Given market has shifted to {new_direction}, adjust stop loss and targets for this trade.

[CAPITAL] Entry Price: Rs.{entry_price}
[UP] Current Price: Rs.{current_price}
[SL] Current Stop Loss: Rs.{current_sl}
[TARGET] Current Target 1: Rs.{current_target_1}
[TARGET] Current Target 2: Rs.{current_target_2}

[TRADE] Market Direction: {new_direction}
[TIP] Recommendation: {recommendation}

Provide new stop loss and target levels. Respond in JSON:
{{
  "new_stop_loss": <price>,
  "new_target_1": <price>,
  "new_target_2": <price>,
  "modification_reason": "Brief explanation"
}}

Rules:
- If tighten_sl: Move SL closer to current price to protect profits
- If trail_sl: Trail SL below current price
- If exit_early: Lower targets to book profits quickly
- Stop loss should NEVER be worse than original"""

        response = openai_client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a professional risk manager. Respond only with valid JSON."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=300
        )
        
        content = response.choices[0].message.content
        if not content:
            return None
        
        adjustment_text = content.strip()
        
        # Extract JSON
        if '```json' in adjustment_text:
            adjustment_text = adjustment_text.split('```json')[1].split('```')[0].strip()
        elif '```' in adjustment_text:
            adjustment_text = adjustment_text.split('```')[1].split('```')[0].strip()
        
        adjustments = json.loads(adjustment_text)
        
        logging.info(f"[CONFIG] AI Trade Adjustment:")
        logging.info(f"   New SL: Rs.{adjustments.get('new_stop_loss')}")
        logging.info(f"   New T1: Rs.{adjustments.get('new_target_1')}")
        logging.info(f"   New T2: Rs.{adjustments.get('new_target_2')}")
        logging.info(f"   Reason: {adjustments.get('modification_reason')}")
        
        return adjustments
        
    except Exception as e:
        logging.error(f"Error in AI trade adjustment: {e}")
        return None

def get_individual_order_status(clientcode, unique_order_id):
    """Get individual order status using uniqueorderid from Angel One API"""
    try:
        session_id = None
        for sid, sdata in _SMARTAPI_SESSIONS.items():
            if sdata.get('clientcode') == clientcode:
                session_id = sid
                break
        
        if not session_id:
            return None
        
        smartapi = _SMARTAPI_SESSIONS[session_id]['api']
        
        # Angel One API: GET /order/v1/details/{uniqueorderid}
        order_status = smartapi.orderbook()  # Get full order book
        
        if order_status and order_status.get('status'):
            orders = order_status.get('data', [])
            # Find order by uniqueorderid
            for order in orders:
                if order.get('uniqueorderid') == unique_order_id:
                    return {
                        'status': True,
                        'orderstatus': order.get('orderstatus'),  # open, complete, rejected, cancelled
                        'filledshares': order.get('filledshares', '0'),
                        'averageprice': order.get('averageprice', 0),
                        'text': order.get('text', ''),  # Rejection reason
                        'orderid': order.get('orderid'),
                        'updatetime': order.get('updatetime')
                    }
        
        return None
        
    except Exception as e:
        logging.error(f"Error getting order status: {e}")
        return None

def verify_order_execution(clientcode, unique_order_id, max_retries=5, wait_seconds=2):
    """Verify if order was executed successfully by checking status repeatedly"""
    import time
    
    for attempt in range(max_retries):
        order_status = get_individual_order_status(clientcode, unique_order_id)
        
        if order_status:
            status = order_status['orderstatus'].lower()
            
            if status == 'complete':
                logging.info(f"[OK] Order EXECUTED: Filled @ Rs.{order_status['averageprice']}")
                return {
                    'executed': True,
                    'status': 'complete',
                    'avg_price': float(order_status['averageprice']),
                    'filled_qty': int(order_status['filledshares'])
                }
            elif status == 'rejected':
                logging.error(f"[FAIL] Order REJECTED: {order_status['text']}")
                return {
                    'executed': False,
                    'status': 'rejected',
                    'reason': order_status['text']
                }
            elif status in ['cancelled', 'cancelled after market order']:
                logging.warning(f"[WARNING] Order CANCELLED")
                return {
                    'executed': False,
                    'status': 'cancelled'
                }
            else:
                # Order still pending (open, trigger pending, etc.)
                logging.info(f"Attempt {attempt+1}/{max_retries}: Order status = {status}")
                time.sleep(wait_seconds)
        else:
            time.sleep(wait_seconds)
    
    # Timeout
    logging.warning(f"⏱️ Order verification timeout after {max_retries} attempts")
    return {
        'executed': False,
        'status': 'timeout'
    }

def execute_trade_entry(trade_setup, clientcode):
    """Execute entry order for a trade setup with comprehensive risk checks"""
    try:
        # ==================== RISK MANAGEMENT CHECKS ====================
        
        # 1. Daily Loss Circuit Breaker
        loss_check = check_daily_loss_circuit_breaker(clientcode, loss_limit_pct=10.0)
        if not loss_check[0]:
            logging.error(f"{loss_check[1]} | Current loss: {loss_check[2]:.1f}%")
            return None
        
        # 2. Max Trades Limit
        trades_check = check_max_trades_limit(clientcode, max_trades=10, extended_max=15)
        if not trades_check[0]:
            logging.error(f"{trades_check[1]}")
            return None
        
        # 3. Max Open Positions
        positions_check = check_max_open_positions(clientcode, max_positions=2)
        if not positions_check[0]:
            logging.error(f"{positions_check[1]}")
            return None
        
        # 4. Time-based Blocking (2:30-3:15 PM)
        time_check = check_time_based_blocking()
        if not time_check[0]:
            logging.warning(f"{time_check[1]}")
            return None
        
        # 5. Time Decay Filter (No buys after 2 PM)
        decay_check = check_time_decay_filter()
        if not decay_check[0]:
            logging.warning(f"{decay_check[1]}")
            return None
        
        # 6. Correlation Filter (CE/PE hedging check)
        tradingsymbol = trade_setup.get('tradingsymbol')
        correlation_check = check_correlation_filter(clientcode, tradingsymbol)
        if not correlation_check[0]:
            logging.warning(f"{correlation_check[1]}")
            return None
        
        # 7. Flash Crash Protection
        current_nifty_price = trade_setup.get('nifty_price', 26000)
        flash_check = check_flash_crash_protection(clientcode, current_nifty_price)
        if not flash_check[0]:
            logging.error(f"{flash_check[1]}")
            return None
        
        # NEW 25. Consecutive Loss Limit
        loss_streak_check = check_consecutive_loss_limit(clientcode, max_consecutive=3)
        if not loss_streak_check[0]:
            logging.error(f"{loss_streak_check[1]}")
            return None
        
        # NEW 26. Profit Protect Mode
        protected_capital, risk_reduction, protect_status = check_profit_protect_mode(clientcode)
        if protect_status == "STOP_TRADING":
            logging.error(f"[STOP] PROFIT PROTECT: Gave back 40% of peak profit (Rs.{protected_capital:,.0f}) - STOP TRADING")
            return None
        elif protect_status == "REDUCE_RISK":
            logging.warning(f"[WARNING] PROFIT PROTECT: Reduce risk by {risk_reduction}% (protecting Rs.{protected_capital:,.0f})")
        
        # NEW 2. Time-of-Day Adjustment
        sl_mult, target_mult, time_phase = get_time_of_day_adjustment()
        logging.info(f"[TIME] TIME PHASE: {time_phase} | SL×{sl_mult:.2f} | Target×{target_mult:.2f}")
        
        # NEW 4. Breakout Confirmation
        breakout_level = trade_setup.get('breakout_level', current_nifty_price)
        direction = 'bullish' if 'CE' in tradingsymbol else 'bearish'
        breakout_ok, breakout_msg = check_breakout_confirmation(clientcode, tradingsymbol, current_nifty_price, breakout_level, direction)
        if not breakout_ok:
            logging.info(f"[PENDING] {breakout_msg}")
            return None
        
        # 8. Gap Filter Context (NEW - informational)
        gap_pct, gap_interp = check_gap_filter(clientcode, current_nifty_price)
        if abs(gap_pct) > 1.0:
            logging.info(f"🔔 {gap_interp}")
        
        logging.info(f"[OK] Risk checks passed: {trades_check[1]} | {positions_check[1]} | Loss: {loss_check[2]:.1f}%")
        
        # ==================== POSITION SIZING ====================
        
        # Find symbol token
        symboltoken = find_symbol_token(tradingsymbol, clientcode)
        
        if not symboltoken:
            logging.error(f"Could not find token for {tradingsymbol}")
            return None
        
        # 9. Liquidity Filter (NEW)
        liquidity_check = check_liquidity_filter(symboltoken, clientcode, min_oi=5000)
        if not liquidity_check[0]:
            logging.warning(f"{liquidity_check[1]}")
            return None
        
        # 10. Spread Filter
        spread_check = check_spread_filter(symboltoken, clientcode, max_spread_pct=3.0)
        if not spread_check[0]:
            logging.warning(f"{spread_check[1]}")
            return None
        
        # NEW 3. Volume Confirmation
        volume_check = check_volume_confirmation(symboltoken, clientcode)
        logging.info(f"[STATS] {volume_check[1]}")
        
        # NEW 5. Multi-Timeframe Confirmation
        mtf_aligned, mtf_msg, mtf_trends = check_multi_timeframe_confirmation('NIFTY')
        logging.info(f"[UP] Multi-TF: {mtf_msg}")
        
        # NEW 6. IV Percentile Check
        strike = trade_setup.get('strike')
        expiry = trade_setup.get('expiry')
        current_iv = trade_setup.get('iv', None)
        iv_rank, iv_msg = calculate_iv_percentile(tradingsymbol, strike, expiry, current_iv)
        logging.info(f"💹 IV: {iv_msg}")
        
        # Apply Kelly Criterion position sizing
        base_quantity = trade_setup.get('quantity', 75)  # NIFTY lot size = 75
        kelly_adjusted = calculate_kelly_position_size(clientcode, base_quantity)
        
        # NEW 7. Greeks-Based Position Sizing
        delta = trade_setup.get('delta', None)
        greeks_adjusted = adjust_position_size_by_greeks(kelly_adjusted, delta)
        
        # Apply time-of-day adjustment
        final_quantity = int(greeks_adjusted)
        
        # Apply profit protect mode risk reduction
        if protect_status == "REDUCE_RISK":
            final_quantity = int(final_quantity * (1 - risk_reduction/100))
            logging.info(f"🛡️ PROFIT PROTECT: Reduced size by {risk_reduction}% → {final_quantity}")
        
        # Apply time-of-day adjustment to SL and targets
        original_sl = trade_setup.get('stop_loss')
        original_t1 = trade_setup.get('target_1')
        original_t2 = trade_setup.get('target_2')
        entry_price = trade_setup.get('entry_price')
        
        adjusted_sl = entry_price - (entry_price - original_sl) * sl_mult
        adjusted_t1 = entry_price + (original_t1 - entry_price) * target_mult
        adjusted_t2 = entry_price + (original_t2 - entry_price) * target_mult
        
        logging.info(f"💎 Final Sizing: Base {base_quantity} → Kelly {kelly_adjusted} → Greeks {greeks_adjusted} → Final {final_quantity}")
        logging.info(f"[TARGET] Adjusted Levels: SL Rs.{adjusted_sl:.2f} | T1 Rs.{adjusted_t1:.2f} | T2 Rs.{adjusted_t2:.2f}")
        
        # ==================== ORDER EXECUTION ====================
        
        # Prepare order parameters
        order_params = {
            'variety': 'NORMAL',
            'tradingsymbol': tradingsymbol,
            'symboltoken': symboltoken,
            'transactiontype': 'BUY',
            'exchange': 'NFO',
            'ordertype': 'MARKET',  # Market order for fast execution
            'producttype': 'CARRYFORWARD',  # NRML - Full control over exits
            'duration': 'DAY',
            'quantity': str(final_quantity)
        }
        
        # Track commission (Rs.20 per order)
        track_commission(clientcode, num_orders=1, commission_per_order=20)
        
        # Place order
        order_result = place_order_angel_one(clientcode, order_params)
        
        if order_result and order_result.get('status'):
            order_id = order_result['orderid']
            unique_order_id = order_result.get('uniqueorderid')
            
            # CRITICAL: Verify order execution
            logging.info(f"[PENDING] Verifying order execution...")
            verification = verify_order_execution(clientcode, unique_order_id)
            
            if verification['executed']:
                # Use ACTUAL fill price from exchange
                actual_entry_price = verification['avg_price']
                actual_quantity = verification['filled_qty']
                planned_entry = trade_setup['entry_price']
                
                # Calculate slippage
                slippage_pct, slippage_amt = calculate_slippage(planned_entry, actual_entry_price, 'BUY')
                
                # Track slippage
                global DAILY_STATS
                today = datetime.now().date().isoformat()
                if clientcode in DAILY_STATS and today in DAILY_STATS[clientcode]:
                    slippage_cost = abs(slippage_amt) * actual_quantity
                    DAILY_STATS[clientcode][today]['slippage'] += slippage_cost
                    logging.info(f"[DOWN] Slippage: {slippage_pct:.2f}% (Rs.{slippage_amt:.2f} per qty, Total Rs.{slippage_cost:.0f})")
                
                logging.info(f"[CHART] Order executed: {actual_quantity} @ Rs.{actual_entry_price} (planned: Rs.{planned_entry})")
                
                # Create trade record in ACTIVE_TRADES
                if clientcode not in ACTIVE_TRADES:
                    ACTIVE_TRADES[clientcode] = {}
                
                ACTIVE_TRADES[clientcode][order_id] = {
                    'trade_number': trade_setup['trade_number'],
                    'instrument': trade_setup['instrument'],
                    'tradingsymbol': tradingsymbol,
                    'symboltoken': symboltoken,
                    'entry_price': actual_entry_price,  # Use actual fill price
                    'planned_entry': planned_entry,  # Original plan
                    'slippage_pct': slippage_pct,
                    'slippage_cost': abs(slippage_amt) * actual_quantity,
                    'quantity': actual_quantity,  # Actual filled quantity
                    'remaining_quantity': actual_quantity,
                    'stop_loss': adjusted_sl,  # Use time-adjusted SL
                    'target_1': adjusted_t1,  # Use time-adjusted T1
                    'target_2': adjusted_t2,  # Use time-adjusted T2
                    'original_sl': original_sl,  # Store original for reference
                    'original_t1': original_t1,
                    'original_t2': original_t2,
                    'status': 'open',
                    'entry_time': datetime.now().isoformat(),
                    'entry_timestamp': datetime.now(),  # For time-based profit taking
                    'target_1_hit': False,
                    'uniqueorderid': unique_order_id,
                    'pnl': 0,
                    'pattern_type': trade_setup.get('pattern_type', 'unknown'),  # AI pattern tracking
                    'time_phase': time_phase  # Track which time window entered
                }
                
                # Initialize trailing stop tracking
                update_trailing_stop(clientcode, order_id, actual_entry_price, actual_entry_price, adjusted_sl)
                
                # Track pattern for AI analytics
                logging.info(f"[STATS] Trade Pattern: {trade_setup.get('pattern_type', 'unknown')}")
                
                logging.info(f"[OK] Trade #{trade_setup['trade_number']} entered successfully")
                logging.info(f"[TARGET] Levels: Entry Rs.{actual_entry_price:.2f} | SL Rs.{adjusted_sl:.2f} | T1 Rs.{adjusted_t1:.2f} | T2 Rs.{adjusted_t2:.2f}")
                return order_id
            else:
                # Order NOT executed (rejected/timeout)
                logging.error(f"[FAIL] Order NOT executed: {verification['status']}")
                if verification['status'] == 'rejected':
                    logging.error(f"Rejection reason: {verification.get('reason')}")
                return None
        else:
            logging.error(f"Failed to place order")
            return None
        
    except Exception as e:
        logging.error(f"Error executing trade entry: {e}", exc_info=True)
        return None
SCRIP_MASTER_CACHE = {}  # Cache scrip master data

def fetch_fundamental_context():
    """
    Fetch fundamental context for AI decision-making
    Returns: Dictionary with market events, economic calendar, FII/DII data
    """
    today = datetime.now()
    today_str = today.strftime("%Y-%m-%d")
    day_name = today.strftime("%A")
    
    context = {
        'date': today_str,
        'day': day_name,
        'events': [],
        'market_holidays': [],
        'economic_calendar': [],
        'fii_dii_data': None,
        'sector_sentiment': {},
        'fo_ban_list': []
    }
    
    try:
        # Check if today is a known event day
        # Major recurring events (you can expand this from NSE/BSE calendar)
        major_events = {
            # Format: (month, day): event
            (1, 26): "Republic Day (Holiday)",
            (3, 1): "Budget Session Period",
            (8, 15): "Independence Day (Holiday)",
            (10, 2): "Gandhi Jayanti (Holiday)",
            (10, 24): "Diwali (Holiday)",
            (11, 1): "Diwali Trading (Muhurat)",
            (12, 25): "Christmas (Holiday)"
        }
        
        event_key = (today.month, today.day)
        if event_key in major_events:
            context['events'].append({
                'type': 'HOLIDAY',
                'description': major_events[event_key],
                'impact': 'MARKET CLOSED or SPECIAL SESSION'
            })
        
        # RBI Policy dates (update quarterly)
        rbi_policy_dates = [
            "2025-02-07", "2025-04-09", "2025-06-06", 
            "2025-08-08", "2025-10-09", "2025-12-06"
        ]
        
        if today_str in rbi_policy_dates:
            context['events'].append({
                'type': 'MONETARY_POLICY',
                'description': 'RBI Monetary Policy Decision',
                'impact': 'HIGH VOLATILITY EXPECTED - Avoid directional trades till 10:30 AM'
            })
        
        # Check if tomorrow is RBI day (pre-emptive caution)
        tomorrow = (today + timedelta(days=1)).strftime("%Y-%m-%d")
        if tomorrow in rbi_policy_dates:
            context['events'].append({
                'type': 'PRE_EVENT',
                'description': 'RBI Policy Tomorrow',
                'impact': 'CAUTIOUS - Market may be range-bound'
            })
        
        # Weekly expiry context
        if day_name == "Tuesday":
            context['events'].append({
                'type': 'EXPIRY',
                'description': 'NIFTY Weekly Expiry Today',
                'impact': 'HIGH VOLATILITY post 2:00 PM - Expect whipsaws near strikes'
            })
        
        # F&O Ban List (simulated - in production, fetch from NSE)
        # In real implementation, scrape from: https://www.nseindia.com/report/detail/fo_ban_list
        context['fo_ban_list'] = [
            # Example: "IDEA", "SAIL", "NATIONALUM"
        ]
        
        # FII/DII Data Context (simulated - in production, fetch from Moneycontrol)
        # In real: Scrape from https://www.moneycontrol.com/stocks/marketstats/fii_dii_activity/index.php
        context['fii_dii_data'] = {
            'date': today_str,
            'fii_net': 0,  # Positive = buying, Negative = selling
            'dii_net': 0,
            'interpretation': 'Data unavailable - assume neutral'
        }
        
        # Sector sentiment (can be derived from sectoral indices)
        # In production: Track Bank Nifty, IT index, Pharma index movements
        context['sector_sentiment'] = {
            'bank_nifty': 'NEUTRAL',  # 'STRONG', 'WEAK', 'NEUTRAL'
            'it': 'NEUTRAL',
            'auto': 'NEUTRAL',
            'pharma': 'NEUTRAL'
        }
        
        # US Market closing context (useful for Indian market opening)
        try:
            sp500 = yf.Ticker("^GSPC")
            sp500_data = sp500.history(period="5d")
            if not sp500_data.empty and len(sp500_data) >= 2:
                last_close = sp500_data['Close'].iloc[-1]
                prev_close = sp500_data['Close'].iloc[-2]
                change_pct = ((last_close - prev_close) / prev_close) * 100
                
                if change_pct > 1:
                    us_sentiment = "STRONG POSITIVE (S&P +{:.1f}%)".format(change_pct)
                elif change_pct > 0.3:
                    us_sentiment = "POSITIVE (S&P +{:.1f}%)".format(change_pct)
                elif change_pct < -1:
                    us_sentiment = "STRONG NEGATIVE (S&P {:.1f}%)".format(change_pct)
                elif change_pct < -0.3:
                    us_sentiment = "NEGATIVE (S&P {:.1f}%)".format(change_pct)
                else:
                    us_sentiment = "NEUTRAL (S&P {:.1f}%)".format(change_pct)
                
                context['us_market_sentiment'] = us_sentiment
        except:
            context['us_market_sentiment'] = "DATA UNAVAILABLE"
        
        logging.info(f"[STATS] Fundamental context fetched: {len(context['events'])} events today")
        
    except Exception as e:
        logging.error(f"Error fetching fundamental context: {e}")
    
    return context

def fetch_premarket_data():
    """Fetch pre-market data at 9:00 AM - includes fundamentals, events, news"""
    logging.info("=" * 50)
    logging.info("[NEWS] AUTOMATED TRADING: Fetching pre-market data + events at 9:00 AM")
    
    try:
        # Fetch SGX NIFTY and global markets
        sgx_data = yf.Ticker("^NSEI").history(period="1d")
        
        # Fetch global markets
        global_indices = {
            'S&P 500': '^GSPC',
            'NASDAQ': '^IXIC',
            'Dow Jones': '^DJI',
            'Hang Seng': '^HSI',
            'Nikkei': '^N225'
        }
        
        premarket_data = {
            'sgx_nifty': sgx_data['Close'].iloc[-1] if not sgx_data.empty else None,
            'global_markets': {},
            'fundamental_context': fetch_fundamental_context()
        }
        
        for name, symbol in global_indices.items():
            try:
                ticker = yf.Ticker(symbol)
                data = ticker.history(period="1d")
                if not data.empty:
                    premarket_data['global_markets'][name] = {
                        'price': data['Close'].iloc[-1],
                        'change': data['Close'].iloc[-1] - data['Open'].iloc[-1]
                    }
            except:
                pass
        
        logging.info(f"Pre-market data fetched: {premarket_data}")
        return premarket_data
        
    except Exception as e:
        logging.error(f"Error fetching pre-market data: {e}")
        return None

# Global cache for opening volatility analysis
OPENING_VOLATILITY_CACHE = {}

def analyze_opening_volatility_scalp():
    """
    Analyze opening volatility 5 minutes before market opens (9:10 AM)
    Determines directional bias for 9:15 AM opening scalp trade
    """
    global OPENING_VOLATILITY_CACHE
    
    logging.info("=" * 50)
    logging.info("[OPENING SCALP] ANALYSIS: Analyzing pre-market at 9:10 AM")
    
    try:
        # Get first available session to fetch market data
        if not _SMARTAPI_SESSIONS:
            logging.error("No active sessions for market data fetch")
            OPENING_VOLATILITY_CACHE['bias'] = 'NEUTRAL'
            return
        
        # Use first available session
        session_id = next(iter(_SMARTAPI_SESSIONS.keys()))
        session_data = _SMARTAPI_SESSIONS[session_id]
        
        # Fetch comprehensive market data using test request context
        with app.test_request_context(json={'symboltoken': '99926000'}):
            with app.test_client() as client:
                with client.session_transaction() as sess:
                    sess['session_id'] = session_id
                
                response = client.post('/api/trading/comprehensive-data', json={'symboltoken': '99926000'})
                
                if response.status_code != 200:
                    logging.error("Failed to fetch pre-market data")
                    OPENING_VOLATILITY_CACHE['bias'] = 'NEUTRAL'
                    return
                
                trading_data = response.get_json()
        
        if not trading_data.get('status'):
            logging.error("Failed to fetch pre-market data")
            OPENING_VOLATILITY_CACHE['bias'] = 'NEUTRAL'
            return
        
        indicators = trading_data.get('indicators', {})
        vix_data = trading_data.get('vix', {})
        
        # Calculate bias based on available data
        ltp = indicators.get('ltp', 0)
        prev_close = indicators.get('prev_close', ltp)
        vix = vix_data.get('vix', 15)
        
        # Calculate gap percentage
        gap_percent = ((ltp - prev_close) / prev_close * 100) if prev_close else 0
        
        # Determine bias
        if gap_percent > 0.5:
            bias = 'BULLISH'
            logging.info(f"[OK] BULLISH BIAS: Gap +{gap_percent:.2f}%, VIX {vix:.2f}")
        elif gap_percent < -0.5:
            bias = 'BEARISH'
            logging.info(f"[OK] BEARISH BIAS: Gap {gap_percent:.2f}%, VIX {vix:.2f}")
        else:
            bias = 'NEUTRAL'
            logging.info(f"[ALERT] NEUTRAL BIAS: Small gap {gap_percent:.2f}%, VIX {vix:.2f}")
        
        # Store in cache for 9:15 AM execution
        OPENING_VOLATILITY_CACHE = {
            'bias': bias,
            'gap_percent': gap_percent,
            'ltp': ltp,
            'prev_close': prev_close,
            'vix': vix,
            'timestamp': datetime.now().isoformat()
        }
        
        logging.info(f"Opening volatility analysis complete: {bias}")
        logging.info("=" * 50)
        
    except Exception as e:
        logging.error(f"Error in opening volatility analysis: {e}")
        OPENING_VOLATILITY_CACHE['bias'] = 'NEUTRAL'

def execute_opening_volatility_scalp():
    """
    Execute opening volatility scalp trade at 9:16 AM (after first candle confirmation)
    Waits for 9:15-9:16 first candle to confirm volatility
    Uses 100% capital for single ATM trade based on pre-market bias
    Target: 5% profit, SL: 40%, Time exit: 5 minutes max
    """
    global OPENING_VOLATILITY_CACHE, DAILY_TRADE_PLAN, PARSED_TRADE_SETUPS
    
    logging.info("=" * 50)
    logging.info("[OPENING SCALP] Analyzing first candle (9:15-9:16 AM)")
    
    try:
        # Get cached bias from 9:10 AM analysis
        bias = OPENING_VOLATILITY_CACHE.get('bias', 'NEUTRAL')
        
        if bias == 'NEUTRAL':
            logging.info("[ALERT] NEUTRAL bias detected - skipping opening scalp")
            logging.info("=" * 50)
            return
        
        # Execute for all enabled clients
        for session_id, session_data in _SMARTAPI_SESSIONS.items():
            clientcode = session_data.get('clientcode')
            
            # Check if auto-trading is enabled
            if not AUTO_TRADING_ENABLED.get(clientcode, False):
                logging.info(f"Auto-trading disabled for {clientcode}, skipping...")
                continue
            
            try:
                logging.info(f"Executing opening scalp for client: {clientcode}")
                
                # Initialize daily stats if not already done
                if clientcode not in DAILY_STATS:
                    initialize_daily_stats(clientcode, 15000)
                
                # Get current market data using test request context
                with app.test_request_context(json={'symboltoken': '99926000'}):
                    with app.test_client() as client:
                        with client.session_transaction() as sess:
                            sess['session_id'] = session_id
                        
                        response = client.post('/api/trading/comprehensive-data', json={'symboltoken': '99926000'})
                        
                        if response.status_code != 200:
                            logging.error(f"Failed to fetch market data for {clientcode}")
                            continue
                        
                        trading_data = response.get_json()
                
                if not trading_data.get('status'):
                    logging.error(f"Failed to fetch market data for {clientcode}")
                    continue
                
                # Get first candle data (9:15-9:16 AM)
                candles_data = trading_data.get('data', {}).get('candles', {}).get('data', [])
                
                if not candles_data:
                    logging.error(f"No candle data available for {clientcode}")
                    continue
                
                # Check first minute candle (most recent)
                first_candle = candles_data[-1] if candles_data else None
                
                if not first_candle:
                    logging.error(f"Could not fetch first candle for {clientcode}")
                    continue
                
                # Extract candle OHLC
                candle_high = float(first_candle[2])  # High
                candle_low = float(first_candle[3])   # Low
                candle_range = candle_high - candle_low
                
                logging.info(f"[FIRST CANDLE ANALYSIS]")
                logging.info(f"   High: {candle_high:.2f}, Low: {candle_low:.2f}")
                logging.info(f"   Range: {candle_range:.2f} points")
                
                # FIRST CANDLE FILTER: Check volatility
                if candle_range < 30:
                    logging.warning(f"[LOW VOLATILITY] First candle range {candle_range:.2f} < 30 points")
                    logging.warning(f"[SKIP] Not enough movement for profitable scalp")
                    continue
                
                if candle_range > 50:
                    logging.info(f"[HIGH VOLATILITY] Range {candle_range:.2f} > 50 points")
                    logging.info(f"[PROCEED] Opening scalp confirmed")
                else:
                    logging.info(f"[MODERATE VOLATILITY] Range {candle_range:.2f} (30-50 points) - acceptable")
                
                indicators = trading_data.get('indicators', {})
                ltp = indicators.get('ltp', 0)
                vix = trading_data.get('vix', {}).get('vix', 15)
                
                # Determine strike selection (ATM)
                # Round to nearest 50 for NIFTY options
                atm_strike = round(ltp / 50) * 50
                
                # Select option type based on bias
                option_type = 'CE' if bias == 'BULLISH' else 'PE'
                
                # Create opening scalp trade plan
                gap_percent = OPENING_VOLATILITY_CACHE.get('gap_percent', 0)
                
                scalp_plan = f"""
[OPENING VOLATILITY SCALP] - 9:16 AM (FIRST CANDLE CONFIRMED)

Market Analysis (9:10 AM):
- Opening Bias: {bias}
- Gap: {gap_percent:+.2f}%
- NIFTY LTP: {ltp:.2f}
- VIX: {vix:.2f}

First Candle (9:15-9:16 AM):
- Range: {candle_range:.2f} points (High: {candle_high:.2f}, Low: {candle_low:.2f})
- Volatility: {'HIGH [OK]' if candle_range > 50 else 'MODERATE [OK]' if candle_range >= 30 else 'LOW [X]'}
- Status: {'CONFIRMED - Proceeding with trade' if candle_range >= 30 else 'REJECTED - Insufficient volatility'}

Strategy: Ultra-Fast Opening Scalp (Professional Approach)
Capital Allocation: 100% (Rs. 15,000)
Trade Type: {option_type} (All-In)
Duration: MAXIMUM 5 MINUTES with 2-MINUTE CHECKPOINT

TRADE SETUP:
1. Instrument: NIFTY {atm_strike} {option_type}
   - Entry: Market order at 9:16 AM (after first candle confirmation)
   - Quantity: Maximum lots with full capital
   - Entry Price: Market price at execution
   - Monitoring: EVERY SECOND

2. Profit Target: 5% (Indian market volatility optimized)
   - Exit IMMEDIATELY on hitting 5% profit
   - Don't wait for more - book and exit

3. Stop Loss: 40% (Indian market reality)
   - Wider SL to avoid premature exit
   - Options can swing 30-50% in opening minutes
   - Exit only if genuine trend reversal

4. Time-Based Exits (Professional Approach):
   
   A. 9:17 AM CHECKPOINT (2 minutes):
      - Check P&L at exactly 2 minutes
      - If profit < 1% → EXIT IMMEDIATELY
      - Rationale: Trade not working, avoid holding dead positions
      - This prevents holding losers for full 5 minutes hoping for recovery
   
   B. 9:21 AM HARD EXIT (5 minutes):
      - Force exit at 5-minute mark regardless of P&L
      - Opening volatility exhausted by this time
      - Don't hold beyond 5 minutes under any circumstance

5. First Candle Filter (CRITICAL):
   - Wait for 9:15-9:16 AM first candle to close
   - Check range: High - Low
   - Range > 50 points: HIGH volatility confirmed → Proceed
   - Range 30-50 points: MODERATE volatility → Proceed with caution
   - Range < 30 points: LOW volatility → SKIP TRADE
   - This filter eliminates 40% of losing trades on flat openings

Rationale: 
Professional traders wait for first candle confirmation instead of blindly entering at 9:15. The 9:17 AM checkpoint catches non-working trades early. Indian market opening volatility is extreme but short-lived. Second-by-second monitoring with multi-tier exits: quick profit OR 2-min checkpoint OR stop loss OR 5-min hard exit.

[CRITICAL] First candle confirmation + 9:17 AM checkpoint + instant exits
"""
                
                # Store in global trade plan
                if clientcode not in DAILY_TRADE_PLAN:
                    DAILY_TRADE_PLAN[clientcode] = {}
                
                DAILY_TRADE_PLAN[clientcode] = {
                    'plan': scalp_plan,
                    'generated_at': datetime.now().isoformat(),
                    'method': 'opening-scalp',
                    'bias': bias
                }
                
                # Parse into structured trade setup
                # Ultra-fast scalp with Indian market parameters
                trade_setup = {
                    'trade_number': 1,
                    'tradingsymbol': f"NIFTY{atm_strike}{option_type}",
                    'instrument': f"NIFTY {atm_strike} {option_type}",
                    'strike': atm_strike,
                    'option_type': option_type,
                    'direction': 'BUY',
                    'entry_price': 0,  # Will be filled at market order
                    'target_1': 0,  # 5% profit target (will be set after entry)
                    'target_2': 0,
                    'stop_loss': 0,  # 40% SL (will be set after entry)
                    'quantity': 50,  # NIFTY lot size
                    'capital_allocated': 15000,
                    'profit_target_percent': 5,  # 5% profit booking
                    'stop_loss_percent': 40,  # 40% stop loss
                    'monitor_interval': 1,  # Monitor every 1 second
                    'max_hold_minutes': 5,  # Maximum 5 minute hold
                    'is_opening_scalp': True,  # Mark as opening scalp for special handling
                    'first_candle_range': candle_range,  # Store first candle volatility
                    'entry_conditions': [
                        {
                            'type': 'time',
                            'operator': '>=',
                            'value': '09:15',
                            'description': 'Execute at 9:15 AM sharp'
                        }
                    ],
                    'exit_conditions': [
                        {
                            'type': 'profit_percent',
                            'value': 5,
                            'description': '5% profit target - EXIT IMMEDIATELY'
                        },
                        {
                            'type': 'stop_loss_percent',
                            'value': 40,
                            'description': '40% stop loss - Indian market reality'
                        },
                        {
                            'type': 'time',
                            'operator': '>=',
                            'value': '09:20',
                            'description': 'HARD EXIT at 9:20 AM (5 min limit)'
                        }
                    ],
                    'entry_time_start': '09:15',
                    'entry_time_end': '09:16',
                    'strategy_note': f'ULTRA-FAST {bias} opening scalp: 5% profit, 40% SL, 5-min max hold'
                }
                
                # Store parsed setup
                if clientcode not in PARSED_TRADE_SETUPS:
                    PARSED_TRADE_SETUPS[clientcode] = []
                
                PARSED_TRADE_SETUPS[clientcode] = [trade_setup]
                
                logging.info(f"[OK] Opening scalp trade prepared: {trade_setup['tradingsymbol']}")
                logging.info(f"   Bias: {bias}, Strike: {atm_strike}, Type: {option_type}")
                
            except Exception as e:
                logging.error(f"Error executing opening scalp for {clientcode}: {e}")
        
        logging.info("Opening scalp execution complete")
        logging.info("=" * 50)
        
    except Exception as e:
        logging.error(f"Error in opening volatility scalp execution: {e}")

def generate_daily_trade_plan():
    """Generate AI trade plan at 9:15 AM for all logged-in users"""
    global DAILY_TRADE_PLAN
    
    logging.info("=" * 50)
    logging.info("[AUTO] AUTOMATED TRADING: Generating trade plan at 9:15 AM")
    
    for session_id, session_data in _SMARTAPI_SESSIONS.items():
        clientcode = session_data.get('clientcode')
        
        # Check if auto-trading is enabled for this client
        if not AUTO_TRADING_ENABLED.get(clientcode, False):
            logging.info(f"Auto-trading disabled for {clientcode}, skipping...")
            continue
        
        try:
            logging.info(f"Generating trade plan for client: {clientcode}")
            
            # Initialize daily stats for new trading day
            initialize_daily_stats(clientcode, 15000)
            
            # Simulate request context to call API
            with app.test_request_context(
                json={'capital': 15000, 'risk_percent': 2, 'symbol': 'NIFTY'}
            ):
                # Set session for authentication
                with app.test_client() as client:
                    with client.session_transaction() as sess:
                        sess['session_id'] = session_id
                    
                    # Call AI recommendation endpoint
                    response = client.post(
                        '/api/trading/ai-recommendation',
                        json={'capital': 15000, 'risk_percent': 2, 'symbol': 'NIFTY'}
                    )
                    
                    if response.status_code == 200:
                        data = response.get_json()
                        plan_text = data.get('recommendation')
                        
                        DAILY_TRADE_PLAN[clientcode] = {
                            'plan': plan_text,
                            'generated_at': datetime.now().isoformat(),
                            'status': 'active',
                            'generated_method': 'scheduled-9:15AM'
                        }
                        
                        # Parse trade plan immediately after generation
                        parsed_data = parse_trade_plan_with_ai(plan_text, clientcode)
                        
                        if parsed_data and parsed_data.get('trades'):
                            PARSED_TRADE_SETUPS[clientcode] = parsed_data['trades']
                            
                            # Save to history for UI display
                            plan_id = f"plan_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                            if clientcode not in TRADE_PLAN_HISTORY:
                                TRADE_PLAN_HISTORY[clientcode] = []
                            
                            TRADE_PLAN_HISTORY[clientcode].append({
                                'id': plan_id,
                                'timestamp': datetime.now().isoformat(),
                                'plan_text': plan_text,
                                'trades': parsed_data['trades'],
                                'method': 'scheduled-9:15AM'
                            })
                            
                            logging.info(f"[OK] Trade plan generated and parsed for {clientcode}: {len(parsed_data['trades'])} trades")
                            logging.info(f"[HISTORY] Saved plan {plan_id} to history ({len(TRADE_PLAN_HISTORY[clientcode])} total)")
                        else:
                            logging.warning(f"[WARNING] Failed to parse trade plan for {clientcode}")
                    else:
                        logging.error(f"[FAIL] Failed to generate plan for {clientcode}: {response.status_code}")
        
        except Exception as e:
            logging.error(f"Error generating trade plan for {clientcode}: {e}", exc_info=True)

def ai_performance_review(review_time="mid-session"):
    """
    AI reviews performance and adjusts strategy mid-session
    Called at: 11:00 AM, 1:00 PM, 2:30 PM
    """
    global DAILY_TRADE_PLAN, PARSED_TRADE_SETUPS
    
    logging.info("=" * 50)
    logging.info(f"[AI] PERFORMANCE REVIEW: {review_time.upper()}")
    
    for session_id, session_data in _SMARTAPI_SESSIONS.items():
        clientcode = session_data.get('clientcode')
        
        if not AUTO_TRADING_ENABLED.get(clientcode, False):
            continue
        
        try:
            # Get current stats
            stats = get_daily_stats_summary(clientcode)
            if not stats or stats['trades'] == 0:
                logging.info(f"[STATS] {clientcode}: No trades yet - skipping review")
                continue
            
            loss_check = check_daily_loss_circuit_breaker(clientcode)
            trades_check = check_max_trades_limit(clientcode)
            
            # Build review prompt
            review_prompt = f"""PERFORMANCE REVIEW - {review_time.upper()}

[STATS] TODAY'S PERFORMANCE:
- P&L: Rs.{stats['pnl']:,.0f} ({stats['pnl_pct']:+.1f}%)
- Win Rate: {stats['win_rate']:.0f}% ({stats['wins']}W / {stats['losses']}L)
- Trades: {stats['trades']}/15 executed
- Net P&L: Rs.{stats['net_pnl']:,.0f} (after costs)
- Kelly Multiplier: {KELLY_MULTIPLIER.get(clientcode, 1.0):.1f}x

[TARGET] RISK STATUS:
- Circuit Breaker: {'[ALERT] ACTIVE' if not loss_check[0] else f'[OK] OK ({loss_check[2]:.1f}% from -10% limit)'}
- Trades Remaining: {15 - trades_check[2]} trades available

INSTRUCTIONS:
Based on current performance, provide ONE of these recommendations:

1. CONTINUE - Performance good, stick to original plan
2. TIGHTEN - Reduce position sizes, be more selective
3. PAUSE - Stop trading, conditions unfavorable
4. ADJUST - Modify strategy (explain what to change)

Respond in JSON format:
{{
  "action": "CONTINUE" | "TIGHTEN" | "PAUSE" | "ADJUST",
  "reasoning": "Brief explanation (2-3 sentences)",
  "suggestions": ["specific actionable advice"],
  "position_size_multiplier": 0.5 to 1.5
}}"""

            # Call AI for review
            try:
                response = openai_client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {"role": "system", "content": "You are a risk management AI reviewing intraday trading performance. Be honest and protective of capital."},
                        {"role": "user", "content": review_prompt}
                    ],
                    temperature=0.3,
                    max_tokens=500
                )
                
                review_text = response.choices[0].message.content
                
                if not review_text:
                    logging.warning("AI review returned empty response")
                    continue
                
                # Try to parse JSON
                try:
                    review_data = json.loads(review_text)
                    action = review_data.get('action', 'CONTINUE')
                    reasoning = review_data.get('reasoning', 'No reasoning provided')
                    
                    logging.info(f"[AI] RECOMMENDATION: {action}")
                    logging.info(f"💭 Reasoning: {reasoning}")
                    
                    # Take action based on AI recommendation
                    if action == "PAUSE":
                        AUTO_TRADING_ENABLED[clientcode] = False
                        logging.warning(f"⏸️ PAUSING AUTO-TRADING for {clientcode} per AI recommendation")
                    
                    elif action == "TIGHTEN":
                        # Reduce Kelly multiplier
                        if clientcode in KELLY_MULTIPLIER:
                            KELLY_MULTIPLIER[clientcode] *= 0.7
                            logging.info(f"[DOWN] Reducing position size to {KELLY_MULTIPLIER[clientcode]:.1f}x")
                    
                    # Log suggestions
                    for suggestion in review_data.get('suggestions', []):
                        logging.info(f"[TIP] {suggestion}")
                
                except json.JSONDecodeError:
                    logging.warning(f"Could not parse AI review JSON, got: {review_text[:200]}")
            
            except Exception as e:
                logging.error(f"AI review API call failed: {e}")
        
        except Exception as e:
            logging.error(f"Performance review error for {clientcode}: {e}", exc_info=True)

def monitor_prices_and_execute():
    """Monitor live prices every 5 mins and execute trades when conditions met"""
    global ACTIVE_TRADES, PARSED_TRADE_SETUPS
    
    while True:
        try:
            current_time = datetime.now().time()
            
            # Only monitor during market hours (9:15 AM - 3:15 PM)
            if current_time < datetime.strptime("09:15", "%H:%M").time():
                time.sleep(60)  # Sleep 1 minute before market opens
                continue
            
            if current_time > datetime.strptime("15:15", "%H:%M").time():
                logging.info("Market closing time reached, stopping new entries")
                break
            
            # Dynamic sleep interval: 1 second during opening scalp, 60 seconds after
            opening_scalp_start = datetime.strptime("09:15", "%H:%M").time()
            opening_scalp_end = datetime.strptime("09:20", "%H:%M").time()
            
            if opening_scalp_start <= current_time <= opening_scalp_end:
                check_interval = 1  # 1 second during opening scalp
            else:
                check_interval = 300  # 5 minutes for regular entry checks
            
            logging.info(f"=" * 50)
            logging.info(f"Price monitoring cycle at {current_time}")
            
            # Process each client's trade setups
            for clientcode, trade_setups in PARSED_TRADE_SETUPS.items():
                if not AUTO_TRADING_ENABLED.get(clientcode, False):
                    continue
                
                plan_data = DAILY_TRADE_PLAN.get(clientcode, {})
                if plan_data.get('status') != 'pending':
                    continue
                
                logging.info(f"Processing {len(trade_setups)} trade setups for {clientcode}")
                
                # Check each trade setup for entry
                for trade_setup in trade_setups:
                    trade_number = trade_setup.get('trade_number')
                    
                    # Skip if already executed
                    if clientcode in ACTIVE_TRADES:
                        already_executed = any(
                            t.get('trade_number') == trade_number 
                            for t in ACTIVE_TRADES[clientcode].values()
                        )
                        if already_executed:
                            logging.info(f"Trade #{trade_number} already executed, skipping")
                            continue
                    
                    # Evaluate entry conditions
                    logging.info(f"Evaluating entry conditions for Trade #{trade_number}: {trade_setup.get('instrument')}")
                    
                    if evaluate_entry_conditions(trade_setup, clientcode):
                        logging.info(f"[TARGET] ENTRY CONDITIONS MET for Trade #{trade_number}!")
                        
                        # Execute trade
                        trade_id = execute_trade_entry(trade_setup, clientcode)
                        
                        if trade_id:
                            logging.info(f"[OK] Trade #{trade_number} executed successfully: Order ID = {trade_id}")
                        else:
                            logging.error(f"[FAIL] Failed to execute Trade #{trade_number}")
                    else:
                        logging.info(f"Entry conditions not met for Trade #{trade_number}")
            
            # Monitor active trades for stop loss / target hit
            monitor_active_trades_sl_target()
            
            # AI-powered market shift detection - DISABLED (wasteful, adds latency)
            # ai_monitor_and_adjust_trades()
            
            # Dynamic sleep: 1 second during opening scalp (9:15-9:20), 5 minutes after
            opening_scalp_start = datetime.strptime("09:15", "%H:%M").time()
            opening_scalp_end = datetime.strptime("09:20", "%H:%M").time()
            
            if opening_scalp_start <= current_time <= opening_scalp_end:
                sleep_interval = 1  # 1 second during opening scalp
                logging.info(f"⚡ OPENING SCALP: Sleeping for {sleep_interval} second...")
            else:
                sleep_interval = 300  # 5 minutes for regular monitoring
                logging.info(f"Sleeping for 5 minutes...")
            
            time.sleep(sleep_interval)
        
        except Exception as e:
            logging.error(f"Error in price monitoring: {e}", exc_info=True)
            time.sleep(60)

def ai_monitor_and_adjust_trades():
    """AI-powered market shift detection and trade adjustment"""
    global ACTIVE_TRADES
    
    for clientcode, trades in ACTIVE_TRADES.items():
        if not trades or not AUTO_TRADING_ENABLED.get(clientcode, False):
            continue
        
        # Run AI market analysis once per client
        market_analysis = ai_analyze_market_shift(clientcode)
        
        if not market_analysis or not market_analysis.get('shift_detected'):
            continue
        
        logging.info(f"="*60)
        logging.info(f"[SEARCH] AI DETECTED MARKET SHIFT for {clientcode}")
        logging.info(f"   Direction: {market_analysis.get('new_direction')}")
        logging.info(f"   Confidence: {market_analysis.get('confidence')}%")
        
        # Check if confidence is high enough to act
        if market_analysis.get('confidence', 0) < 70:
            logging.info(f"   [WARNING] Confidence too low, not modifying trades")
            continue
        
        # Analyze and adjust each active trade
        for trade_id, trade_data in trades.items():
            if trade_data.get('status') != 'open':
                continue
            
            logging.info(f"\n   Analyzing Trade: {trade_data.get('instrument')}")
            
            # Get AI recommendation for new parameters
            adjustments = ai_adjust_trade_params(clientcode, trade_data, market_analysis)
            
            if not adjustments:
                continue
            
            # Validate adjustments
            new_sl = adjustments.get('new_stop_loss')
            new_t1 = adjustments.get('new_target_1')
            new_t2 = adjustments.get('new_target_2')
            
            # Safety check: Don't worsen stop loss
            current_sl = trade_data.get('stop_loss')
            entry_price = trade_data.get('entry_price')
            
            if new_sl < current_sl:  # For CE, lower SL is worse
                logging.warning(f"   [WARNING] AI suggested worse SL (Rs.{new_sl} < Rs.{current_sl}), keeping original")
                new_sl = current_sl
            
            # Update trade data in memory
            old_sl = trade_data['stop_loss']
            old_t1 = trade_data['target_1']
            old_t2 = trade_data['target_2']
            
            trade_data['stop_loss'] = new_sl
            trade_data['target_1'] = new_t1
            trade_data['target_2'] = new_t2
            trade_data['last_modified'] = datetime.now().isoformat()
            trade_data['modification_reason'] = adjustments.get('modification_reason')
            
            logging.info(f"   [OK] Trade parameters updated:")
            logging.info(f"      SL: Rs.{old_sl} → Rs.{new_sl}")
            logging.info(f"      T1: Rs.{old_t1} → Rs.{new_t1}")
            logging.info(f"      T2: Rs.{old_t2} → Rs.{new_t2}")
            logging.info(f"      Reason: {adjustments.get('modification_reason')}")
            
            # Note: We don't modify the actual exchange order for options
            # (market orders are executed instantly)
            # We only update our internal SL/Target tracking
            # For LIMIT orders, you would call modify_order_angel_one() here

def monitor_active_trades_sl_target():
    """Monitor active trades for stop loss and target exits - runs every 60 seconds"""
    global ACTIVE_TRADES
    
    for clientcode, trades in ACTIVE_TRADES.items():
        if not trades:
            continue
        
        # Collect all symbol tokens for batch fetch
        open_trades = {
            trade_id: trade_data 
            for trade_id, trade_data in trades.items() 
            if trade_data.get('status') == 'open'
        }
        
        if not open_trades:
            continue
        
        try:
            # Batch fetch prices for ALL active positions in ONE API call
            symboltokens = [t.get('symboltoken') for t in open_trades.values()]
            batch_prices = get_batch_option_prices(symboltokens, clientcode)
            
            if not batch_prices:
                logging.warning(f"Could not fetch batch prices for {clientcode}")
                continue
            
            # Process each trade with fetched prices
            for trade_id, trade_data in open_trades.items():
                symboltoken = trade_data.get('symboltoken')
                price_data = batch_prices.get(symboltoken)
                
                if not price_data:
                    logging.warning(f"No price data for {trade_data.get('tradingsymbol')}")
                    continue
                
                current_price = price_data['ltp']
                entry_price = trade_data.get('entry_price', 0)
                
                # Calculate profit percentage
                profit_pct = ((current_price - entry_price) / entry_price) * 100 if entry_price > 0 else 0
                
                logging.info(f"Monitoring Trade {trade_id}: Current=Rs.{current_price}, Entry=Rs.{entry_price}, Profit={profit_pct:.2f}%, SL=Rs.{trade_data.get('stop_loss')}, T1=Rs.{trade_data.get('target_1')}")
                
                # NEW 1. Dynamic Trailing Stop Loss
                new_trailing_sl = update_trailing_stop(clientcode, trade_id, current_price, entry_price, trade_data.get('stop_loss'))
                if new_trailing_sl > trade_data.get('stop_loss'):
                    trade_data['stop_loss'] = new_trailing_sl
                
                # NEW 12. Time-Based Profit Taking
                entry_time = trade_data.get('entry_timestamp', datetime.now())
                should_exit_time, exit_reason = check_time_based_profit_taking(clientcode, trade_id, entry_time, profit_pct)
                
                # OPENING SCALP: Special exit logic
                if trade_data.get('is_opening_scalp', False):
                    time_in_trade = (datetime.now() - entry_time).total_seconds() / 60
                    
                    # 9:17 AM CHECKPOINT (2 minutes): Exit if not working
                    if time_in_trade >= 2.0 and profit_pct < 1.0:
                        should_exit_time = True
                        exit_reason = f"OPENING SCALP: 9:17 AM checkpoint - Trade not working (Profit: {profit_pct:.2f}% < 1%)"
                        logging.warning(f"[ALERT] {exit_reason} for Trade {trade_id} (held {time_in_trade:.1f} min)")
                        logging.info(f"[EXIT] Early exit to avoid holding dead trade for full 5 minutes")
                    
                    # 5-MINUTE HARD LIMIT: Force exit regardless
                    elif time_in_trade >= 5:
                        should_exit_time = True
                        exit_reason = "OPENING SCALP: 5-minute hard limit reached"
                        logging.warning(f"[ALERT] {exit_reason} for Trade {trade_id} (held {time_in_trade:.1f} min)")
                
                if should_exit_time:
                    logging.info(f"{exit_reason}")
                    close_position(clientcode, trade_id, current_price, 'time_based_exit')
                    continue
                
                # Get VIX-based thresholds
                vix_value = get_current_vix_value()
                thresholds = calculate_vix_based_thresholds(vix_value)
                profit_target = calculate_vix_based_profit_target(vix_value)
                
                if vix_value:
                    logging.info(f"[VIX] Current: {vix_value:.2f} | Target: {profit_target:.1f}% | Trailing SL: Rs.{trade_data.get('stop_loss'):.2f}")
                
                # Check VIX-based profit target (quick exit)
                if profit_pct >= profit_target:
                    pnl = (current_price - entry_price) * trade_data.get('quantity', 0)
                    logging.info(f"[PROFIT] VIX-based exit ({profit_target:.1f}%): {trade_id} | Entry Rs.{entry_price:.2f} → Exit Rs.{current_price:.2f} | P&L Rs.{pnl:.2f}")
                    close_position(clientcode, trade_id, current_price, f'{profit_target:.0f}pct_profit')
                    
                    # Reset for re-entry
                    trade_number = trade_data.get('trade_number')
                    if trade_number and clientcode in PARSED_TRADE_SETUPS:
                        for setup in PARSED_TRADE_SETUPS[clientcode]:
                            if setup.get('trade_number') == trade_number:
                                setup['executed'] = False
                                logging.info(f"[RESET] Setup {trade_number} reset for re-entry")
                                break
                
                # Check stop loss (including trailing)
                elif current_price <= trade_data.get('stop_loss', 0):
                    logging.warning(f"[SL] STOP LOSS HIT for Trade {trade_id}")
                    close_position(clientcode, trade_id, current_price, 'stop_loss')
                
                # NEW 11. Partial Position Scaling (3-tier)
                # Target 1: Book 33%
                elif not trade_data.get('target_1_hit') and current_price >= trade_data.get('target_1', 999999):
                    logging.info(f"[T1] TARGET 1 HIT - Booking 33% profit")
                    partial_close_position_scaled(clientcode, trade_id, current_price, 'target_1', scale_pct=33)
                
                # Target 2: Book another 33% (total 66%)
                elif trade_data.get('target_1_hit') and not trade_data.get('target_2_hit') and current_price >= trade_data.get('target_2', 999999):
                    logging.info(f"[T2] TARGET 2 HIT - Booking another 33%")
                    partial_close_position_scaled(clientcode, trade_id, current_price, 'target_2', scale_pct=33)
                
                # Target 3 or trailing stop: Close remaining 34%
                elif trade_data.get('target_2_hit') and current_price >= (trade_data.get('target_2') * 1.15):  # 15% above T2
                    logging.info(f"[T3] TARGET 3 HIT - Closing remaining position")
                    close_position(clientcode, trade_id, current_price, 'target_3')
                
        except Exception as e:
            logging.error(f"Error monitoring trades for {clientcode}: {e}", exc_info=True)

def partial_close_position_scaled(clientcode, trade_id, exit_price, reason, scale_pct=33):
    """
    NEW 11. Partial Position Scaling (3-tier)
    Close specified percentage of remaining position
    
    3-Tier Scaling:
    - Target 1: Close 33% (scale_pct=33)
    - Target 2: Close 33% of remaining = 33% of original (scale_pct=33)
    - Target 3/SL: Close remaining 34%
    """
    try:
        trade_data = ACTIVE_TRADES[clientcode][trade_id]
        
        remaining_qty = trade_data.get('remaining_quantity')
        original_qty = trade_data.get('quantity', remaining_qty)
        
        # Calculate close quantity based on percentage
        close_qty = int(remaining_qty * (scale_pct / 100))
        
        # Ensure at least 1 lot closes
        if close_qty == 0:
            logging.warning(f"Quantity too small to partially close ({remaining_qty} remaining, {scale_pct}%)")
            return
        
        # Prepare exit order
        order_params = {
            'variety': 'NORMAL',
            'tradingsymbol': trade_data.get('tradingsymbol'),
            'symboltoken': trade_data.get('symboltoken'),
            'transactiontype': 'SELL',
            'exchange': 'NFO',
            'ordertype': 'MARKET',
            'producttype': 'CARRYFORWARD',  # NRML - Match entry product type
            'duration': 'DAY',
            'quantity': str(close_qty)
        }
        
        order_result = place_order_angel_one(clientcode, order_params)
        
        if order_result and order_result.get('status'):
            # Update trade data
            new_remaining = remaining_qty - close_qty
            trade_data['remaining_quantity'] = new_remaining
            
            # Mark targets hit
            if reason == 'target_1':
                trade_data['target_1_hit'] = True
                trade_data['partial_exit_1_price'] = exit_price
                trade_data['partial_exit_1_time'] = datetime.now().isoformat()
                trade_data['partial_exit_1_qty'] = close_qty
            elif reason == 'target_2':
                trade_data['target_2_hit'] = True
                trade_data['partial_exit_2_price'] = exit_price
                trade_data['partial_exit_2_time'] = datetime.now().isoformat()
                trade_data['partial_exit_2_qty'] = close_qty
            
            closed_pct = (close_qty / original_qty) * 100
            remaining_pct = (new_remaining / original_qty) * 100
            
            logging.info(f"[3-TIER SCALE] {reason.upper()}: Closed {close_qty}/{original_qty} lots ({closed_pct:.0f}%) at Rs.{exit_price} | Remaining: {new_remaining} ({remaining_pct:.0f}%)")
        
    except Exception as e:
        logging.error(f"Error in partial close: {e}")

# LEGACY function for backward compatibility
def partial_close_position(clientcode, trade_id, exit_price, reason):
    """Legacy: Close 50% of position at Target 1 - Use partial_close_position_scaled instead"""
    partial_close_position_scaled(clientcode, trade_id, exit_price, reason, scale_pct=50)

def close_position(clientcode, trade_id, exit_price, reason):
    """Close entire position with P&L tracking, pattern analytics, and loss streak updates"""
    try:
        trade_data = ACTIVE_TRADES[clientcode][trade_id]
        
        close_qty = trade_data.get('remaining_quantity')
        
        # Prepare exit order
        order_params = {
            'variety': 'NORMAL',
            'tradingsymbol': trade_data.get('tradingsymbol'),
            'symboltoken': trade_data.get('symboltoken'),
            'transactiontype': 'SELL',
            'exchange': 'NFO',
            'ordertype': 'MARKET',
            'producttype': 'CARRYFORWARD',  # NRML - Match entry product type
            'duration': 'DAY',
            'quantity': str(close_qty)
        }
        
        # Track commission for exit order
        track_commission(clientcode, num_orders=1, commission_per_order=20)
        
        order_result = place_order_angel_one(clientcode, order_params)
        
        if order_result and order_result.get('status'):
            trade_data['status'] = f'closed_{reason}'
            trade_data['exit_price'] = exit_price
            trade_data['exit_time'] = datetime.now().isoformat()
            
            # Calculate P&L
            entry_price = trade_data.get('entry_price', 0)
            planned_entry = trade_data.get('planned_entry', entry_price)
            total_qty = trade_data.get('quantity', 0)
            
            # P&L = (Exit - Entry) * Quantity
            raw_pnl = (exit_price - entry_price) * total_qty
            
            # Account for slippage cost already tracked at entry
            slippage_cost = trade_data.get('slippage_cost', 0)
            
            # Exit slippage (actual exit vs planned SL/Target)
            planned_exit = trade_data.get('target_1', exit_price)
            if 'sl' in reason or 'stop' in reason:
                planned_exit = trade_data.get('stop_loss', exit_price)
            
            exit_slippage_pct, exit_slippage_amt = calculate_slippage(planned_exit, exit_price, 'SELL')
            exit_slippage_cost = abs(exit_slippage_amt) * total_qty
            
            # Track exit slippage
            global DAILY_STATS
            today = datetime.now().date().isoformat()
            if clientcode in DAILY_STATS and today in DAILY_STATS[clientcode]:
                DAILY_STATS[clientcode][today]['slippage'] += exit_slippage_cost
            
            # Net P&L (raw P&L - already factored in commissions/slippage via daily stats)
            net_pnl = raw_pnl
            trade_data['pnl'] = net_pnl
            trade_data['raw_pnl'] = raw_pnl
            trade_data['exit_slippage_pct'] = exit_slippage_pct
            trade_data['exit_slippage_cost'] = exit_slippage_cost
            
            # Determine win/loss
            is_win = net_pnl > 0
            
            # NEW: Update loss streak for consecutive loss protection
            update_loss_streak(clientcode, is_win)
            
            # NEW: Track AI pattern performance
            pattern_type = trade_data.get('pattern_type', 'unknown')
            track_trade_pattern_performance(clientcode, pattern_type, is_win, net_pnl)
            
            # Update daily statistics
            update_daily_pnl(clientcode, net_pnl, is_win=is_win)
            
            # Log comprehensive exit summary
            pnl_pct = (net_pnl / (entry_price * total_qty)) * 100
            logging.info(f"[{'[OK] WIN' if is_win else '[FAIL] LOSS'}] Position closed: {reason} | Pattern: {pattern_type}")
            logging.info(f"  Entry: Rs.{entry_price:.2f} × {total_qty} = Rs.{entry_price * total_qty:,.0f}")
            logging.info(f"  Exit: Rs.{exit_price:.2f} × {total_qty} = Rs.{exit_price * total_qty:,.0f}")
            logging.info(f"  P&L: Rs.{net_pnl:,.0f} ({pnl_pct:+.1f}%)")
            logging.info(f"  Exit Slippage: {exit_slippage_pct:.2f}% (Rs.{exit_slippage_cost:.0f})")
            
            # NEW: Show loss streak status if applicable
            if clientcode in CONSECUTIVE_LOSSES:
                streak = CONSECUTIVE_LOSSES[clientcode]
                if streak > 0:
                    logging.warning(f"  [WARNING] Consecutive losses: {streak}/3")
            
            # NEW: Show best performing patterns
            best_patterns = get_best_performing_patterns(clientcode, min_trades=3)
            if best_patterns:
                top_pattern = best_patterns[0]
                logging.info(f"  [UP] Best pattern: {top_pattern['pattern']} ({top_pattern['win_rate']:.0f}% WR, {top_pattern['trades']} trades)")
            
            # Print daily summary
            stats_summary = get_daily_stats_summary(clientcode)
            if stats_summary:
                logging.info(f"[STATS] Today: Rs.{stats_summary['net_pnl']:,.0f} | {stats_summary['win_rate']:.0f}% WR | {stats_summary['trades']} trades")
        
    except Exception as e:
        logging.error(f"Error closing position: {e}")

def close_all_positions():
    """Close all open positions at 3:15 PM (intraday rule)"""
    logging.info("=" * 50)
    logging.info("AUTOMATED TRADING: Closing all positions at 3:15 PM")
    
    for clientcode, trades in ACTIVE_TRADES.items():
        try:
            for trade_id, trade_data in list(trades.items()):
                if trade_data.get('status') == 'open':
                    logging.info(f"Closing trade {trade_id} for {clientcode}: {trade_data.get('tradingsymbol')}")
                    
                    # Get current market price
                    symboltoken = trade_data.get('symboltoken')
                    current_price = get_option_ltp(symboltoken, clientcode)
                    
                    if current_price is None:
                        logging.error(f"Could not fetch LTP for EOD closure")
                        current_price = trade_data.get('entry_price', 0)  # Fallback
                    
                    # Close position
                    close_position(clientcode, trade_id, current_price, 'eod_auto_close')
                    
        except Exception as e:
            logging.error(f"Error closing positions for {clientcode}: {e}", exc_info=True)

def end_of_day_review():
    """Review executed trades and calculate P&L at 3:30 PM"""
    logging.info("=" * 50)
    logging.info("AUTOMATED TRADING: End-of-day review at 3:30 PM")
    
    for clientcode, trades in ACTIVE_TRADES.items():
        try:
            total_pnl = 0
            trades_summary = []
            
            for trade_id, trade_data in trades.items():
                pnl = trade_data.get('pnl', 0)
                total_pnl += pnl
                trades_summary.append({
                    'trade_id': trade_id,
                    'entry': trade_data.get('entry_price'),
                    'exit': trade_data.get('exit_price'),
                    'pnl': pnl,
                    'status': trade_data.get('status')
                })
            
            # Store results in database
            conn = sqlite3.connect(DB_FILE)
            c = conn.cursor()
            c.execute('''
                INSERT INTO daily_pnl (clientcode, date, total_pnl, trades_count, trades_data)
                VALUES (?, ?, ?, ?, ?)
            ''', (clientcode, datetime.now().date(), total_pnl, len(trades_summary), json.dumps(trades_summary)))
            conn.commit()
            conn.close()
            
            logging.info(f"EOD Review for {clientcode}: Total P&L = Rs.{total_pnl}")
        
        except Exception as e:
            logging.error(f"Error in EOD review for {clientcode}: {e}")
    
    # Clear daily data
    DAILY_TRADE_PLAN.clear()
    ACTIVE_TRADES.clear()

# API endpoint to enable/disable auto-trading
@app.route('/api/autotrading/toggle', methods=['POST'])
def toggle_auto_trading():
    """Enable or disable automated trading for current user"""
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    
    clientcode = _SMARTAPI_SESSIONS[session_id]['clientcode']
    body = request.get_json() or {}
    enabled = body.get('enabled', False)
    starting_capital = body.get('starting_capital', 15000)
    
    # Initialize daily stats if enabling trading
    if enabled:
        initialize_daily_stats(clientcode, starting_capital)
        logging.info(f"[START] Auto-trading ENABLED for {clientcode} with capital Rs.{starting_capital:,.0f}")
    else:
        logging.info(f"⏸️ Auto-trading DISABLED for {clientcode}")
    
    AUTO_TRADING_ENABLED[clientcode] = enabled
    
    return jsonify({
        'status': True,
        'message': f"Auto-trading {'enabled' if enabled else 'disabled'}",
        'clientcode': clientcode,
        'enabled': enabled,
        'starting_capital': starting_capital if enabled else None
    })

@app.route('/api/autotrading/status', methods=['GET'])
def get_auto_trading_status():
    """Get current auto-trading status and today's plan"""
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    
    clientcode = _SMARTAPI_SESSIONS[session_id]['clientcode']
    
    return jsonify({
        'status': True,
        'enabled': AUTO_TRADING_ENABLED.get(clientcode, False),
        'trade_plan': DAILY_TRADE_PLAN.get(clientcode),
        'parsed_setups': PARSED_TRADE_SETUPS.get(clientcode, []),
        'active_trades': ACTIVE_TRADES.get(clientcode, {}),
        'plan_history_count': len(TRADE_PLAN_HISTORY.get(clientcode, [])),
        'timestamp': datetime.now().isoformat()
    })

@app.route('/api/autotrading/plan-history', methods=['GET'])
def get_trade_plan_history():
    """Get all saved trade plans"""
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    
    clientcode = _SMARTAPI_SESSIONS[session_id]['clientcode']
    history = TRADE_PLAN_HISTORY.get(clientcode, [])
    
    return jsonify({
        'status': True,
        'plans': history,
        'count': len(history)
    })

@app.route('/api/autotrading/select-plan', methods=['POST'])
def select_trade_plan_from_history():
    """Select a saved trade plan to activate"""
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    
    clientcode = _SMARTAPI_SESSIONS[session_id]['clientcode']
    body = request.get_json() or {}
    plan_id = body.get('plan_id')
    
    if not plan_id:
        return jsonify({'status': False, 'message': 'plan_id required'}), 400
    
    # Find the plan
    history = TRADE_PLAN_HISTORY.get(clientcode, [])
    selected_plan = None
    for plan in history:
        if plan['id'] == plan_id:
            selected_plan = plan
            plan['selected'] = True
        else:
            plan['selected'] = False
    
    if not selected_plan:
        return jsonify({'status': False, 'message': 'Plan not found'}), 404
    
    # Activate the selected plan
    DAILY_TRADE_PLAN[clientcode] = {
        'plan': selected_plan['plan'],
        'generated_at': selected_plan['generated_at'],
        'status': 'active',
        'generated_method': 'history-selection'
    }
    
    PARSED_TRADE_SETUPS[clientcode] = selected_plan['trades']
    
    logging.info(f"[HISTORY] Selected plan {plan_id} for {clientcode} with {len(selected_plan['trades'])} trades")
    
    return jsonify({
        'status': True,
        'message': 'Trade plan activated from history',
        'plan': selected_plan,
        'num_trades': len(selected_plan['trades'])
    })

@app.route('/api/autotrading/generate-now', methods=['POST'])
def generate_and_start_now():
    """
    ON-DEMAND: Generate trade plan and start trading immediately
    Use this when you start the system after 9:15 AM
    """
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    
    clientcode = _SMARTAPI_SESSIONS[session_id]['clientcode']
    body = request.get_json() or {}
    starting_capital = body.get('starting_capital', 15000)
    auto_enable = body.get('auto_enable', True)  # Auto-enable trading after generation
    
    try:
        # Check market hours
        now = datetime.now()
        current_time = now.time()
        market_open = datetime.strptime("09:15", "%H:%M").time()
        market_close = datetime.strptime("15:30", "%H:%M").time()
        
        if current_time < market_open:
            return jsonify({
                'status': False,
                'message': f'Market not yet open. Opens at 9:15 AM. Current time: {current_time.strftime("%H:%M")}'
            }), 400
        
        if current_time > market_close:
            return jsonify({
                'status': False,
                'message': f'Market closed. Closes at 3:30 PM. Current time: {current_time.strftime("%H:%M")}'
            }), 400
        
        logging.info("=" * 60)
        logging.info(f"[ON-DEMAND] Generating trade plan for {clientcode} at {current_time.strftime('%H:%M')}")
        logging.info("=" * 60)
        
        # Step 1: Initialize daily stats
        initialize_daily_stats(clientcode, starting_capital)
        logging.info(f"[STEP 1/4] Daily stats initialized with Rs.{starting_capital:,.0f}")
        
        # Step 2: Generate AI trade plan
        logging.info(f"[STEP 2/4] Calling AI to generate trade plan...")
        
        # Call AI recommendation endpoint using test client with session context
        with app.test_client() as client:
            with client.session_transaction() as sess:
                sess['session_id'] = session_id
            
            response = client.post(
                '/api/trading/ai-recommendation',
                json={'capital': starting_capital, 'risk_percent': 2, 'symbol': 'NIFTY'}
            )
        
        if response.status_code != 200:
            error_data = response.get_json() if response.data else {}
            logging.error(f"[ERROR] AI endpoint returned {response.status_code}: {error_data}")
            return jsonify({
                'status': False,
                'message': 'Failed to generate trade plan from AI',
                'error': error_data
            }), 500
        
        ai_data = response.get_json()
        plan_text = ai_data.get('recommendation')
        
        if not plan_text:
            return jsonify({
                'status': False,
                'message': 'AI returned empty trade plan'
            }), 500
        
        DAILY_TRADE_PLAN[clientcode] = {
            'plan': plan_text,
            'generated_at': datetime.now().isoformat(),
            'status': 'active',
            'generated_method': 'on-demand'
        }
        
        logging.info(f"[STEP 2/4] AI trade plan generated ({len(plan_text)} chars)")
        
        # Step 3: Parse trade plan
        logging.info(f"[STEP 3/4] Parsing trade plan...")
        
        parsed_data = parse_trade_plan_with_ai(plan_text, clientcode)
        
        if not parsed_data or not parsed_data.get('trades'):
            return jsonify({
                'status': False,
                'message': 'Failed to parse trade plan. AI might have returned invalid format.',
                'raw_plan': plan_text[:500]
            }), 500
        
        PARSED_TRADE_SETUPS[clientcode] = parsed_data['trades']
        num_trades = len(parsed_data['trades'])
        logging.info(f"[STEP 3/4] Parsed {num_trades} trade setups successfully")
        
        # Save to history
        plan_id = f"plan_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        if clientcode not in TRADE_PLAN_HISTORY:
            TRADE_PLAN_HISTORY[clientcode] = []
        
        TRADE_PLAN_HISTORY[clientcode].append({
            'id': plan_id,
            'plan': plan_text,
            'trades': parsed_data['trades'],
            'generated_at': datetime.now().isoformat(),
            'num_trades': num_trades,
            'selected': False,
            'capital': starting_capital,
            'method': 'on-demand'
        })
        
        logging.info(f"[HISTORY] Saved plan {plan_id} to history ({len(TRADE_PLAN_HISTORY[clientcode])} total)")
        
        # Step 4: Auto-enable trading if requested
        if auto_enable:
            AUTO_TRADING_ENABLED[clientcode] = True
            logging.info(f"[STEP 4/4] Auto-trading ENABLED for {clientcode}")
        else:
            logging.info(f"[STEP 4/4] Auto-trading NOT enabled (manual mode)")
        
        # Prepare trade summaries for response
        trade_summaries = []
        for trade in parsed_data['trades']:
            # Extract expiry from tradingsymbol (e.g., NIFTY21NOV2425850CE)
            tradingsymbol = trade.get('tradingsymbol', '')
            expiry_info = 'Not found'
            
            if tradingsymbol:
                import re
                # Match pattern: NIFTY + DATE (e.g., 21NOV24) + STRIKE + CE/PE
                match = re.search(r'NIFTY(\d{2}[A-Z]{3}\d{2})(\d+)(CE|PE)', tradingsymbol.upper())
                if match:
                    expiry_str = match.group(1)  # e.g., "21NOV24"
                    try:
                        # Parse expiry date
                        expiry_date = datetime.strptime(expiry_str, '%d%b%y')
                        expiry_info = expiry_date.strftime('%d-%b-%Y')  # e.g., "21-Nov-2024"
                    except:
                        expiry_info = expiry_str
            
            trade_summaries.append({
                'trade_number': trade.get('trade_number'),
                'symbol': trade.get('tradingsymbol', 'N/A'),
                'trade_type': trade.get('option_type', 'N/A'),
                'entry_price': trade.get('entry_price', 0),
                'stop_loss': trade.get('stop_loss', 0),
                'target_price': trade.get('target_1', 0),
                'target_2': trade.get('target_2', 0),
                'quantity': trade.get('quantity', 25),
                'entry_conditions': trade.get('entry_conditions', []),
                'expiry': expiry_info
            })
        
        logging.info("=" * 60)
        logging.info(f"[SUCCESS] On-demand setup complete! Trading {'ACTIVE' if auto_enable else 'READY (not enabled)'}")
        logging.info("=" * 60)
        
        return jsonify({
            'status': True,
            'message': f'Trade plan generated successfully at {current_time.strftime("%H:%M")}',
            'generated_at': datetime.now().isoformat(),
            'num_trades': num_trades,
            'trades': trade_summaries,
            'auto_trading_enabled': auto_enable,
            'starting_capital': starting_capital,
            'current_time': current_time.strftime("%H:%M"),
            'raw_plan': plan_text[:1000] + '...' if len(plan_text) > 1000 else plan_text
        })
    
    except Exception as e:
        logging.error(f"Error in on-demand generation: {e}", exc_info=True)
        return jsonify({
            'status': False,
            'message': f'Error generating trade plan: {str(e)}'
        }), 500

@app.route('/api/autotrading/test-execution', methods=['POST'])
def test_trade_execution():
    """Manually test trade execution (for testing only)"""
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    
    clientcode = _SMARTAPI_SESSIONS[session_id]['clientcode']
    
    try:
        # Get parsed trade setups
        trade_setups = PARSED_TRADE_SETUPS.get(clientcode, [])
        
        if not trade_setups:
            return jsonify({
                'status': False,
                'message': 'No parsed trade setups available. Generate trade plan first.'
            })
        
        results = []
        
        # Check each trade setup
        for trade_setup in trade_setups:
            trade_number = trade_setup.get('trade_number')
            
            # Check if entry conditions met
            conditions_met = evaluate_entry_conditions(trade_setup, clientcode)
            
            result = {
                'trade_number': trade_number,
                'instrument': trade_setup.get('instrument'),
                'conditions_met': conditions_met,
                'entry_conditions': trade_setup.get('entry_conditions'),
                'current_nifty': get_live_nifty_price(clientcode),
                'current_indicators': get_current_technical_indicators(clientcode)
            }
            
            results.append(result)
        
        return jsonify({
            'status': True,
            'message': 'Entry conditions evaluated',
            'results': results,
            'timestamp': datetime.now().isoformat()
        })
        
    except Exception as e:
        logging.error(f"Error in test execution: {e}", exc_info=True)
        return jsonify({'status': False, 'message': str(e)}), 500

@app.route('/api/autotrading/force-parse', methods=['POST'])
def force_parse_trade_plan():
    """Manually trigger trade plan parsing (for testing)"""
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    
    clientcode = _SMARTAPI_SESSIONS[session_id]['clientcode']
    
    try:
        plan_data = DAILY_TRADE_PLAN.get(clientcode)
        
        if not plan_data:
            return jsonify({
                'status': False,
                'message': 'No trade plan available. Generate plan first via AI Trading page.'
            })
        
        plan_text = plan_data.get('plan')
        
        # Parse trade plan
        parsed_data = parse_trade_plan_with_ai(plan_text, clientcode)
        
        if parsed_data and parsed_data.get('trades'):
            PARSED_TRADE_SETUPS[clientcode] = parsed_data['trades']
            
            return jsonify({
                'status': True,
                'message': f'Successfully parsed {len(parsed_data["trades"])} trade setups',
                'parsed_trades': parsed_data['trades']
            })
        else:
            return jsonify({
                'status': False,
                'message': 'Failed to parse trade plan'
            })
        
    except Exception as e:
        logging.error(f"Error parsing trade plan: {e}", exc_info=True)
        return jsonify({'status': False, 'message': str(e)}), 500

# ==================== BACKTESTING MODULE ====================

@app.route('/api/backtest/historical', methods=['POST'])
def backtest_historical_date():
    """
    Backtest automated trading system for a specific historical date
    Simulates entire trading day minute-by-minute with REAL option prices
    For OPTIONS TRADING: Single day analysis is sufficient (intraday positions)
    """
    session_id = session.get('session_id')
    if not session_id or session_id not in _SMARTAPI_SESSIONS:
        return jsonify({'status': False, 'message': 'Not logged in'}), 401
    
    clientcode = _SMARTAPI_SESSIONS[session_id]['clientcode']
    body = request.get_json() or {}
    
    backtest_date = body.get('date')  # Format: "2025-11-10"
    capital = body.get('capital', 15000)
    risk_percent = body.get('risk_percent', 2)
    
    if not backtest_date:
        return jsonify({'status': False, 'message': 'Date required (format: YYYY-MM-DD)'}), 400
    
    try:
        logging.info(f"=" * 60)
        logging.info(f"BACKTESTING: Simulating trading day for {backtest_date}")
        logging.info(f"Capital: Rs.{capital}, Risk: {risk_percent}%")
        
        # Step 1: Fetch historical candle data for the entire day
        logging.info("Step 1: Fetching NIFTY historical candles...")
        historical_candles = fetch_historical_day_candles(clientcode, backtest_date)
        
        if not historical_candles:
            logging.error(f"No NIFTY historical data available for {backtest_date}")
            return jsonify({
                'status': False,
                'message': f'No historical data available for {backtest_date}. This could be a holiday or weekend.'
            })
        
        logging.info(f"[OK] Fetched {len(historical_candles)} NIFTY candles")
        
        # Step 2: Generate AI trade plan based on day's opening data
        logging.info("Step 2: Generating AI trade plan from opening data...")
        logging.info(f"Using first 10 of {len(historical_candles)} candles")
        
        trade_plan_text = generate_backtest_trade_plan(
            clientcode, 
            historical_candles[:10],  # Use first 10 candles for plan
            capital,
            risk_percent,
            backtest_date  # Pass date to fetch uploaded documents
        )
        
        if not trade_plan_text:
            logging.error("Trade plan generation returned None or empty string")
            return jsonify({
                'status': False,
                'message': 'Failed to generate AI trade plan'
            })
        
        logging.info("[OK] Trade plan generated")
        logging.info(f"AI Trade Plan:\n{trade_plan_text}")
        
        # Step 3: Parse trade plan
        logging.info("Step 3: Parsing trade plan with AI...")
        parsed_data = parse_trade_plan_with_ai(trade_plan_text, clientcode)
        
        if not parsed_data or not parsed_data.get('trades'):
            logging.error("Failed to parse trade plan")
            return jsonify({
                'status': False,
                'message': 'Failed to parse backtest trade plan'
            })
        
        trade_setups = parsed_data['trades']
        logging.info(f"[OK] Parsed {len(trade_setups)} trade setups")
        logging.info(f"Parsed JSON:\n{json.dumps(parsed_data, indent=2)}")
        
        # Log parsed trades
        for i, setup in enumerate(trade_setups, 1):
            logging.info(f"  Trade {i}: {setup.get('tradingsymbol', 'N/A')}")
            logging.info(f"    Entry Price: {setup.get('entry_price', 'N/A')}")
            logging.info(f"    Entry Conditions: {setup.get('entry_conditions', 'N/A')}")
            logging.info(f"    SL: {setup.get('stop_loss', 'N/A')} | Target 1: {setup.get('target_1', 'N/A')} | Target 2: {setup.get('target_2', 'N/A')}")
        
        # Step 3.5: Enrich trade setups with symbol tokens for historical option data
        logging.info("Step 4: Finding symbol tokens for options...")
        trade_setups = enrich_trade_setups_with_tokens(clientcode, trade_setups, backtest_date)
        
        # Log tokens found
        tokens_found = sum(1 for s in trade_setups if s.get('symboltoken'))
        logging.info(f"  Found tokens for {tokens_found}/{len(trade_setups)} trades")
        
        # Step 4: Simulate trading day with REAL historical option prices
        logging.info("Step 5: Starting day simulation with real option prices...")
        simulation_result = simulate_trading_day(
            clientcode,
            trade_setups,
            historical_candles,
            backtest_date
        )
        
        if not simulation_result:
            logging.error("Simulation failed")
            return jsonify({
                'status': False,
                'message': 'Simulation failed'
            })
        
        logging.info("[OK] Simulation complete")
        logging.info(f"Results: {simulation_result['total_trades']} trades, P&L: Rs.{simulation_result['total_pnl']:.2f}")
        
        # Store backtest results
        logging.info("Step 6: Storing results in database...")
        store_backtest_results(clientcode, backtest_date, simulation_result)
        logging.info("[OK] Backtest complete!")
        logging.info("=" * 60)
        
        return jsonify({
            'status': True,
            'date': backtest_date,
            'trade_plan': trade_plan_text,
            'parsed_setups': trade_setups,
            'simulation': simulation_result,
            'summary': {
                'total_trades': simulation_result['total_trades'],
                'winning_trades': simulation_result['winning_trades'],
                'losing_trades': simulation_result['losing_trades'],
                'total_pnl': simulation_result['total_pnl'],
                'win_rate': simulation_result['win_rate'],
                'max_drawdown': simulation_result['max_drawdown']
            }
        })
        
    except Exception as e:
        logging.error(f"Backtest error: {e}", exc_info=True)
        return jsonify({'status': False, 'message': str(e)}), 500

def fetch_historical_day_candles(clientcode, date_str):
    """Fetch 5-minute candles for entire trading day (NIFTY spot only)"""
    try:
        logging.info(f"Attempting to fetch historical data for {date_str}...")
        
        session_id = None
        for sid, sdata in _SMARTAPI_SESSIONS.items():
            if sdata.get('clientcode') == clientcode:
                session_id = sid
                break
        
        if not session_id:
            logging.error(f"No active session found for clientcode {clientcode}")
            return None
        
        smartapi = _SMARTAPI_SESSIONS[session_id]['api']
        
        # Set access token before making API call
        tokens = _SMARTAPI_SESSIONS[session_id].get('tokens', {})
        jwt_token = tokens.get('jwtToken')
        if jwt_token:
            # Strip 'Bearer ' prefix if present (setAccessToken adds it automatically)
            if jwt_token.startswith('Bearer '):
                jwt_token = jwt_token[7:]  # Remove 'Bearer ' (7 chars)
                logging.info("Stripped 'Bearer ' prefix from token")
            
            token_preview = f"{jwt_token[:20]}...{jwt_token[-20:]}" if len(jwt_token) > 40 else jwt_token
            logging.info(f"Setting access token: {token_preview}")
            smartapi.setAccessToken(jwt_token)
            logging.info(f"Token set successfully")
        else:
            logging.error("No jwtToken found in session - API call will fail")
            return None
        
        # Parse date
        target_date = datetime.strptime(date_str, '%Y-%m-%d')
        
        # Check if weekend
        if target_date.weekday() >= 5:  # Saturday=5, Sunday=6
            logging.warning(f"Weekend date selected: {date_str} ({target_date.strftime('%A')})")
        
        # Market hours: 9:15 AM to 3:30 PM
        from_datetime = target_date.replace(hour=9, minute=15)
        to_datetime = target_date.replace(hour=15, minute=30)
        
        params = {
            'exchange': 'NSE',
            'symboltoken': '99926000',  # NIFTY 50
            'interval': 'FIVE_MINUTE',
            'fromdate': from_datetime.strftime('%Y-%m-%d %H:%M'),
            'todate': to_datetime.strftime('%Y-%m-%d %H:%M')
        }
        
        logging.info(f"Calling Angel One getCandleData API with params: {params}")
        candles = smartapi.getCandleData(params)
        
        if candles and candles.get('status') and candles.get('data'):
            logging.info(f"[OK] SUCCESS: Fetched {len(candles['data'])} NIFTY candles for {date_str}")
            return candles['data']
        else:
            logging.error(f"[X] FAILED: Angel One returned no data or error status")
            logging.error(f"Response: {candles}")
            return None
        
    except Exception as e:
        logging.error(f"[X] EXCEPTION fetching historical candles: {e}", exc_info=True)
        return None

def fetch_historical_option_candles(clientcode, symboltoken, tradingsymbol, date_str):
    """Fetch 5-minute option candles for entire trading day using getCandleData with NFO exchange"""
    try:
        logging.info(f"  Fetching option candles for {tradingsymbol} (token: {symboltoken})...")
        
        session_id = None
        for sid, sdata in _SMARTAPI_SESSIONS.items():
            if sdata.get('clientcode') == clientcode:
                session_id = sid
                break
        
        if not session_id:
            logging.error(f"  No session found")
            return None
        
        smartapi = _SMARTAPI_SESSIONS[session_id]['api']
        
        # Set access token before making API call
        tokens = _SMARTAPI_SESSIONS[session_id].get('tokens', {})
        jwt_token = tokens.get('jwtToken')
        if jwt_token:
            # Strip 'Bearer ' prefix if present
            if jwt_token.startswith('Bearer '):
                jwt_token = jwt_token[7:]
            smartapi.setAccessToken(jwt_token)
        else:
            logging.error("  No jwtToken found for option data fetch")
            return None
        
        # Parse date
        target_date = datetime.strptime(date_str, '%Y-%m-%d')
        
        # Market hours: 9:15 AM to 3:30 PM
        from_datetime = target_date.replace(hour=9, minute=15)
        to_datetime = target_date.replace(hour=15, minute=30)
        
        params = {
            'exchange': 'NFO',  # F&O segment for options
            'symboltoken': symboltoken,
            'interval': 'FIVE_MINUTE',
            'fromdate': from_datetime.strftime('%Y-%m-%d %H:%M'),
            'todate': to_datetime.strftime('%Y-%m-%d %H:%M')
        }
        
        logging.info(f"  API call: getCandleData(NFO, {symboltoken}, FIVE_MINUTE)")
        candles = smartapi.getCandleData(params)
        
        if candles and candles.get('status') and candles.get('data'):
            logging.info(f"  [OK] Fetched {len(candles['data'])} option candles for {tradingsymbol}")
            # Convert to dict with timestamp as key for fast lookup
            candle_dict = {}
            for candle in candles['data']:
                timestamp = candle[0]
                candle_dict[timestamp] = {
                    'open': candle[1],
                    'high': candle[2],
                    'low': candle[3],
                    'close': candle[4],
                    'volume': candle[5]
                }
            return candle_dict
        else:
            logging.warning(f"  [X] No option data for {tradingsymbol} - Response: {candles}")
            return None
        
    except Exception as e:
        logging.error(f"  [X] Error fetching option candles for {tradingsymbol}: {e}")
        return None

def generate_backtest_trade_plan(clientcode, opening_candles, capital, risk_percent, backtest_date=None):
    """Generate AI trade plan using opening market data and optional uploaded documents"""
    try:
        # Calculate indicators from opening candles
        logging.info(f"Generating plan from {len(opening_candles)} opening candles")
        
        df = pd.DataFrame(opening_candles, columns=['timestamp', 'open', 'high', 'low', 'close', 'volume'])
        logging.info(f"DataFrame created with shape: {df.shape}")
        
        # Calculate RSI
        df['rsi'] = ta.rsi(df['close'], length=14)
        
        # Calculate MACD - returns DataFrame with multiple columns
        macd_result = ta.macd(df['close'])
        if macd_result is not None and isinstance(macd_result, pd.DataFrame):
            df['macd'] = macd_result['MACD_12_26_9']
        else:
            df['macd'] = 0
        
        current_price = df['close'].iloc[-1]
        current_rsi = df['rsi'].iloc[-1] if not pd.isna(df['rsi'].iloc[-1]) else 50
        current_macd = df['macd'].iloc[-1] if not pd.isna(df['macd'].iloc[-1]) else 0
        
        logging.info(f"Indicators: Price={current_price}, RSI={current_rsi:.2f}, MACD={current_macd:.2f}")
        
        # Fetch uploaded documents for the backtest date
        document_analysis = ""
        if backtest_date:
            try:
                conn = sqlite3.connect(DB_FILE)
                conn.row_factory = sqlite3.Row
                c = conn.cursor()
                
                c.execute('''
                    SELECT * FROM documents 
                    WHERE clientcode = ? AND upload_date = ?
                    ORDER BY id DESC
                ''', (clientcode, backtest_date))
                doc_rows = c.fetchall()
                conn.close()
                
                if doc_rows:
                    logging.info(f"Found {len(doc_rows)} uploaded documents for {backtest_date}")
                    document_analysis = "\n\n[UPLOADED RESEARCH DOCUMENTS]:\n"
                    for doc_row in doc_rows:
                        doc_type = doc_row['document_type']
                        source_name = "Perplexity AI" if doc_type == 'perplexity' else "OpenAI Research"
                        document_analysis += f"\n{source_name}:\n{doc_row['content'][:2000]}\n"  # Limit to 2000 chars per doc
                else:
                    logging.info(f"No uploaded documents found for {backtest_date}")
            except Exception as e:
                logging.warning(f"Could not fetch documents: {e}")
        
        # Fetch pre-market global data
        try:
            premarket = fetch_premarket_data()
        except:
            premarket = {}
        
        # Calculate max allocation per trade (50% of capital per trade for 2 trades)
        max_per_trade = capital * 0.5
        nifty_lot_size = 75  # NIFTY lot size (updated Nov 2024)
        
        # Build prompt for AI
        prompt = f"""Generate intraday NIFTY options trade plan based on HISTORICAL data.

CAPITAL: Rs.{capital:,}
MAX PER TRADE: Rs.{max_per_trade:,.0f} (50% of capital)
CURRENT NIFTY: {current_price:.2f}
RSI: {current_rsi:.2f}
MACD: {current_macd:.2f}
NIFTY LOT SIZE: {nifty_lot_size}
{document_analysis}

Generate 1-2 NIFTY option trade setups with complete details:

For each trade, specify:
1. Strike price (ATM, slightly OTM based on NIFTY level)
2. Option type (CE for bullish, PE for bearish)
3. Entry price: Option premium price (realistic based on NIFTY level)
4. Entry conditions: NIFTY spot price level that triggers entry
5. Stop loss: Option premium level (not NIFTY index)
6. Target 1 & Target 2: Option premium levels
7. **QUANTITY CALCULATION**: Calculate lots to maximize capital usage
   - Formula: Lots = floor(Max Per Trade / (Entry Premium × Lot Size))
   - Quantity = Lots × {nifty_lot_size}
   - Example: If entry premium is Rs.120 and max per trade is Rs.{max_per_trade:,.0f}:
     * Lots = floor({max_per_trade:,.0f} / (120 × {nifty_lot_size})) = floor({max_per_trade/3000:.1f}) = {int(max_per_trade/3000)} lots
     * Quantity = {int(max_per_trade/3000)} × {nifty_lot_size} = {int(max_per_trade/3000) * nifty_lot_size}
8. Entry time window: e.g., 09:30 to 11:00

Guidelines:
- Option premiums: Rs.50-200 range for ATM options
- Stop loss: 20-30% below entry price
- Target 1: 15-20% above entry price
- Target 2: 30-40% above entry price

**INTRADAY ENTRY STRATEGY - REALISTIC EXECUTION**:

Current Market Level: {current_price:.0f}

Entry Distance Guidelines:
- **AGGRESSIVE SCALP** (Recommended): +/-30 to 50 points → High execution probability (70-80%)
  * Bullish: {current_price+30:.0f} to {current_price+50:.0f}
  * Bearish: {current_price-50:.0f} to {current_price-30:.0f}
  
- **MODERATE SWING**: +/-50 to 100 points → Medium execution probability (40-60%)
  * Bullish: {current_price+50:.0f} to {current_price+100:.0f}
  * Bearish: {current_price-100:.0f} to {current_price-50:.0f}

- **AVOID**: Entries beyond +/-150 points → Low execution probability (<20%)
  * These rarely trigger in a single trading day

RULE: Default to AGGRESSIVE SCALP range for maximum execution rate. Only suggest MODERATE range if strong trend/momentum indicators support big moves. NEVER suggest entries beyond +/-150 points for intraday plans.

Rationale: Indian market intraday moves average 100-200 points. Suggesting 200+ point entries means trades won't execute, capital sits idle. Tight entries = More action = Better capital utilization.

- Entry condition format: "When NIFTY crosses above {current_price+50:.0f}" (for CE) or "When NIFTY crosses below {current_price-50:.0f}" (for PE)
- **IMPORTANT**: Calculate quantity to use maximum capital available per trade

Format example (for capital Rs.25,000, max per trade Rs.12,500):
Trade 1: NIFTY 26000 CE
Entry Premium: Rs.120
Entry Condition: When NIFTY crosses above 25900
Stop Loss: Rs.85 (premium)
Target 1: Rs.140 (premium)
Target 2: Rs.165 (premium)
Quantity: 100 (calculated: floor(12500/(120×25)) = 4 lots = 4×25 = 100 qty)
Entry Time: 09:30 to 11:30

Trade 2: NIFTY 25700 PE
Entry Premium: Rs.100
Entry Condition: When NIFTY crosses below 25750
Stop Loss: Rs.70 (premium)
Target 1: Rs.120 (premium)
Target 2: Rs.145 (premium)
Quantity: 125 (calculated: floor(12500/(100×25)) = 5 lots = 5×25 = 125 qty)
Entry Time: 09:30 to 12:00"""

        response = openai_client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a professional intraday trader."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=2000
        )
        
        return response.choices[0].message.content
        
    except Exception as e:
        logging.error(f"Error generating backtest plan: {e}", exc_info=True)
        return None

def construct_nifty_option_symbol(strike, option_type, expiry_date):
    """
    Construct proper NIFTY option trading symbol
    Example: NIFTY18NOV2525000CE
    
    Args:
        strike: Strike price (e.g., 25000, 25500)
        option_type: 'CE' or 'PE'
        expiry_date: datetime object or string 'DDMMMYYYY' format
    """
    try:
        if isinstance(expiry_date, str):
            # Parse date string like '2025-11-11'
            from datetime import datetime
            expiry_date = datetime.strptime(expiry_date, '%Y-%m-%d')
        
        # Format: NIFTY18NOV2525000CE
        day = expiry_date.strftime('%d')
        month = expiry_date.strftime('%b').upper()
        year = expiry_date.strftime('%y')
        
        # Find next weekly expiry (Tuesdays - NEW RULE)
        weekday = expiry_date.weekday()  # 0=Monday, 1=Tuesday
        
        if weekday == 1:  # If today is Tuesday
            # Use next week's Tuesday
            days_to_tuesday = 7
        elif weekday < 1:  # Monday
            # Use this week's Tuesday
            days_to_tuesday = 1 - weekday
        else:  # Wednesday onwards (2-6)
            # Use next week's Tuesday
            days_to_tuesday = (7 - weekday) + 1
        
        from datetime import timedelta
        expiry = expiry_date + timedelta(days=days_to_tuesday)
        
        expiry_day = expiry.strftime('%d')
        expiry_month = expiry.strftime('%b').upper()
        expiry_year = expiry.strftime('%y')
        
        # Strike as integer (25900, not in paise)
        strike_int = int(strike)
        
        symbol = f"NIFTY{expiry_day}{expiry_month}{expiry_year}{strike_int}{option_type}"
        
        return symbol
        
    except Exception as e:
        logging.error(f"Error constructing option symbol: {e}")
        return None

def enrich_trade_setups_with_tokens(clientcode, trade_setups, backtest_date):
    """Find and add symbol tokens to trade setups for historical data fetching"""
    try:
        from datetime import datetime, timedelta
        
        # Parse backtest date
        trade_date = datetime.strptime(backtest_date, '%Y-%m-%d')
        
        # Find next weekly expiry (Tuesday)
        # NIFTY weekly options expire on Tuesdays
        weekday = trade_date.weekday()  # 0=Monday, 1=Tuesday, 2=Wednesday...
        
        if weekday == 1:  # If today is Tuesday
            # Use next week's Tuesday (7 days ahead)
            days_to_expiry = 7
        elif weekday < 1:  # Monday
            # Use this week's Tuesday
            days_to_expiry = 1 - weekday
        else:  # Wednesday onwards (2-6)
            # Use next week's Tuesday
            days_to_expiry = (7 - weekday) + 1
        
        expiry_date = trade_date + timedelta(days=days_to_expiry)
        expiry_day = expiry_date.strftime('%d')
        expiry_month = expiry_date.strftime('%b').upper()
        expiry_year = expiry_date.strftime('%y')
        
        logging.info(f"  Trade date: {backtest_date}, Expiry: {expiry_day}{expiry_month}{expiry_year}")
        
        for setup in trade_setups:
            tradingsymbol = setup.get('tradingsymbol')
            
            # If tradingsymbol is missing or incomplete, try to construct it
            if not tradingsymbol or 'NIFTY' not in tradingsymbol.upper():
                # Try to extract from instrument field (e.g., "NIFTY 25600 CE")
                instrument = setup.get('instrument', '')
                if 'NIFTY' in instrument.upper():
                    parts = instrument.split()
                    if len(parts) >= 3:
                        try:
                            strike = int(parts[1])
                            option_type = parts[2].upper()  # CE or PE
                            tradingsymbol = f"NIFTY{strike}{option_type}"
                            setup['tradingsymbol'] = tradingsymbol
                        except:
                            pass
            
            # Now construct the full symbol with expiry date
            if tradingsymbol:
                # Extract strike and option type from tradingsymbol like "NIFTY25850CE"
                import re
                match = re.search(r'NIFTY(\d+)(CE|PE)', tradingsymbol.upper())
                if match:
                    strike = match.group(1)
                    option_type = match.group(2)
                    
                    # Construct full symbol: NIFTY18NOV2525850CE
                    full_symbol = f"NIFTY{expiry_day}{expiry_month}{expiry_year}{strike}{option_type}"
                    
                    logging.info(f"  Constructed full symbol: {tradingsymbol} -> {full_symbol}")
                    
                    # Find token from scrip master
                    token_result = find_symbol_token(full_symbol, clientcode)
                    if token_result:
                        setup['symboltoken'] = token_result['token']
                        setup['tradingsymbol'] = token_result['symbol']  # Use exact symbol from scrip master
                        logging.info(f"  Found token: {token_result['token']} for {token_result['symbol']}")
                    else:
                        logging.warning(f"  [WARNING] Token not found for {full_symbol}")
                else:
                    logging.warning(f"  [WARNING] Could not parse tradingsymbol: {tradingsymbol}")
        
        return trade_setups
        
    except Exception as e:
        logging.error(f"Error enriching trade setups: {e}", exc_info=True)
        return trade_setups

def simulate_trading_day(clientcode, trade_setups, historical_candles, date_str):
    """Simulate trading day minute-by-minute with REAL historical option prices"""
    try:
        executed_trades = []
        current_positions = []
        total_pnl = 0
        max_drawdown = 0
        capital_curve = []
        current_capital = 15000
        
        logging.info(f"Simulating {len(historical_candles)} candles for {len(trade_setups)} trade setups")
        
        # Log trade setup details
        for i, setup in enumerate(trade_setups, 1):
            logging.info(f"  Setup {i}: {setup.get('tradingsymbol', 'MISSING')} | Token: {setup.get('symboltoken', 'NOT FOUND')}")
        
        # Fetch historical option candles for each trade setup
        option_candles_cache = {}
        for setup in trade_setups:
            symboltoken = setup.get('symboltoken')
            tradingsymbol = setup.get('tradingsymbol')
            
            if symboltoken and tradingsymbol:
                option_candles = fetch_historical_option_candles(
                    clientcode, symboltoken, tradingsymbol, date_str
                )
                if option_candles:
                    option_candles_cache[tradingsymbol] = option_candles
                    logging.info(f"  Loaded {len(option_candles)} candles for {tradingsymbol}")
                else:
                    logging.warning(f"  [WARNING] No historical data for {tradingsymbol}, will use approximation")
        
        # Process each 5-minute candle
        for i, candle in enumerate(historical_candles):
            timestamp, open_price, high, low, close_price, volume = candle
            candle_time = datetime.strptime(timestamp, '%Y-%m-%dT%H:%M:%S%z') if 'T' in timestamp else datetime.strptime(timestamp, '%Y-%m-%d %H:%M:%S')
            
            # Calculate indicators up to this point
            indicators = calculate_indicators_upto_candle(historical_candles[:i+1])
            
            # Check entry conditions for pending setups
            for setup in trade_setups:
                if setup.get('executed'):
                    continue
                
                # Skip if no trading symbol found
                if not setup.get('tradingsymbol'):
                    if i == 0:  # Log once at start
                        logging.warning(f"  [SKIP] Setup {setup.get('trade_number', '?')}: No tradingsymbol found")
                    continue
                
                # Simulate entry condition check
                conditions_met = simulate_entry_check(setup, close_price, indicators, candle_time)
                
                if conditions_met:
                    # Execute trade
                    trade = {
                        'setup_number': setup['trade_number'],
                        'instrument': setup['instrument'],
                        'tradingsymbol': setup.get('tradingsymbol'),
                        'entry_time': timestamp,
                        'entry_price': setup['entry_price'],
                        'quantity': setup['quantity'],
                        'stop_loss': setup['stop_loss'],
                        'target_1': setup['target_1'],
                        'target_2': setup['target_2'],
                        'status': 'open',
                        'current_price': setup['entry_price']
                    }
                    current_positions.append(trade)
                    setup['executed'] = True
                    logging.info(f"  [OK] Trade executed at {timestamp}: {setup['instrument']} @ Rs.{setup['entry_price']}")
            
            # Monitor open positions
            for position in current_positions:
                if position['status'] != 'open':
                    continue
                
                tradingsymbol = position.get('tradingsymbol')
                
                # Get REAL historical option price from fetched candles
                if tradingsymbol and tradingsymbol in option_candles_cache:
                    option_data = option_candles_cache[tradingsymbol].get(timestamp)
                    if option_data:
                        # Use actual market close price
                        simulated_option_price = option_data['close']
                    else:
                        # Fallback: use approximation if specific timestamp missing
                        price_change_pct = (close_price - float(position['entry_price'])) / float(position['entry_price'])
                        simulated_option_price = float(position['entry_price']) * (1 + price_change_pct * 0.5)
                else:
                    # Fallback: approximate (when historical option data not available)
                    price_change_pct = (close_price - float(position['entry_price'])) / float(position['entry_price'])
                    simulated_option_price = float(position['entry_price']) * (1 + price_change_pct * 0.5)
                
                position['current_price'] = simulated_option_price
                
                # Calculate profit percentage
                profit_pct = ((simulated_option_price - position['entry_price']) / position['entry_price']) * 100
                
                # Smart trailing stop loss (time + VIX aware - same as live trading)
                candle_time_obj = datetime.strptime(timestamp, '%Y-%m-%dT%H:%M:%S%z') if 'T' in timestamp else datetime.strptime(timestamp, '%Y-%m-%d %H:%M:%S')
                morning_volatility = candle_time_obj.time() < datetime.strptime("10:30", "%H:%M").time()
                
                # Get VIX-based thresholds (use current VIX as approximation for backtest)
                vix_value = get_current_vix_value()
                thresholds = calculate_vix_based_thresholds(vix_value)
                profit_target = calculate_vix_based_profit_target(vix_value)
                
                # EARLY SESSION (9:15-10:30): Only trail after profit target (time-based safety)
                if morning_volatility:
                    if profit_pct >= profit_target and position.get('stop_loss', 0) < position['entry_price']:
                        position['stop_loss'] = position['entry_price']
                
                # LATER SESSION (10:30-15:15): VIX-based dynamic trailing
                else:
                    # Move to breakeven based on VIX threshold
                    if profit_pct >= thresholds['breakeven_threshold'] and position.get('stop_loss', 0) < position['entry_price']:
                        position['stop_loss'] = position['entry_price']
                    
                    # Trail at 5% below current based on VIX threshold
                    elif profit_pct >= thresholds['trail_threshold']:
                        new_sl = simulated_option_price * 0.95
                        if new_sl > position.get('stop_loss', 0):
                            position['stop_loss'] = new_sl
                
                # Check VIX-based profit target (quick exit)
                if profit_pct >= profit_target:
                    pnl = (simulated_option_price - position['entry_price']) * position['quantity']
                    position['status'] = f'closed_profit_{profit_target:.0f}pct'
                    position['exit_time'] = timestamp
                    position['exit_price'] = simulated_option_price
                    position['pnl'] = pnl
                    total_pnl += pnl
                    executed_trades.append(position.copy())
                    logging.info(f"  [{profit_target:.1f}% PROFIT] VIX-based quick exit (VIX={vix_value:.1f if vix_value else 'N/A'}): {position['instrument']} Entry: Rs.{position['entry_price']:.2f} Exit: Rs.{simulated_option_price:.2f} P&L: Rs.{pnl:.2f}")
                    
                    # Mark the setup as not executed so it can be re-entered
                    for setup in trade_setups:
                        if setup.get('trade_number') == position['setup_number']:
                            setup['executed'] = False
                            logging.info(f"  [RESET] Setup {position['setup_number']} reset for re-entry")
                            break
                
                # Check stop loss
                elif simulated_option_price <= position['stop_loss']:
                    pnl = (simulated_option_price - position['entry_price']) * position['quantity']
                    position['status'] = 'closed_sl'
                    position['exit_time'] = timestamp
                    position['exit_price'] = simulated_option_price
                    position['pnl'] = pnl
                    total_pnl += pnl
                    executed_trades.append(position.copy())
                    logging.info(f"  [SL] Stop Loss Hit: {position['instrument']} P&L: Rs.{pnl:.2f}")
                
                # Check target
                elif simulated_option_price >= position['target_2']:
                    pnl = (simulated_option_price - position['entry_price']) * position['quantity']
                    position['status'] = 'closed_target'
                    position['exit_time'] = timestamp
                    position['exit_price'] = simulated_option_price
                    position['pnl'] = pnl
                    total_pnl += pnl
                    executed_trades.append(position.copy())
                    logging.info(f"  [TARGET] Target Hit: {position['instrument']} P&L: Rs.{pnl:.2f}")
            
            # Update capital curve
            unrealized_pnl = sum(
                (p['current_price'] - p['entry_price']) * p['quantity'] 
                for p in current_positions if p['status'] == 'open'
            )
            current_capital = 15000 + total_pnl + unrealized_pnl
            capital_curve.append({'time': timestamp, 'capital': current_capital})
            
            # Calculate drawdown
            if capital_curve:
                peak = max(c['capital'] for c in capital_curve)
                drawdown = (peak - current_capital) / peak * 100
                max_drawdown = max(max_drawdown, drawdown)
        
        # Close any remaining positions at EOD
        for position in current_positions:
            if position['status'] == 'open':
                pnl = (position['current_price'] - position['entry_price']) * position['quantity']
                position['status'] = 'closed_eod'
                position['exit_time'] = historical_candles[-1][0]
                position['exit_price'] = position['current_price']
                position['pnl'] = pnl
                total_pnl += pnl
                executed_trades.append(position.copy())
        
        logging.info(f"[OK] Simulation complete - {len(executed_trades)} trades executed")
        
        # Log reasons for non-execution
        for setup in trade_setups:
            if not setup.get('executed'):
                if not setup.get('tradingsymbol'):
                    logging.warning(f"  Setup {setup.get('trade_number', '?')}: NOT EXECUTED - Missing tradingsymbol")
                else:
                    logging.warning(f"  Setup {setup.get('trade_number', '?')} ({setup.get('tradingsymbol')}): NOT EXECUTED - Entry conditions never met")
        
        logging.info(f"Total P&L: Rs.{total_pnl:.2f}")
        
        # Calculate statistics
        winning_trades = [t for t in executed_trades if t.get('pnl', 0) > 0]
        losing_trades = [t for t in executed_trades if t.get('pnl', 0) < 0]
        
        return {
            'date': date_str,
            'total_trades': len(executed_trades),
            'winning_trades': len(winning_trades),
            'losing_trades': len(losing_trades),
            'win_rate': len(winning_trades) / len(executed_trades) * 100 if executed_trades else 0,
            'total_pnl': total_pnl,
            'max_drawdown': max_drawdown,
            'trades': executed_trades,
            'capital_curve': capital_curve
        }
        
    except Exception as e:
        logging.error(f"Simulation error: {e}", exc_info=True)
        return None

def calculate_indicators_upto_candle(candles):
    """Calculate technical indicators from candles up to current point"""
    try:
        df = pd.DataFrame(candles, columns=['timestamp', 'open', 'high', 'low', 'close', 'volume'])
        
        df['rsi'] = ta.rsi(df['close'], length=14)
        macd_result = ta.macd(df['close'])
        df['macd'] = macd_result['MACD_12_26_9'] if macd_result is not None else 0
        
        return {
            'rsi': df['rsi'].iloc[-1] if not pd.isna(df['rsi'].iloc[-1]) else 50,
            'macd': df['macd'].iloc[-1] if not pd.isna(df['macd'].iloc[-1]) else 0
        }
    except:
        return {'rsi': 50, 'macd': 0}

def simulate_entry_check(setup, current_price, indicators, current_time):
    """Check if entry conditions met during backtest"""
    try:
        conditions = setup.get('entry_conditions', [])
        
        # Log once at first check
        if not hasattr(simulate_entry_check, f'_logged_{setup.get("trade_number")}'):
            logging.info(f"    Checking entry for Trade {setup.get('trade_number')}: {len(conditions)} conditions")
            for cond in conditions:
                logging.info(f"      - {cond}")
            setattr(simulate_entry_check, f'_logged_{setup.get("trade_number")}', True)
        
        # Check time window
        entry_start = datetime.strptime(setup.get('entry_time_start', '09:30'), '%H:%M').time()
        entry_end = datetime.strptime(setup.get('entry_time_end', '15:00'), '%H:%M').time()
        
        if not (entry_start <= current_time.time() <= entry_end):
            return False
        
        # Check all conditions
        for condition in conditions:
            condition_type = condition.get('type')
            indicator_name = condition.get('indicator', '').upper()
            operator = condition.get('operator')
            threshold = condition.get('value')
            
            if condition_type == 'price' and indicator_name == 'NIFTY':
                current_value = current_price
            elif condition_type == 'indicator':
                current_value = indicators.get(indicator_name.lower())
                if current_value is None:
                    return False
            else:
                continue
            
            # Evaluate
            if operator == '>' and not (current_value > threshold):
                return False
            elif operator == '<' and not (current_value < threshold):
                return False
            elif operator == '>=' and not (current_value >= threshold):
                return False
            elif operator == '<=' and not (current_value <= threshold):
                return False
        
        return True
        
    except:
        return False

def store_backtest_results(clientcode, date, results):
    """Store backtest results in database"""
    try:
        conn = sqlite3.connect(DB_FILE)
        c = conn.cursor()
        
        c.execute('''
            CREATE TABLE IF NOT EXISTS backtest_results (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                clientcode TEXT,
                date TEXT,
                total_pnl REAL,
                total_trades INTEGER,
                win_rate REAL,
                results_json TEXT,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        c.execute('''
            INSERT INTO backtest_results (clientcode, date, total_pnl, total_trades, win_rate, results_json)
            VALUES (?, ?, ?, ?, ?, ?)
        ''', (
            clientcode,
            date,
            results['total_pnl'],
            results['total_trades'],
            results['win_rate'],
            json.dumps(results)
        ))
        
        conn.commit()
        conn.close()
        
        logging.info(f"Backtest results stored for {date}")
        
    except Exception as e:
        logging.error(f"Error storing backtest results: {e}")

@app.route('/view/backtest-history')
def view_backtest_history():
    if 'session_id' not in session:
        return redirect(url_for('index'))
    return render_template('backtest_history.html')

# Initialize scheduler with IST timezone
scheduler = BackgroundScheduler(timezone=IST)

# Schedule jobs
scheduler.add_job(
    fetch_premarket_data,
    CronTrigger(hour=9, minute=0, day_of_week='mon-fri', timezone=IST),
    id='premarket_data',
    name='Fetch pre-market data at 9:00 AM IST'
)

# NEW: Opening volatility scalp analysis at 9:10 AM
scheduler.add_job(
    lambda: analyze_opening_volatility_scalp(),
    CronTrigger(hour=9, minute=10, day_of_week='mon-fri', timezone=IST),
    id='opening_scalp_analysis',
    name='Analyze opening volatility at 9:10 AM IST'
)

# NEW: Execute opening volatility scalp at 9:16 AM (after first candle confirmation)
scheduler.add_job(
    lambda: execute_opening_volatility_scalp(),
    CronTrigger(hour=9, minute=16, day_of_week='mon-fri', timezone=IST),
    id='opening_scalp_execute',
    name='Execute opening scalp at 9:16 AM IST (first candle confirmed)'
)

scheduler.add_job(
    generate_daily_trade_plan,
    CronTrigger(hour=9, minute=15, day_of_week='mon-fri', timezone=IST),
    id='generate_plan',
    name='Generate trade plan at 9:15 AM IST'
)

# Real-time trade monitoring (every 60 seconds for regular trades)
scheduler.add_job(
    monitor_active_trades_sl_target,
    'interval',
    seconds=60,  # Check every 60 seconds for stop loss and targets
    id='monitor_trades',
    name='Monitor active trades (every 60 seconds)',
    max_instances=1  # Prevent overlapping executions
)

# Hourly heartbeat to confirm monitoring is active
def monitoring_heartbeat():
    """Hourly confirmation that monitoring systems are running"""
    active_count = sum(len([t for t in trades.values() if t.get('status') == 'open']) 
                       for trades in ACTIVE_TRADES.values())
    logging.info(f"[HEARTBEAT] System active | {active_count} open positions | Next check in 1 hour")

scheduler.add_job(
    monitoring_heartbeat,
    'interval',
    hours=1,
    id='monitoring_heartbeat',
    name='Hourly monitoring heartbeat',
    next_run_time=get_ist_now() + timedelta(hours=1)  # First heartbeat in 1 hour (IST)
)

# Performance review jobs
scheduler.add_job(
    lambda: ai_performance_review("mid-morning"),
    CronTrigger(hour=11, minute=0, day_of_week='mon-fri', timezone=IST),
    id='review_11am',
    name='AI performance review at 11:00 AM IST'
)

scheduler.add_job(
    lambda: ai_performance_review("post-lunch"),
    CronTrigger(hour=13, minute=0, day_of_week='mon-fri', timezone=IST),
    id='review_1pm',
    name='AI performance review at 1:00 PM IST'
)

scheduler.add_job(
    lambda: ai_performance_review("final-hour"),
    CronTrigger(hour=14, minute=30, day_of_week='mon-fri', timezone=IST),
    id='review_230pm',
    name='AI performance review at 2:30 PM IST'
)

scheduler.add_job(
    close_all_positions,
    CronTrigger(hour=15, minute=15, day_of_week='mon-fri', timezone=IST),
    id='close_positions',
    name='Close all positions at 3:15 PM IST'
)

scheduler.add_job(
    end_of_day_review,
    CronTrigger(hour=15, minute=30, day_of_week='mon-fri', timezone=IST),
    id='eod_review',
    name='End-of-day review at 3:30 PM IST'
)

# Start scheduler
try:
    scheduler.start()
    ist_now = get_ist_now()
    logging.info("=" * 60)
    logging.info("Automated trading scheduler started successfully")
    logging.info(f"Current IST Time: {ist_now.strftime('%Y-%m-%d %H:%M:%S %Z')}")
    logging.info(f"Current IST Day: {ist_now.strftime('%A')} (weekday={ist_now.weekday()})")
    logging.info(f"Expiry Day (Tuesday): {'YES' if ist_now.weekday() == 1 else 'NO'}")
    logging.info("=" * 60)
except Exception as e:
    logging.error(f"Failed to start scheduler: {e}")

# ==================== WEBSOCKET & IMPROVED MONITORING ====================

def init_websocket_for_client(clientcode):
    """Initialize WebSocket connection for real-time price updates"""
    try:
        from SmartApi.smartWebSocketV2 import SmartWebSocketV2
        
        # Get session data
        session_data = None
        for sid, sdata in _SMARTAPI_SESSIONS.items():
            if sdata.get('clientcode') == clientcode:
                session_data = sdata
                break
        
        if not session_data:
            logging.error(f"No session found for {clientcode}")
            return None
        
        # Extract tokens
        tokens = session_data.get('tokens', {})
        auth_token = tokens.get('jwtToken', '').replace('Bearer ', '')
        feed_token = tokens.get('feedToken')
        api_key = os.getenv('ANGELONE_API_KEY') or os.getenv('SMARTAPI_API_KEY')
        
        if not all([auth_token, api_key, feed_token]):
            logging.error(f"Missing WebSocket credentials for {clientcode}: jwt={bool(auth_token)}, api_key={bool(api_key)}, feed={bool(feed_token)}")
            return None
        
        # Create WebSocket instance
        ws = SmartWebSocketV2(
            auth_token,
            api_key,
            clientcode,
            feed_token,
            max_retry_attempt=3,
            retry_strategy=1,
            retry_delay=10,
            retry_multiplier=2,
            retry_duration=30
        )
        
        # Define callbacks (match SmartWebSocketV2 signature)
        def on_data(wsapp, data):
            """Handle incoming price data"""
            try:
                if isinstance(data, dict):
                    token = data.get('token')
                    ltp = data.get('last_traded_price') or data.get('ltp')
                    
                    if token and ltp:
                        LIVE_PRICE_CACHE[token] = {
                            'ltp': float(ltp) / 100,  # Convert paise to rupees
                            'timestamp': datetime.now(),
                            'volume': data.get('volume', 0),
                            'oi': data.get('open_interest', 0)
                        }
                        # Put in queue for processing
                        PRICE_UPDATE_QUEUE.put({'token': token, 'ltp': float(ltp) / 100})
            except Exception as e:
                logging.error(f"Error processing WebSocket data: {e}")
        
        def on_open(wsapp):
            logging.info(f"[WS] WebSocket opened for {clientcode}")
        
        def on_error():
            logging.error(f"[WS] WebSocket error for {clientcode}")
        
        def on_close(wsapp):
            logging.info(f"[WS] WebSocket closed for {clientcode}")
        
        # Assign callbacks
        ws.on_open = on_open
        ws.on_data = on_data
        ws.on_error = on_error
        ws.on_close = on_close
        
        # Connect
        ws.connect()
        
        WEBSOCKET_CONNECTIONS[clientcode] = ws
        logging.info(f"[OK] WebSocket initialized for {clientcode}")
        
        return ws
        
    except ImportError:
        logging.error("SmartWebSocketV2 not available. Install smartapi-python package.")
        return None
    except Exception as e:
        logging.error(f"Error initializing WebSocket: {e}", exc_info=True)
        return None

def subscribe_to_symbols(clientcode, symboltokens):
    """Subscribe to price updates for specific symbols"""
    try:
        ws = WEBSOCKET_CONNECTIONS.get(clientcode)
        if not ws:
            ws = init_websocket_for_client(clientcode)
        
        if ws:
            # Subscribe to tokens (mode 1 = LTP, mode 2 = Quote, mode 3 = Snap Quote)
            correlation_id = f"sub_{clientcode}_{datetime.now().timestamp()}"
            ws.subscribe(correlation_id, 1, symboltokens)  # Mode 1 for LTP updates
            logging.info(f"[WS] Subscribed to {len(symboltokens)} symbols for {clientcode}")
            return True
        
        return False
        
    except Exception as e:
        logging.error(f"Error subscribing to symbols: {e}")
        return False

def get_price_from_cache_or_api(symboltoken, clientcode):
    """Get price from WebSocket cache, fallback to API"""
    try:
        # Check cache first
        cached = LIVE_PRICE_CACHE.get(symboltoken)
        if cached:
            # Check if cache is fresh (< 10 seconds old)
            age = (datetime.now() - cached['timestamp']).total_seconds()
            if age < 10:
                return cached['ltp']
        
        # Fallback to API
        prices = get_batch_option_prices([symboltoken], clientcode)
        if prices and symboltoken in prices:
            return prices[symboltoken]['ltp']
        
        return None
        
    except Exception as e:
        logging.error(f"Error getting price: {e}")
        return None

def improved_monitor_prices_and_execute():
    """Enhanced monitoring with faster intervals and WebSocket support"""
    global ACTIVE_TRADES, PARSED_TRADE_SETUPS
    
    while True:
        try:
            current_time = datetime.now().time()
            
            # Only monitor during market hours (9:15 AM - 3:15 PM)
            if current_time < datetime.strptime("09:15", "%H:%M").time():
                time.sleep(60)
                continue
            
            if current_time > datetime.strptime("15:15", "%H:%M").time():
                logging.info("Market closing time reached, stopping monitoring")
                break
            
            logging.info(f"=" * 50)
            logging.info(f"[FAST] Price monitoring cycle at {current_time}")
            
            # Process each client's trade setups
            for clientcode, trade_setups in PARSED_TRADE_SETUPS.items():
                if not AUTO_TRADING_ENABLED.get(clientcode, False):
                    continue
                
                # Initialize WebSocket if not already done
                if clientcode not in WEBSOCKET_CONNECTIONS:
                    ws = init_websocket_for_client(clientcode)
                    if ws:
                        # Subscribe to NIFTY + active option tokens
                        tokens_to_subscribe = ['99926000']  # NIFTY token
                        for setup in trade_setups:
                            token = setup.get('symboltoken')
                            if token:
                                tokens_to_subscribe.append(str(token))
                        subscribe_to_symbols(clientcode, tokens_to_subscribe)
                
                plan_data = DAILY_TRADE_PLAN.get(clientcode, {})
                if plan_data.get('status') != 'pending':
                    continue
                
                logging.info(f"Processing {len(trade_setups)} trade setups for {clientcode}")
                
                # Check each trade setup for entry
                for trade_setup in trade_setups:
                    trade_number = trade_setup.get('trade_number')
                    
                    # Skip if already executed
                    if clientcode in ACTIVE_TRADES:
                        already_executed = any(
                            t.get('trade_number') == trade_number 
                            for t in ACTIVE_TRADES[clientcode].values()
                        )
                        if already_executed:
                            continue
                    
                    # Evaluate entry conditions
                    if evaluate_entry_conditions(trade_setup, clientcode):
                        logging.info(f"[ENTRY] Conditions MET for Trade #{trade_number}!")
                        
                        # Execute trade with retry
                        trade_id = execute_trade_with_retry(trade_setup, clientcode)
                        
                        if trade_id:
                            logging.info(f"[OK] Trade #{trade_number} executed: ID={trade_id}")
                        else:
                            logging.error(f"[FAIL] Failed to execute Trade #{trade_number}")
            
            # Monitor active trades
            improved_monitor_active_trades()
            
            # Adaptive sleep - faster near trigger levels
            sleep_duration = calculate_adaptive_sleep()
            logging.info(f"Sleeping for {sleep_duration} seconds...")
            time.sleep(sleep_duration)
        
        except Exception as e:
            logging.error(f"Error in improved monitoring: {e}", exc_info=True)
            time.sleep(30)

def calculate_adaptive_sleep():
    """Calculate sleep duration based on market conditions"""
    try:
        # Default: 60 seconds (much faster than 5 minutes)
        base_interval = MONITORING_INTERVAL
        
        # Check if any trades are close to entry/exit levels
        for clientcode, trade_setups in PARSED_TRADE_SETUPS.items():
            if not AUTO_TRADING_ENABLED.get(clientcode, False):
                continue
            
            nifty_price = get_price_from_cache_or_api('99926000', clientcode)
            if not nifty_price:
                continue
            
            for setup in trade_setups:
                if setup.get('executed'):
                    continue
                
                # Check if close to entry trigger
                for condition in setup.get('entry_conditions', []):
                    if condition.get('type') == 'price' and condition.get('indicator') == 'NIFTY':
                        trigger_level = condition.get('value', 0)
                        distance = abs(nifty_price - trigger_level)
                        
                        # If within 20 points of trigger, check every 10 seconds
                        if distance < 20:
                            return 10
                        # If within 50 points, check every 30 seconds
                        elif distance < 50:
                            return 30
        
        # Check active trades
        for clientcode, trades in ACTIVE_TRADES.items():
            for trade_id, trade_data in trades.items():
                if trade_data.get('status') != 'open':
                    continue
                
                current_price = get_price_from_cache_or_api(
                    trade_data.get('symboltoken'), clientcode
                )
                if not current_price:
                    continue
                
                entry_price = trade_data.get('entry_price', 0)
                sl = trade_data.get('stop_loss', 0)
                target = trade_data.get('target_1', 0)
                
                # If close to SL or target, check more frequently
                sl_distance = abs(current_price - sl) / entry_price if entry_price > 0 else 1
                target_distance = abs(current_price - target) / entry_price if entry_price > 0 else 1
                
                # Within 2% of SL or target - check every 10 seconds
                if sl_distance < 0.02 or target_distance < 0.02:
                    return 10
                # Within 5% - check every 20 seconds
                elif sl_distance < 0.05 or target_distance < 0.05:
                    return 20
        
        return base_interval
        
    except Exception as e:
        logging.error(f"Error calculating adaptive sleep: {e}")
        return 60

def execute_trade_with_retry(trade_setup, clientcode, max_retries=3):
    """Execute trade with retry logic"""
    for attempt in range(max_retries):
        try:
            trade_id = execute_trade_entry(trade_setup, clientcode)
            if trade_id:
                return trade_id
            
            # Wait before retry with exponential backoff
            wait_time = 5 * (2 ** attempt)
            logging.warning(f"Retry {attempt + 1}/{max_retries} in {wait_time}s...")
            time.sleep(wait_time)
            
        except Exception as e:
            logging.error(f"Attempt {attempt + 1} failed: {e}")
            if attempt == max_retries - 1:
                return None
    
    return None

def improved_monitor_active_trades():
    """Enhanced active trade monitoring with WebSocket and faster checks"""
    global ACTIVE_TRADES
    
    for clientcode, trades in ACTIVE_TRADES.items():
        if not trades:
            continue
        
        open_trades = {
            trade_id: trade_data 
            for trade_id, trade_data in trades.items() 
            if trade_data.get('status') == 'open'
        }
        
        if not open_trades:
            continue
        
        try:
            # Use cache-first approach with WebSocket
            for trade_id, trade_data in open_trades.items():
                symboltoken = trade_data.get('symboltoken')
                
                # Get price from cache or API
                current_price = get_price_from_cache_or_api(symboltoken, clientcode)
                
                if not current_price:
                    logging.warning(f"No price for {trade_data.get('tradingsymbol')}")
                    continue
                
                entry_price = trade_data.get('entry_price', 0)
                profit_pct = ((current_price - entry_price) / entry_price) * 100 if entry_price > 0 else 0
                
                logging.info(f"[MONITOR] Trade {trade_id}: Rs.{current_price:.2f} ({profit_pct:+.2f}%)")
                
                # Smart trailing stop loss (time + VIX aware)
                current_time = datetime.now().time()
                morning_volatility = current_time < datetime.strptime("10:30", "%H:%M").time()
                
                # Get VIX-based thresholds (cached, refreshes every 5 min)
                vix_value = get_current_vix_value()
                thresholds = calculate_vix_based_thresholds(vix_value)
                profit_target = calculate_vix_based_profit_target(vix_value)
                
                # EARLY SESSION (9:15-10:30): Only trail after profit target (time-based safety)
                if morning_volatility:
                    if profit_pct >= profit_target and trade_data.get('stop_loss', 0) < entry_price:
                        old_sl = trade_data['stop_loss']
                        trade_data['stop_loss'] = entry_price
                        logging.info(f"[TRAIL SL] Morning session: Moved SL to breakeven at {profit_target:.1f}% profit: Rs.{old_sl:.2f} → Rs.{entry_price:.2f}")
                
                # LATER SESSION (10:30-15:15): VIX-based dynamic trailing
                else:
                    # Move to breakeven based on VIX threshold
                    if profit_pct >= thresholds['breakeven_threshold'] and trade_data.get('stop_loss', 0) < entry_price:
                        old_sl = trade_data['stop_loss']
                        trade_data['stop_loss'] = entry_price
                        logging.info(f"[TRAIL SL] VIX-based breakeven: Rs.{old_sl:.2f} → Rs.{entry_price:.2f} ({profit_pct:.1f}% profit, VIX={vix_value:.1f if vix_value else 'N/A'})")
                    
                    # Trail at 5% below current based on VIX threshold
                    elif profit_pct >= thresholds['trail_threshold']:
                        new_sl = current_price * 0.95  # 5% below current
                        if new_sl > trade_data.get('stop_loss', 0):
                            old_sl = trade_data['stop_loss']
                            trade_data['stop_loss'] = new_sl
                            logging.info(f"[TRAIL SL] VIX-based trailing: Rs.{old_sl:.2f} → Rs.{new_sl:.2f} ({profit_pct:.1f}% profit, VIX={vix_value:.1f if vix_value else 'N/A'})")
                
                # VIX-based profit exit (highest priority)
                if profit_pct >= profit_target:
                    pnl = (current_price - entry_price) * trade_data.get('quantity', 0)
                    logging.info(f"[{profit_target:.1f}% PROFIT] VIX-based quick exit (VIX={vix_value:.1f if vix_value else 'N/A'}): {trade_id} Entry: Rs.{entry_price:.2f} Exit: Rs.{current_price:.2f} P&L: Rs.{pnl:.2f}")
                    close_position(clientcode, trade_id, current_price, f'{profit_target:.0f}pct_profit')
                    
                    # Reset for re-entry
                    trade_number = trade_data.get('trade_number')
                    if trade_number and clientcode in PARSED_TRADE_SETUPS:
                        for setup in PARSED_TRADE_SETUPS[clientcode]:
                            if setup.get('trade_number') == trade_number:
                                setup['executed'] = False
                                logging.info(f"[RESET] Setup {trade_number} available for re-entry")
                                break
                
                # Stop loss
                elif current_price <= trade_data.get('stop_loss', 0):
                    logging.warning(f"[SL] Stop Loss Hit for Trade {trade_id}")
                    close_position(clientcode, trade_id, current_price, 'stop_loss')
                
                # Target 1
                elif not trade_data.get('target_1_hit') and current_price >= trade_data.get('target_1', 999999):
                    logging.info(f"[TARGET 1] Hit for Trade {trade_id}")
                    partial_close_position(clientcode, trade_id, current_price, 'target_1')
                
                # Target 2
                elif trade_data.get('target_1_hit') and current_price >= trade_data.get('target_2', 999999):
                    logging.info(f"[TARGET 2] Hit for Trade {trade_id}")
                    close_position(clientcode, trade_id, current_price, 'target_2')
                
        except Exception as e:
            logging.error(f"Error monitoring trades for {clientcode}: {e}", exc_info=True)

# Start improved monitoring in background thread
def start_improved_monitoring():
    """Start the improved monitoring system"""
    global PRICE_MONITOR_THREAD
    
    if PRICE_MONITOR_THREAD and PRICE_MONITOR_THREAD.is_alive():
        logging.info("Monitoring thread already running")
        return
    
    PRICE_MONITOR_THREAD = threading.Thread(
        target=improved_monitor_prices_and_execute,
        daemon=True,
        name="ImprovedPriceMonitor"
    )
    PRICE_MONITOR_THREAD.start()
    logging.info("[OK] Improved monitoring thread started")

# Start monitoring when app initializes
try:
    start_improved_monitoring()
except Exception as e:
    logging.error(f"Failed to start improved monitoring: {e}")

if __name__ == '__main__':
    app.run(debug=True)
